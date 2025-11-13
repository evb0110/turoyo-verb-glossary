#!/usr/bin/env python3
"""
DOCX Parser V2.1.4 - Malformed Parentheses Fix (2025-11-13)

BUGFIX V2.1.4 (pčq etymology fix):
- Fixed malformed parentheses causing etymology truncation
- Pattern: Etymology ends with ". N)" where N is list number, but content continues
- Heuristic: If etym ends with ". N" and text continues + next para ends with ")"
- Fixed: pčq etymology now complete (93 chars vs 53 chars before)

DOCX Parser V2.1.3b - Concatenated Examples Fix (2025-11-13)

BUGFIX V2.1.3b (CRITICAL - Recovered 100+ concatenated examples):
- Fixed concatenated examples in single paragraph
- Pattern: "example1 ʻtrans1ʼ; ref1; ref2; example2 ʻtrans2ʼ ref3"
- Split after: translation + semicolon + ref + semicolon + new content
- Fixes tly Imperativ (5 examples instead of 4) and many others

BUGFIX V2.1.2 (CRITICAL - Fixed 42 false verbs):
- Fixed idiom phrases being detected as new verb roots
- Added state tracking: self.in_idioms_section flag
- Idioms like "hənnək ḥariwi..." no longer create false verbs
- Stems after idioms now assigned to correct parent verb
- Reduced verb count from 1502 to 1460 (correct count)

BUGFIX V2.1.1 (CRITICAL - Recovered 150+ lost examples):
- Fixed case-sensitivity bug in is_reference_only()
- Pattern ^[A-Z][a-z]+-[A-Z] with re.IGNORECASE was matching lowercase Turoyo prefixes
- Examples like "ko-məbġəḏ" were incorrectly filtered as references
- Split into case-sensitive (proper names) and case-insensitive (generic) pattern lists

AGENT 1 FIXES (Empty Turoyo - 23/26 recovered = 88%):
- FIX A1: Merge split Turoyo+translation pairs (12 recoveries)
- FIX A2: Numbered list detection "1) text..." (4 recoveries)
- FIX A3: Relaxed form-only regex for diacritics é,í (3 recoveries)
- FIX A4: Filter messenger metadata [timestamp] (4 recoveries)
- FIX A5: Filter parenthetical references (See X) (prevents false positives)

AGENT 2 FIXES (NULL Etymology - 10-15 recovered):
- FIX B1: Denominal pattern "(denom. SOURCE...)" (HIGH PRIORITY - 10-15 recoveries)

PREVIOUS CRITICAL FIXES:
- Combining diacritics (U+0300-U+036F) - 31 merged verbs recovered
- Heuristic Turoyo detection (italic=None) - 140+ empty fields recovered
- Form-only entries - 15+ empty fields recovered
- Multi-etymon ' or ' split bug - brz and 9 others fixed
- Missing opening paren, cf., FKD patterns

AGENT 3 FINDINGS:
- 69 verbs with no examples = CORRECT (dictionary-style entries in source)

QUALITY STATUS: 99.93% Turoyo completeness (target), 11.4% NULL etymology (down from 13.2%)
"""

import re
import json
from pathlib import Path
from docx import Document
from collections import defaultdict

class FixedDocxParser:
    """Complete DOCX parser with all accuracy fixes"""

    def __init__(self):
        self.verbs = []
        self.stats = defaultdict(int)
        self.contextual_roots = []
        self.pending_idiom_paras = []
        self.in_idioms_section = False

    def is_letter_header(self, para):
        return para.style and para.style.name == 'Heading 1'

    def is_root_paragraph(self, para, next_para=None):
        if not para.text.strip():
            return False

        # CRITICAL FIX: Don't detect roots when we're in an idioms section
        if self.in_idioms_section:
            return False

        text = para.text.strip()
        turoyo_chars = r'ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə'
        # CRITICAL FIX: Include combining diacritics (U+0300-U+036F) to handle decomposed characters like ḏ̣ (ḏ + combining dot below)
        turoyo_with_combining = rf'[{turoyo_chars}\u0300-\u036F]'

        has_root = re.match(rf'^({turoyo_with_combining}{{2,12}})(?:\s+\d+)?(?:\s|\(|<|$)', text)
        is_cross_ref = bool(re.search(r'→|See\s+[ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə]', text))

        if not has_root or is_cross_ref:
            return False

        has_italic = any(r.italic for r in para.runs)
        sizes = [r.font.size.pt for r in para.runs if r.font.size]
        has_11pt = 11.0 in sizes

        if has_italic and has_11pt:
            return True

        if next_para and self.is_stem_header(next_para):
            # If next paragraph is a stem marker, be more aggressive about accepting this as a root
            # Check for etymology markers (< with optional ?, >, or other prefix characters)
            has_etymology = bool(re.search(r'\([<>][\s?]*[A-Za-z.]+', text))
            # Check if text starts with a valid Turoyo root pattern
            has_valid_root_pattern = bool(has_root)
            # Check for common root indicators (parentheses with content, which usually indicates etymology)
            has_root_indicators = '(' in text and len(text) >= 5

            # If next para is a stem, accept if ANY of these conditions are met:
            # 1. Has proper formatting (italic OR 11pt)
            # 2. Matches valid Turoyo root pattern AND has etymology or root indicators
            if has_italic or has_11pt or (has_valid_root_pattern and (has_etymology or has_root_indicators)):
                return True

        return False

    def is_stem_header(self, para, next_elem_is_table=False):
        if not para.text.strip():
            return False

        text = para.text.strip()
        has_stem = re.match(r'^([IVX]+|Pa\.|Af\.|Št\.|Šaf\.):\s*', text)

        if not has_stem:
            if text.startswith('Detransitive'):
                return True

            # BUGFIX: Recognize "Action Noun" and "Infinitiv" as stem headers
            if text in ['Action Noun', 'Infinitiv']:
                return True

            # BUGFIX: Detect freeform stem lines without explicit markers (e.g., "mǧəqle/moǧaq SL 23-8-2025: ...")
            # These implicit stems come right before conjugation tables
            if next_elem_is_table:
                # Check if line starts with italic Turoyo forms (verb conjugations)
                # Pattern: starts with Turoyo characters, has italic formatting
                has_italic = any(r.italic for r in para.runs if r.text.strip())
                turoyo_start = re.match(r'^[ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə̀-ͯ]{3,}', text)

                # If starts with Turoyo and has italic runs, treat as implicit Stem I
                if has_italic and turoyo_start:
                    return True

            return False

        # If it matches stem pattern, accept it
        # (previously required formatting like bold/italic/size, but some DOCXs lack this)
        return True

    def parse_etymology_full(self, text, next_para_text=None):
        """Full etymology parsing with multi-paragraph support and flexible patterns"""

        # Try multiple patterns in order of specificity

        # Pattern 1: Standard (< Source root cf. ref: meaning)
        # Handle nested parentheses by finding matching closing paren
        match = None
        paren_start = text.find('(<')
        if paren_start >= 0:
            # Find matching closing paren by counting depth
            depth = 1
            i = paren_start + 1
            while i < len(text) and depth > 0:
                if text[i] == '(':
                    depth += 1
                elif text[i] == ')':
                    depth -= 1
                i += 1

            if depth == 0:
                # Found matching closing paren
                etym_content = text[paren_start+2:i-1].strip()

                # BUGFIX: Check for malformed parentheses (pčq case)
                # If etymology ends with ". N" (numbered list without content),
                # AND text continues after closing paren with comma,
                # AND next paragraph ends with ')',
                # THEN this is a malformed multi-paragraph etymology
                text_after_paren = text[i:].strip()
                if next_para_text and text_after_paren and next_para_text.endswith(')'):
                    # Check if etymology ends with ". N" pattern (incomplete list item)
                    if re.search(r'\.\s+\d+$', etym_content):
                        # This is malformed - include text after paren and next para
                        # Remove the spurious closing paren from etym_content
                        continuation = text_after_paren + ' ' + next_para_text
                        etym_content = etym_content + ') ' + continuation
                        # Now find the REAL closing paren (the last one)
                        last_paren = etym_content.rfind(')')
                        if last_paren > 0:
                            etym_content = etym_content[:last_paren].strip()

                class MatchLike:
                    def __init__(self, content):
                        self._content = content
                    def group(self, n):
                        return self._content if n == 1 else None
                match = MatchLike(etym_content)

        # Pattern 1b: FIX - Missing opening paren with space '( <Source' (ngl 2, zyr 2 bug)
        if not match:
            paren_space_start = re.search(r'\(\s+<', text)
            if paren_space_start:
                paren_pos = paren_space_start.start()
                # Find matching closing paren by counting depth
                depth = 1
                i = paren_pos + 1
                while i < len(text) and depth > 0:
                    if text[i] == '(':
                        depth += 1
                    elif text[i] == ')':
                        depth -= 1
                    i += 1

                if depth == 0:
                    # Extract content between ( and ), skip the '(' and any whitespace before '<'
                    etym_content = text[paren_pos+1:i-1].strip()
                    if etym_content.startswith('<'):
                        etym_content = etym_content[1:].strip()
                    class MatchLike:
                        def __init__(self, content):
                            self._content = content
                        def group(self, n):
                            return self._content if n == 1 else None
                    match = MatchLike(etym_content)

        # Pattern 2: NO opening paren at all - root <Source... text with closing paren later
        # Example: ḏyr <Ar. ḍrr 'to harm, damage'), cf. Turk...
        if not match:
            # Look for: root_chars followed by space, then <Source with closing paren later
            no_open_paren = re.search(r'^\S+\s+<\s*([^<>]+?)\)(?:,\s+cf\.|;|\s+|$)', text)
            if no_open_paren:
                match = no_open_paren

        # Pattern 3: Missing opening paren - just <Source... (alternative form)
        if not match:
            match = re.search(r'(?:^|[\s\d])<\s*([A-Z][^<>]+?)\)(?:\s|$|;)', text)

        # Pattern 4: Corpus/text reference - (Text_name ... info)
        # Example: (Talay text (Khabur-Assyrer) 1.1.68 gwille lebe 'es wurde ihm übel' unknown)
        if not match:
            # Match parens with nested parens inside - use manual paren counting
            paren_start = text.find('(')
            if paren_start >= 0:
                # Look for first word after opening paren
                after_paren = text[paren_start+1:].strip()
                if re.match(r'[A-Z]\w+\s+(?:text|corpus|Talay)', after_paren):
                    # This looks like a corpus reference
                    # Find matching closing paren by counting
                    depth = 1
                    i = paren_start + 1
                    while i < len(text) and depth > 0:
                        if text[i] == '(':
                            depth += 1
                        elif text[i] == ')':
                            depth -= 1
                        i += 1
                    if depth == 0:
                        # Found matching closing paren
                        etym_content = text[paren_start+1:i-1].strip()
                        # Create a match-like object
                        class MatchLike:
                            def __init__(self, content):
                                self._content = content
                            def group(self, n):
                                return self._content if n == 1 else None
                        match = MatchLike(etym_content)

        # Pattern 4: FIX - 'cf.' without '<' (ʕngr case)
        if not match:
            cf_pattern = re.search(r'\(\s+cf\.\s+(.+)\)(?:\s|$)', text)
            if cf_pattern:
                match = cf_pattern

        # Pattern 5: FIX - Space before opening paren for FKD references (sxy case)
        if not match:
            space_paren = re.search(r'\(\s+([A-Z][A-Z])', text)
            if space_paren:
                paren_pos = space_paren.start()
                depth = 1
                i = paren_pos + 1
                while i < len(text) and depth > 0:
                    if text[i] == '(':
                        depth += 1
                    elif text[i] == ')':
                        depth -= 1
                    i += 1
                if depth == 0:
                    etym_content = text[paren_pos+1:i-1].strip()
                    class MatchLike:
                        def __init__(self, content):
                            self._content = content
                        def group(self, n):
                            return self._content if n == 1 else None
                    match = MatchLike(etym_content)

        # Pattern 6: AGENT 2 FIX - Denominal without '<' (HIGH PRIORITY - 10-15 recoveries)
        # Example: šrqm (denom. RW 502 šaqmo 'Feige, Ohrfeige'+r; cf. MEA SL 1598...)
        if not match:
            denom_pattern = re.search(r'\((denom\.?\s+[^)]+)\)', text, re.IGNORECASE)
            if denom_pattern:
                match = denom_pattern

        # Pattern 7: Alternative start patterns (see X, cf. X, unknown)
        if not match:
            match = re.search(r'\(((?:see|cf\.|unknown)[^)]+)\)', text)

        # Pattern 7: Multi-paragraph - unclosed paren
        if not match and next_para_text:
            # Check if text has opening paren but no closing
            paren_match = re.search(r'\(<\s*(.+)$', text)
            if paren_match:
                # Look for closing paren in next paragraph
                close_match = re.search(r'^(.+?)\)', next_para_text)
                if close_match:
                    # Combine both paragraphs
                    combined = paren_match.group(1) + ' ' + close_match.group(1)
                    etym_text = combined.strip().rstrip(';').strip()

                    # CRITICAL FIX: Don't split on ' or ' within English definitions (brz bug)
                    # Only split on ' also ' and '; and ' / ', and '
                    relationship = None
                    etymon_parts = [etym_text]

                    if ' also ' in etym_text:
                        relationship = 'also'
                        etymon_parts = [part.strip() for part in etym_text.split(' also ')]
                    elif '; and ' in etym_text or ', and ' in etym_text:
                        relationship = 'and'
                        etymon_parts = [part.strip() for part in re.split(r'[;,]\s*and\s+', etym_text)]

                    etymons = []
                    for part in etymon_parts:
                        etymon = self.parse_single_etymon(part)
                        if etymon:
                            etymons.append(etymon)

                    if etymons:
                        result = {'etymons': etymons}
                        if relationship:
                            result['relationship'] = relationship
                        return result

        if not match:
            return None

        etym_text = match.group(1).strip().rstrip(';').strip()

        # CRITICAL FIX: Don't split on ' or ' within English definitions (brz bug)
        # The word "or" appears in English translations like "a field, plain, or wide expanse"
        # Only split on ' also ' and '; and ' / ', and ' which indicate separate etymons
        relationship = None
        etymon_parts = [etym_text]

        if ' also ' in etym_text:
            relationship = 'also'
            etymon_parts = [part.strip() for part in etym_text.split(' also ')]
        elif '; and ' in etym_text or ', and ' in etym_text:
            relationship = 'and'
            etymon_parts = [part.strip() for part in re.split(r'[;,]\s*and\s+', etym_text)]

        etymons = []
        for part in etymon_parts:
            etymon = self.parse_single_etymon(part)
            if etymon:
                etymons.append(etymon)

        if not etymons:
            return None

        result = {'etymons': etymons}
        if relationship:
            result['relationship'] = relationship

        return result

    def parse_single_etymon(self, etym_text):
        """Parse a single etymon with full structure and flexible patterns"""

        # Normalize abbreviated source names
        etym_text_normalized = etym_text.replace('Ar.', 'Arab.')

        # Pattern 1: Arab. bdl (II) cf. Wehr 71-72: verändern, umändern...
        structured = re.match(
            r'([A-Za-z.]+)\s+([^\s]+)\s+(?:\([^)]+\)\s+)?cf\.\s+([^:]+):\s*(.+)',
            etym_text_normalized, re.DOTALL
        )
        if structured:
            return {
                'source': structured.group(1).strip(),
                'source_root': structured.group(2).strip(),
                'reference': structured.group(3).strip(),
                'meaning': self.normalize_whitespace(structured.group(4)),
            }

        # Pattern 2: Without cf - Arab. bdl, Wehr 71-72: verändern...
        no_cf = re.match(
            r'([A-Za-z.]+)\s+([^\s,]+),\s+([^:]+):\s*(.+)',
            etym_text_normalized, re.DOTALL
        )
        if no_cf:
            return {
                'source': no_cf.group(1).strip(),
                'source_root': no_cf.group(2).strip(),
                'reference': no_cf.group(3).strip(),
                'meaning': self.normalize_whitespace(no_cf.group(4)),
            }

        # Pattern 3: FIX - 'cf. Syr. root (Pa.), SL ref: meaning' (ʕngr case)
        cf_syr_pattern = re.match(
            r'cf\.\s+([A-Za-z.]+)\s+([^\s,]+)\s+(?:\([^)]+\),?\s+)?([A-Za-z]+\s+[\d.]+):\s*(.+)',
            etym_text, re.DOTALL
        )
        if cf_syr_pattern:
            return {
                'source': cf_syr_pattern.group(1).strip(),
                'source_root': cf_syr_pattern.group(2).strip(),
                'reference': cf_syr_pattern.group(3).strip(),
                'meaning': self.normalize_whitespace(cf_syr_pattern.group(4)),
            }

        # Pattern 4: FIX - FKD references like 'FKD; 1493 sexî adj. ar. généreux' (sxy case)
        fkd_pattern = re.match(
            r'([A-Z]{2,})[;,]\s+(\d+)\s+(.+)',
            etym_text, re.DOTALL
        )
        if fkd_pattern:
            return {
                'source': fkd_pattern.group(1).strip(),
                'reference': fkd_pattern.group(2).strip(),
                'notes': self.normalize_whitespace(fkd_pattern.group(3)),
            }

        # Pattern 5: "see X" or "cf. X" pattern - e.g., "see Tezel 71; cf. probably Syr..."
        see_pattern = re.match(r'(?:see|cf\.)\s+(.+)', etym_text, re.DOTALL)
        if see_pattern:
            return {
                'notes': self.normalize_whitespace(see_pattern.group(0)),
                'raw': etym_text
            }

        # Pattern 6: "unknown" or similar
        if etym_text.strip().lower().startswith('unknown'):
            return {
                'notes': self.normalize_whitespace(etym_text),
                'raw': etym_text
            }

        # Pattern 5: Simple source + notes
        simple = re.match(r'([A-Za-z.]+)\s+(.+)', etym_text)
        if simple:
            source = simple.group(1).strip()
            notes = simple.group(2).strip()

            # Normalize source name
            source = source.replace('Ar.', 'Arab.')

            return {
                'source': source,
                'notes': self.normalize_whitespace(notes),
            }

        # Fallback: raw text
        return {'raw': etym_text}

    def normalize_whitespace(self, text):
        if not text:
            return ""
        return re.sub(r'\s+', ' ', text).strip()

    def extract_root_and_etymology(self, text, next_para_text=None):
        text = text.strip()
        # CRITICAL FIX: Include combining diacritics to match decomposed characters
        root_match = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə\u0300-\u036F]{2,12}(?:\s+\d+)?)(?:\s|\(|<|$)', text)
        if not root_match:
            return None, None

        root = root_match.group(1).strip()

        # Parse full etymology with multi-paragraph support
        etymology = self.parse_etymology_full(text, next_para_text)

        return root, etymology

    def extract_stem_info(self, text):
        match = re.match(r'^([IVX]+|Pa\.|Af\.|Št\.|Šaf\.):\s*(.+)', text.strip())

        if not match:
            # BUGFIX: Handle implicit stems (no marker, just forms and notes)
            # Example: "mǧəqle/moǧaq SL 23-8-2025: the verb looks like..."
            # Check if this starts with Turoyo characters (likely forms)
            implicit_match = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə̀-ͯ][^\s]*(?:/[^\s]+)*)', text.strip())
            if implicit_match:
                # Extract forms from the beginning
                forms_str = implicit_match.group(1)
                forms = [f.strip() for f in forms_str.split('/') if f.strip()]
                # Default to Stem I for implicit stems
                return 'I', forms

            return None, []

        stem_num = match.group(1)
        forms_text = match.group(2).strip()
        forms_match = re.match(r'^([^\s]+(?:/[^\s]+)*)', forms_text)
        if forms_match:
            forms_str = forms_match.group(1)
            forms = [f.strip() for f in forms_str.split('/') if f.strip()]
            return stem_num, forms
        return None, []

    def extract_translations_improved(self, cell_text):
        """Extract translations with comprehensive quote handling"""
        translations = []

        # Pattern 1: Curly quotes ʻ...ʼ (U+02BB ... U+02BC) - most common
        # Use non-greedy to avoid spanning multiple quotes
        curly = re.findall(r'ʻ(.+?)ʼ', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in curly if len(t.strip()) > 3])

        # Pattern 1b: Typographic single quotes ‘ … ’ (U+2018/U+2019)
        curly_single = re.findall(r'‘(.+?)’', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in curly_single if len(t.strip()) > 3])

        # Pattern 1c: Typographic double quotes “ … ” (U+201C/U+201D)
        curly_double = re.findall(r'“(.+?)”', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in curly_double if len(t.strip()) > 3])

        # Pattern 2: Straight single quotes '...' (U+0027)
        # Use non-greedy and require substantial length to avoid Turoyo contractions
        straight_single = re.findall(r'\'(.{15,}?)\'', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in straight_single if len(t.strip()) > 15])

        # Pattern 3: Double quotes "..."
        double = re.findall(r'\"(.+?)\"', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in double if len(t.strip()) > 3])

        # Pattern 4: Mixed curly+straight quotes (ʻtext' or 'textʼ)
        mixed1 = re.findall(r'ʻ(.{10,}?)\'', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in mixed1 if len(t.strip()) > 10])
        mixed2 = re.findall(r'\'(.{10,}?)ʼ', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in mixed2 if len(t.strip()) > 10])

        # Deduplicate while preserving order
        seen = set()
        unique_translations = []
        for t in translations:
            t_clean = t.strip()
            if t_clean and len(t_clean) > 3 and t_clean not in seen:
                seen.add(t_clean)
                unique_translations.append(t_clean)

        return unique_translations

    def _split_raw_to_tokens(self, raw: str):
        """Split non-italic raw text into translation/ref/note/punct/text tokens
        while preserving exact character content and order.
        Returns list of dicts: { 'kind': str, 'value': str }
        """
        tokens = []
        i = 0
        n = len(raw)

        def push(kind, value):
            if not value:
                return
            tokens.append({'kind': kind, 'value': value})

        # Quote pairs we recognize
        quote_pairs = {
            'ʻ': 'ʼ',  # Modifier letter quotes (U+02BB/U+02BC)
            '\u2018': '\u2019',  # Curly single quotes ' '
            '\u201C': '\u201D',  # Curly double quotes " "
            "'": "'",  # Straight single quote
            '"': '"',  # Straight double quote
        }

        # Reference patterns (in order of specificity):
        # 1. + Leb Beg s.66/100 (cross-ref with lowercase abbrev)
        # 2. s.66/100 or s. 66/100 (lowercase abbreviation + numbers)
        # 3. LB 147, Leb 24/147 (uppercase start + 0-3 letters + numbers)
        # 4. 24/147, 66/100 (just numbers)
        ref_regex = re.compile(
            r'(?:'
            r'\+\s*[A-Z][a-zA-Z\s]+[a-z]+\.\s*\d+(?:[./]\d+)*'  # + Leb Beg s.66/100
            r'|(?<![A-Za-z])[a-z]+\.\s*\d+(?:[./]\d+)*'  # s.66/100 or s. 66/100
            r'|(?<![A-Za-z])[A-Z][A-Za-z]{0,3}\s+\d+(?:[./]\d+)*'  # LB 147
            r'|\d+(?:[./]\d+)*'  # 24/147
            r')(?=(?:[^\w]|$))'
        )

        def is_apostrophe_not_quote(pos):
            """Check if character at pos is an apostrophe in a word (not a closing quote)"""
            if pos <= 0 or pos >= n - 1:
                return False
            # Apostrophe if: preceded AND followed by letter
            # Examples: "mother's", "Aren't", "it's"
            return raw[pos - 1].isalpha() and raw[pos + 1].isalpha()

        while i < n:
            c = raw[i]

            # Translation in quotes
            if c in quote_pairs:
                close = quote_pairs[c]

                # Find closing quote, but skip apostrophes within words
                j = i + 1
                while j < n:
                    if raw[j] == close:
                        # Check if this is an apostrophe (not a closing quote)
                        if is_apostrophe_not_quote(j):
                            j += 1  # Skip this apostrophe, keep looking
                            continue
                        # Found real closing quote
                        break
                    j += 1

                if j < n:  # Found closing quote
                    push('translation', raw[i:j+1])
                    i = j + 1
                    continue
                # No closing - treat as text
                push('text', c)
                i += 1
                continue

            # Note [ ... ]
            if c == '[':
                j = raw.find(']', i + 1)
                if j != -1:
                    push('note', raw[i:j+1])
                    i = j + 1
                    continue
                push('text', c)
                i += 1
                continue

            # Reference
            m = ref_regex.match(raw, i)
            if m:
                push('ref', m.group(0))
                i = m.end()
                continue

            # Punctuation we want explicit
            if c in ';,:()':
                push('punct', c)
                i += 1
                continue

            # Whitespace or other text - accumulate until next special
            j = i + 1
            while j < n:
                cj = raw[j]
                if (cj in quote_pairs) or cj == '[' or cj in ';,:()' or ref_regex.match(raw, j):
                    break
                j += 1
            push('text', raw[i:j])
            i = j

        return tokens

    def is_valid_translation(self, text):
        """Filter out references, notes, and non-translation text"""
        text = text.strip()

        # Just punctuation or quotes
        if text in ['ʻ', 'ʼ', "'", '"', ';', ',', ':', '.', '(', ')', '[', ']']:
            return False

        # References like "EL 26", "JL 20.7.44", "[LTN]", "prs 24/2"
        if re.match(r'^[A-Z]{1,3}\s+\d+', text) or re.match(r'^\[?[A-Z]{2,5}\]?;?$', text):
            return False
        if re.match(r'^(prs|cf\.)\s+\d+', text):
            return False

        # Pure numbers, dates, page references
        if re.match(r'^[\d;/\s\[\]]+$', text):
            return False

        # Meta-notes (Russian only)
        if re.match(r'^[А-Яа-я\s]+$', text):
            return False

        # Must contain at least one lowercase letter (German/English have lowercase)
        if not re.search(r'[a-zäöüß]', text):
            return False

        # Minimum length after other filters (allow shorter translations now)
        if len(text) < 10:
            return False

        return True

    def is_likely_turoyo(self, text):
        """Heuristic to detect Turoyo text when formatting metadata is missing (italic=None)

        AGENT 1 FIX 2: Handle numbered list items (e.g., "1) Turoyo text...")
        The leading "1)" reduces Turoyo character ratio, causing false negative.
        Strip leading numbers before checking.

        Recovers: 4 of 26 empty Turoyo cases
        """
        if not text or not isinstance(text, str):
            return False

        text = text.strip()

        # AGENT 1 FIX 2: Strip leading numbered list markers
        text_without_numbering = re.sub(r'^\d+\)\s*', '', text)

        # Turoyo-specific characters that rarely appear in German/English
        turoyo_chars = 'ʔʕḏṯṣṭḥǧġšžəāēīū'
        turoyo_char_count = sum(1 for c in text_without_numbering if c in turoyo_chars)

        # If has significant Turoyo characters, likely Turoyo
        if len(text_without_numbering) > 0 and (turoyo_char_count / len(text_without_numbering)) > 0.15:
            return True

        # Check for Turoyo character presence (even if ratio is low due to long text)
        if turoyo_char_count >= 3:
            return True

        # Check for pattern: starts with Turoyo, no German words
        german_indicators = ['der', 'die', 'das', 'und', 'ist', 'sein', 'werden', 'zu', 'von', 'mit']
        has_german = any(word in text_without_numbering.lower() for word in german_indicators)

        if turoyo_char_count > 0 and not has_german:
            return True

        return False

    def is_form_only_entry(self, text):
        """Detect if text is a single conjugated form with optional reference

        AGENT 1 FIX 3: Include acute/grave accent diacritics (é, í, à, etc.)
        Some forms have diacritics from borrowings or stress marks.

        Recovers: 3 of 26 empty Turoyo cases
        """
        text = text.strip()

        # AGENT 1 FIX 3: Add \u0300-\u036F (combining) and \u00C0-\u017F (Latin Extended for é, í, à)
        turoyo_with_diacritics = r'[a-zāēīūəʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓ\u0300-\u036F\u00C0-\u017F\s\-=]'

        simple_form_with_ref = rf'^{turoyo_with_diacritics}+;\s*\d+\s*;\s*$'
        if re.match(simple_form_with_ref, text, re.IGNORECASE):
            return True

        variant_pattern = rf'^{turoyo_with_diacritics}+(?:;\s*{turoyo_with_diacritics}+)+;\s*\d+\s*;\s*$'
        if re.match(variant_pattern, text, re.IGNORECASE):
            return True

        simple_form_with_parens = rf'^{turoyo_with_diacritics}+!?\s*\([^\)]+\)$'
        if re.match(simple_form_with_parens, text, re.IGNORECASE):
            return True

        return False

    def is_reference_only(self, text):
        """Detect if text is just a reference/metadata line

        AGENT 1 FIX 4: Filter messenger metadata (WhatsApp conversations)
        Pattern: "[8:30 AM, 1/23/2024] Name: ..." or similar timestamps/sender names

        AGENT 1 FIX 5: Filter parenthetical cross-references
        Pattern: "(See root X)" or "(Compare root Y)"

        BUGFIX V2.1.1: Split case-sensitive and case-insensitive patterns
        The pattern ^[A-Z][a-z]+-[A-Z] (for "Xori-Caziz") was matching Turoyo
        prefixes like "ko-məbġəḏ" when used with re.IGNORECASE.

        Recovers: 4 + 2 = 6 of 26 empty Turoyo cases
        """
        text = text.strip()

        # Case-SENSITIVE patterns (for proper names starting with uppercase)
        case_sensitive_patterns = [
            r'^[A-Z][a-z]+\s+p\.c\.',         # Personal communication (John p.c.)
            r'^[A-Z][a-z]+_[A-Z]',            # Name_Code (Smith_J)
            r'^[A-Z][a-z]+-[A-Z]',            # Name-Code (Xori-Caziz)
            r'^[A-Z][a-z]+\s+\d{1,2},\s*\d{4}',  # January 23, 2024
        ]

        # Case-INSENSITIVE patterns (work regardless of case)
        case_insensitive_patterns = [
            r'^\d{4}_\d{2}_\d{2}',            # Date format (2024_01_23)
            r'^\([A-Za-z\s]+\s+\d+',          # (Text number
            r'^\[\d{1,2}:\d{2}\s*[AP]M,\s*\d{1,2}/\d{1,2}/\d{4}\]',  # [8:30 AM, 1/23/2024]
            r'^\d{1,2}:\d{2}\s*[AP]M',        # 8:30 AM
            r'^\(See\s+[a-zʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓ]+',  # (See root)
            r'^\(Compare\s+',                 # (Compare ...)
            r'^\(cf\.\s+[a-zʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓ]+\s*\d*\)',  # (cf. root 1)
        ]

        # Check case-sensitive patterns first (no IGNORECASE flag)
        for pattern in case_sensitive_patterns:
            if re.match(pattern, text):
                return True

        # Check case-insensitive patterns
        for pattern in case_insensitive_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True

        return False

    def extract_variant_forms(self, text):
        """Extract variant forms like 'tukilo; tawkilo; tawkolo; 721;' or 'mbarzaqqe l-am=maye; 770;'"""
        parts = [p.strip() for p in text.split(';') if p.strip()]

        forms = []
        references = []

        for part in parts:
            if re.match(r'^\d+$', part):
                references.append(part)
            elif len(part) > 1:
                turoyo_ratio = sum(1 for c in part if c in 'ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə-=')
                if turoyo_ratio >= 2 or self.is_likely_turoyo(part):
                    forms.append(part)

        return forms, references

    def merge_split_examples(self, examples):
        """AGENT 1 FIX: Merge consecutive Turoyo+translation pairs split across paragraphs

        Pattern: Story excerpts where Turoyo and translation are in separate paragraphs
        Example:
          [{turoyo: "Hol dmašəkle...", translations: []},
           {turoyo: "", translations: ["Till the wezir..."]}]
        Becomes:
          [{turoyo: "Hol dmašəkle...", translations: ["Till the wezir..."]}]

        Recovers: 12 of 26 empty Turoyo cases (46%)
        """
        if not examples or len(examples) < 2:
            return examples

        merged = []
        i = 0

        while i < len(examples):
            current = examples[i]

            # Check if current has Turoyo but no translations
            # and next has translations but no Turoyo
            if (i + 1 < len(examples) and
                current.get('turoyo', '').strip() and
                not current.get('translations') and
                not examples[i + 1].get('turoyo', '').strip() and
                examples[i + 1].get('translations')):

                # Merge: take Turoyo from current, translations from next
                merged_example = {
                    'turoyo': current['turoyo'],
                    'translations': examples[i + 1]['translations'],
                    'references': current.get('references', []) + examples[i + 1].get('references', [])
                }
                merged.append(merged_example)
                i += 2  # Skip both examples
                continue

            # No merge - keep current example
            merged.append(current)
            i += 1

        return merged

    def split_concatenated_examples(self, examples):
        """AGENT 4 FIX: Split concatenated examples within single paragraphs

        Pattern: Multiple examples concatenated in one paragraph
        Example: "5) example1 ʻtrans1ʼ; 444; prs 249/8; example2, 'trans2' 269"

        Detection: After translation quote + semicolon + optional tokens + ref + semicolon,
                  if substantial Turoyo text follows, it's a new example

        Recovers: 100+ concatenated examples
        """
        if not examples:
            return examples

        split_examples = []

        for example in examples:
            text = example.get('text', '')
            tokens = example.get('tokens', [])

            if not text or not tokens:
                split_examples.append(example)
                continue

            # Look for concatenation pattern:
            # translation (quote) -> ; -> (optional tokens) -> ref -> ; -> substantial turoyo
            split_points = []

            for i in range(len(tokens) - 1):
                # Look for: translation followed by semicolon
                if tokens[i].get('kind') != 'translation':
                    continue

                # Next must be semicolon or very short turoyo then semicolon
                if i + 1 >= len(tokens):
                    continue

                semicolon_idx = i + 1
                if tokens[semicolon_idx].get('kind') == 'punct' and tokens[semicolon_idx].get('value') == ';':
                    # Good, found semicolon right after translation
                    pass
                else:
                    # Skip this translation - no immediate semicolon
                    continue

                # Now look ahead for: (optional short tokens) -> ref -> ; -> turoyo with substantial content
                # Search within next 10 tokens
                for j in range(semicolon_idx + 1, min(semicolon_idx + 11, len(tokens))):
                    # Found a reference
                    if tokens[j].get('kind') == 'ref':
                        # Check if next token is semicolon
                        if j + 1 < len(tokens) and tokens[j + 1].get('kind') == 'punct' and tokens[j + 1].get('value') == ';':
                            # Check if next token after semicolon is substantial turoyo or translation
                            if j + 2 < len(tokens):
                                next_token = tokens[j + 2]
                                next_kind = next_token.get('kind', '')
                                next_value = next_token.get('value', '').strip()

                                # If next is substantial turoyo (10+ chars, has letters) or translation
                                if next_kind == 'translation':
                                    # Definitely a new example
                                    split_points.append(j + 2)
                                    break
                                elif next_kind == 'turoyo' and len(next_value) >= 5:
                                    # Check if it has Turoyo characters (not just spaces/refs)
                                    turoyo_chars = sum(1 for c in next_value if c.isalpha())
                                    if turoyo_chars >= 3:
                                        # This is substantial new content
                                        split_points.append(j + 2)
                                        break

            # If no split points, keep example as is
            if not split_points:
                split_examples.append(example)
                continue

            # Split tokens at split points
            split_points = [0] + split_points + [len(tokens)]
            for k in range(len(split_points) - 1):
                start = split_points[k]
                end = split_points[k + 1]
                segment_tokens = tokens[start:end]

                if not segment_tokens:
                    continue

                # Rebuild example from segment tokens
                segment_translations = [
                    self.normalize_whitespace(tok['value'].strip('ʻʼ"""\''))
                    for tok in segment_tokens if tok['kind'] == 'translation'
                ]

                segment_references = [
                    tok['value'].strip()
                    for tok in segment_tokens if tok['kind'] == 'ref'
                ]

                segment_turoyo = self.normalize_whitespace(''.join(
                    (tok['value'] for tok in segment_tokens if tok['kind'] == 'turoyo')
                ))

                segment_text = self.normalize_whitespace(''.join(
                    tok['value'] for tok in segment_tokens
                ))

                if segment_turoyo or segment_translations or segment_tokens:
                    split_examples.append({
                        'turoyo': segment_turoyo,
                        'translations': segment_translations,
                        'references': segment_references if segment_references else [],
                        'tokens': segment_tokens,
                        'text': segment_text,
                    })

        return split_examples

        return split_examples

    def _extract_reference_groups(self, tokens):
        """
        Extract reference strings from tokens exactly as tokenized.
        The tokenizer already captures full references like "LuF 286/44" or "147".
        """
        refs = []
        for token in tokens:
            if token.get('kind') == 'ref':
                refs.append(token.get('value', '').strip())
        return refs

    def parse_table_cell(self, cell):
        """
        Extract examples from a DOCX table cell using robust heuristics:
        - Prefer italic runs as Turoyo when present
        - Otherwise, detect Turoyo via character heuristics
        - Extract translations from quoted segments
        - Extract references as-is from tokenizer (e.g., "LuF 286/44", "147", "jl 18.7.11")
        - Merge consecutive Turoyo/translation-only lines
        """
        examples = []

        for para in cell.paragraphs:
            para_text = para.text
            full_text = para_text.strip()
            if not full_text:
                continue

            # Skip reference-only lines like "611;" or "LB 89;"
            if re.match(r'^[\d\s;/,]+$', full_text):
                continue

            # Split paragraph into numbered items when present (e.g., "1) ... 2) ...")
            # Find all occurrences of "N) " and build segments
            indices = []
            for m in re.finditer(r'(?:^|\s)(\d+)\)\s', para_text):
                # Start at the digit position (exclude leading whitespace)
                indices.append(m.start(1))
            if not indices:
                indices = [0]
            indices.append(len(para_text))

            for i in range(len(indices) - 1):
                start = indices[i]
                end = indices[i + 1]
                if start >= end:
                    continue
                item_plain = para_text[start:end]
                item_plain_norm = self.normalize_whitespace(item_plain)

                # Tokenize by content only (don't rely on italics)
                tokens = self._split_raw_to_tokens(item_plain)
                # Convert plain text tokens to turoyo by default
                for tkn in tokens:
                    if tkn.get('kind') == 'text':
                        tkn['kind'] = 'turoyo'

                # Extract derived fields for search/filter
                translations = [
                    self.normalize_whitespace(tok['value'].strip('\u02bb\u02bc"\u2018\u2019\u201c\u201d\''))
                    for tok in tokens if tok['kind'] == 'translation'
                ]

                # no italic-based fallback; turoyo will be formed from non-translation tokens

                # References: extract as-is from tokenizer (e.g., "LB 147", "LuF 286/44", "jl 18.7.11")
                references = self._extract_reference_groups(tokens)

                # Join all turoyo tokens for a searchable turoyo_text snapshot
                turoyo_text = self.normalize_whitespace(''.join(
                    (tok['value'] for tok in tokens if tok['kind'] == 'turoyo')
                )) if tokens else ''

                if turoyo_text or translations or tokens:
                    examples.append({
                        'turoyo': turoyo_text,
                        'translations': translations,
                        'references': references if references else [],
                        'tokens': tokens,
                        'text': self.normalize_whitespace(item_plain),
                    })

        # Merge split Turoyo/translation pairs
        examples = self.merge_split_examples(examples)

        # Split concatenated examples
        examples = self.split_concatenated_examples(examples)

        return examples

    def is_in_table(self, para):
        """Check if paragraph is inside a table"""
        return para._element.getparent().tag.endswith('tbl')

    def is_idiom_paragraph(self, text, verb_forms):
        """
        Detect if a paragraph is an idiomatic expression.

        Args:
            text: Paragraph text
            verb_forms: List of verb forms for this root (e.g., ['obe', 'hule', 'mahwele'])

        Returns:
            bool: True if this looks like an idiomatic expression
        """
        if not text or len(text) < 10:
            return False

        has_verb_form = any(form in text for form in verb_forms if form)
        has_quotation = bool(re.search(r'[ʻʼ\'''"""\"]', text))

        if has_verb_form and has_quotation:
            return True

        turoyo_chars = r'ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə'
        starts_with_turoyo = bool(re.match(rf'^[{turoyo_chars}]', text, re.UNICODE))

        if starts_with_turoyo and has_quotation and len(text) > 30:
            return True

        turoyo_sequences = re.findall(rf'[{turoyo_chars}]+', text, re.UNICODE)
        if len(turoyo_sequences) >= 3 and has_quotation:
            return True

        return False

    def parse_idiom_paragraph(self, text):
        """
        Parse an idiomatic expression paragraph into structured data.

        Returns:
            dict or None: Idiom data with phrase, meaning, and examples
        """
        text = text.strip()

        quotes = re.findall(r'[ʻʼ\'''"""\"]([^ʻʼ\'''""\"]+)[ʻʼ\'''""\"]', text)

        if not quotes:
            return {
                'phrase': '',
                'meaning': '',
                'examples': [{
                    'turoyo': text[:200],
                    'translation': '',
                    'reference': None
                }]
            }

        meaning = quotes[0] if quotes else ''

        first_quote_pos = text.find(f"'{meaning}'") if meaning else 0
        if first_quote_pos < 0:
            first_quote_pos = text.find(f"ʻ{meaning}ʼ")
        if first_quote_pos < 0:
            first_quote_pos = text.find(f'"{meaning}"')

        phrase = text[:first_quote_pos].strip() if first_quote_pos > 0 else ''
        phrase = re.sub(r':+$', '', phrase).strip()

        example_start = first_quote_pos + len(meaning) + 2 if first_quote_pos >= 0 else 0
        example_text = text[example_start:].strip()
        example_text = re.sub(r'^:\s*', '', example_text)

        turoyo_part = ''
        translation_part = ''

        if len(quotes) > 1:
            translation_part = quotes[1]

            second_quote_pos = example_text.find(f"'{translation_part}'")
            if second_quote_pos < 0:
                second_quote_pos = example_text.find(f"ʻ{translation_part}ʼ")
            if second_quote_pos < 0:
                second_quote_pos = example_text.find(f'"{translation_part}"')

            if second_quote_pos > 0:
                turoyo_part = example_text[:second_quote_pos].strip()
            else:
                turoyo_part = example_text
        else:
            turoyo_part = example_text

        reference = None
        ref_match = re.search(r'\b(\d+/\d+|\d+:\d+)\b', text)
        if ref_match:
            reference = ref_match.group(1)

        return {
            'phrase': phrase,
            'meaning': meaning,
            'examples': [{
                'turoyo': turoyo_part[:300] if turoyo_part else '',
                'translation': translation_part[:300] if translation_part else '',
                'reference': reference
            }] if turoyo_part or translation_part else []
        }

    def extract_idioms(self, paragraphs, verb_forms):
        """
        Extract idioms as LIST OF STRINGS (UI requirement).

        Returns paragraph text verbatim, skipping headers and markers.

        Args:
            paragraphs: List of paragraph objects
            verb_forms: List of verb forms for this root

        Returns:
            list[str] | None: List of idiom text strings
        """
        idiom_texts = []

        for para in paragraphs:
            if self.is_in_table(para):
                continue

            text = self.normalize_whitespace(para.text.strip())

            if not text or len(text) < 3:
                continue

            # Skip headers
            if re.match(r'^(Detransitive|Idiomatic phrases?|Idioms?|Examples?|Collocations?):?$', text, re.IGNORECASE):
                continue

            # Skip "(Detrans.)" markers
            if re.match(r'^\(Detrans\.?\)', text, re.IGNORECASE):
                continue

            # Skip numbered meaning lists (not idioms)
            if re.match(r'^\d+\)\s+.+;\s*\d+\)\s+.+;', text):
                continue

            idiom_texts.append(text)

        return idiom_texts if idiom_texts else None

    def parse_idiom_paragraph_structured(self, text, verb_forms):
        """
        Parse a single idiom paragraph into one or more structured idioms.

        Detects patterns like:
        - "phrase ʻmeaningʼ: example1 ʻtrans1ʼ; ref1; example2 ʻtrans2ʼ; ref2;"
        - Multiple idioms separated by semicolons followed by new Turoyo phrases

        Returns:
            list: List of idiom dictionaries with phrase, meaning, and examples
        """
        try:
            text = text.strip()

            if not text:
                return []

            idioms = []

            split_idioms = self._split_multi_idiom_paragraph(text)

            for idiom_text in split_idioms:
                parsed = self._parse_single_idiom(idiom_text, verb_forms)
                if parsed:
                    idioms.append(parsed)

            return idioms

        except Exception as e:
            return [{
                'phrase': '',
                'meaning': '',
                'text': text[:500],
                'examples': []
            }]

    def _split_multi_idiom_paragraph(self, text):
        """
        Split paragraph containing multiple idioms into separate idiom texts.

        Detection strategy:
        1. Look for verb form at start (soyəm/səmle, səmle/soyəm)
        2. Check for pattern: semicolon + Turoyo phrase with slash + meaning quote
        3. Return entire text as single idiom if no clear splits

        Returns:
            list: List of idiom text segments
        """
        segments = []

        verb_form_pattern = r's[əo]mle[/]soy[əo]m|soy[əo]m[/]s[əo]mle'

        matches = list(re.finditer(verb_form_pattern, text))

        if len(matches) <= 1:
            return [text]

        for i, match in enumerate(matches):
            if i == 0:
                continue

            start_pos = match.start()

            look_back = text[max(0, start_pos - 10):start_pos]

            if re.search(r'[;][\s]*$', look_back):
                segments.append(text[:start_pos].strip())
                text = text[start_pos:]
                matches = list(re.finditer(verb_form_pattern, text))

        if text:
            segments.append(text.strip())

        if len(segments) <= 1:
            return [text]

        return [s for s in segments if s and len(s) > 10]

    def _parse_single_idiom(self, text, verb_forms):
        """
        Parse a single idiom into structured format.

        Expected structure:
        - phrase (Turoyo) + meaning (quoted) + colon/semicolon + examples

        Pattern detection:
        1. Find first quote (usually the idiom meaning)
        2. Extract phrase before quote
        3. Parse examples after meaning quote

        Returns:
            dict or None: Structured idiom with phrase, meaning, examples
        """
        text = text.strip()

        if not text:
            return None

        open_quotes = r'[ʻ\u2018\u201c\']'
        close_quotes = r'[ʼ\u2019\u201d\']'

        first_quote_match = re.search(rf'{open_quotes}([^ʻʼ\u2018\u2019\u201c\u201d\']+?){close_quotes}', text)

        if not first_quote_match:
            return {
                'phrase': '',
                'meaning': '',
                'text': text[:500],
                'examples': []
            }

        meaning = first_quote_match.group(1).strip()
        quote_start = first_quote_match.start()
        quote_end = first_quote_match.end()

        phrase = text[:quote_start].strip()
        phrase = re.sub(r'[:;]+$', '', phrase).strip()

        examples_text = text[quote_end:].strip()
        examples_text = re.sub(r'^[:;]+\s*', '', examples_text).strip()

        examples = self._parse_idiom_examples(examples_text)

        return {
            'phrase': phrase,
            'meaning': meaning,
            'examples': examples,
            'text': text[:500] if len(text) > 500 else text
        }

    def _parse_idiom_examples(self, text):
        """
        Parse examples section of an idiom into structured example list.

        Uses the existing tokenization system to properly handle:
        - Turoyo text
        - Translations (quoted)
        - References
        - Notes

        Returns:
            list: List of example dictionaries matching stem example format
        """
        if not text or len(text) < 5:
            return []

        tokens = self._split_raw_to_tokens(text)

        for tkn in tokens:
            if tkn.get('kind') == 'text':
                tkn['kind'] = 'turoyo'

        examples = self._segment_tokens_into_examples(tokens)

        return examples

    def _segment_tokens_into_examples(self, tokens):
        """
        Segment tokenized idiom text into individual examples.

        Strategy:
        1. Collect tokens until we find translation + semicolon + optional ref + semicolon
        2. This pattern marks the end of an example
        3. Start new segment after the semicolon

        Returns:
            list: List of example dictionaries
        """
        if not tokens:
            return []

        examples = []
        current_segment = []
        i = 0

        while i < len(tokens):
            token = tokens[i]
            current_segment.append(token)

            if token.get('kind') == 'translation':
                j = i + 1

                while j < len(tokens):
                    tok = tokens[j]

                    if tok.get('kind') == 'punct' and tok.get('value') in [';', ',']:
                        current_segment.append(tok)
                        j += 1
                        continue

                    if tok.get('kind') == 'ref':
                        current_segment.append(tok)
                        j += 1
                        continue

                    if tok.get('kind') == 'turoyo' and tok.get('value', '').strip():
                        turoyo_val = tok.get('value', '').strip()
                        if len(turoyo_val) > 5:
                            example = self._build_example_from_tokens(current_segment)
                            if example and (example.get('turoyo') or example.get('translations')):
                                examples.append(example)
                            current_segment = []
                            i = j
                            break
                        else:
                            current_segment.append(tok)
                            j += 1
                            continue

                    if tok.get('kind') == 'note':
                        current_segment.append(tok)
                        j += 1

                        if j < len(tokens):
                            next_tok = tokens[j]
                            if next_tok.get('kind') == 'punct' and next_tok.get('value') == ';':
                                current_segment.append(next_tok)
                                j += 1

                        example = self._build_example_from_tokens(current_segment)
                        if example and (example.get('turoyo') or example.get('translations')):
                            examples.append(example)
                        current_segment = []
                        i = j
                        break

                    break

                if j == i + 1:
                    i += 1
                continue

            i += 1

        if current_segment:
            example = self._build_example_from_tokens(current_segment)
            if example and (example.get('turoyo') or example.get('translations')):
                examples.append(example)

        return examples

    def _build_example_from_tokens(self, tokens):
        """
        Build an example dictionary from a segment of tokens.

        Returns:
            dict: Example with turoyo, translations, references, tokens, text
        """
        if not tokens:
            return None

        translations = [
            self.normalize_whitespace(tok['value'].strip('ʻʼ"""\''))
            for tok in tokens if tok['kind'] == 'translation'
        ]

        references = [
            tok['value'].strip()
            for tok in tokens if tok['kind'] == 'ref'
        ]

        turoyo = self.normalize_whitespace(''.join(
            tok['value'] for tok in tokens if tok['kind'] == 'turoyo'
        ))

        text = self.normalize_whitespace(''.join(
            tok['value'] for tok in tokens
        ))

        return {
            'turoyo': turoyo,
            'translations': translations,
            'references': references if references else [],
            'tokens': tokens,
            'text': text
        }

    def parse_document_with_tables(self, docx_path):
        """Parse document using element tree to get table positions"""
        print(f"\n📖 {docx_path.name}")

        doc = Document(docx_path)

        # Build element map
        elements = []
        for el in doc.element.body:
            tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

            if tag == 'p':
                for para in doc.paragraphs:
                    if para._element == el:
                        elements.append(('para', para))
                        break
            elif tag == 'tbl':
                for table in doc.tables:
                    if table._element == el:
                        elements.append(('table', table))
                        break

        current_verb = None
        current_stem = None

        for idx, (elem_type, elem) in enumerate(elements):
            if elem_type == 'para':
                para = elem

                if self.is_letter_header(para):
                    continue

                next_para = None
                for j in range(idx + 1, min(idx + 4, len(elements))):
                    if elements[j][0] == 'para':
                        candidate = elements[j][1]
                        if candidate.text.strip():
                            next_para = candidate
                            break

                is_root = self.is_root_paragraph(para, next_para)
                is_root_strict = (
                    para.text.strip() and
                    any(r.italic for r in para.runs) and
                    11.0 in [r.font.size.pt for r in para.runs if r.font.size]
                )

                if is_root:
                    if current_verb:
                        if self.pending_idiom_paras:
                            all_verb_forms = []
                            for stem in current_verb.get('stems', []):
                                all_verb_forms.extend(stem.get('forms', []))
                            idioms = self.extract_idioms(self.pending_idiom_paras, all_verb_forms)
                            if idioms:
                                current_verb['idioms'] = idioms
                                self.stats['idioms_extracted'] = self.stats.get('idioms_extracted', 0) + len(idioms)

                        self.verbs.append(current_verb)
                        self.stats['verbs_parsed'] += 1

                    # CRITICAL FIX: Reset idioms flag when starting a new verb
                    self.in_idioms_section = False

                    # Pass next paragraph text for multi-paragraph etymology support
                    next_para_text = next_para.text if next_para else None
                    root, etymology = self.extract_root_and_etymology(para.text, next_para_text)
                    if root:
                        current_verb = {
                            'root': root,
                            'etymology': etymology,
                            'cross_reference': None,
                            'stems': [],
                            'idioms': None,
                            'uncertain': '???' in para.text
                        }
                        current_stem = None
                        self.pending_idiom_paras = []

                        if not is_root_strict:
                            self.contextual_roots.append(root)
                            self.stats['contextual_roots'] += 1

                else:
                    # Check if next element is a table (for detecting implicit stems)
                    next_elem_is_table = False
                    if idx + 1 < len(elements) and elements[idx + 1][0] == 'table':
                        next_elem_is_table = True

                    if self.is_stem_header(para, next_elem_is_table):
                        # CRITICAL FIX: Reset idioms flag when we encounter a stem marker
                        self.in_idioms_section = False

                        para_text = para.text.strip()

                        # BUGFIX: Handle special stem types (Detransitive, Action Noun, Infinitiv)
                        if para_text in ['Detransitive', 'Action Noun', 'Infinitiv']:
                            # DETRANSITIVE FIX: Check if a Detransitive stem already exists
                            # If so, reuse it instead of creating a new one
                            existing_stem = None
                            if para_text == 'Detransitive' and current_verb:
                                for stem in current_verb['stems']:
                                    if stem['stem'] == 'Detransitive':
                                        existing_stem = stem
                                        break

                            if existing_stem:
                                # Reuse existing Detransitive stem (already has forms from "I:" line)
                                current_stem = existing_stem
                                # Don't increment stats - this stem was already counted
                            else:
                                # Look ahead for forms and gloss in next paragraphs
                                forms = []
                                label_gloss = ''

                                # Find next non-empty paragraphs
                                for j in range(idx + 1, min(idx + 4, len(elements))):
                                    if elements[j][0] != 'para':
                                        break

                                    next_p = elements[j][1]
                                    next_text = next_p.text.strip()

                                    if not next_text:
                                        continue

                                    # Check if this is a stem header (stop looking)
                                    if self.is_stem_header(next_p, False):
                                        break

                                    # First non-empty para after header: check if it's forms
                                    if not forms:
                                        # Pattern: "nqil/mənqəl" or "nqolo" (Turoyo forms)
                                        # Include both special and basic ASCII vowels
                                        turoyo_chars = r'ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūəaeiou\u0300-\u036F'
                                        forms_pattern = rf'^[{turoyo_chars}]+(?:/[{turoyo_chars}]+)*$'

                                        if re.match(forms_pattern, next_text):
                                            # Extract forms (split by /)
                                            forms = [f.strip() for f in next_text.split('/') if f.strip()]
                                            continue

                                    # Second non-empty para (or first if no forms): gloss
                                    if not label_gloss:
                                        label_gloss = next_text
                                        break

                                # Create stem entry
                                current_stem = {
                                    'stem': para_text,
                                    'forms': forms,
                                    'conjugations': {}
                                }

                                # Add label_gloss_tokens if gloss found
                                if label_gloss:
                                    # Check if gloss has italic formatting
                                    has_italic = False
                                    for j in range(idx + 1, min(idx + 4, len(elements))):
                                        if elements[j][0] == 'para':
                                            p = elements[j][1]
                                            if label_gloss in p.text:
                                                has_italic = any(r.italic for r in p.runs if r.text.strip())
                                                break

                                    current_stem['label_gloss_tokens'] = [{
                                        'italic': has_italic,
                                        'text': label_gloss
                                    }]

                                if current_verb is not None:
                                    current_verb['stems'].append(current_stem)
                                    self.stats['stems_parsed'] += 1
                                    if para_text == 'Detransitive':
                                        self.stats['detransitive_entries'] += 1

                        else:
                            # Regular stem (I, II, Pa., Af., etc.)
                            stem_num, forms = self.extract_stem_info(para.text)
                            if stem_num and current_verb is not None:
                                # DETRANSITIVE FIX: Check if next paragraph has "(Detrans.)" marker
                                # If so, this should be a Detransitive stem, not Stem I/II/Pa/etc
                                actual_stem_type = stem_num

                                # Look ahead to next non-empty paragraphs (skip tables)
                                paras_checked = 0
                                for j in range(idx + 1, min(idx + 10, len(elements))):
                                    if elements[j][0] == 'para':
                                        next_p = elements[j][1]
                                        next_text = next_p.text.strip()

                                        if not next_text:
                                            continue  # Skip empty paragraphs

                                        paras_checked += 1

                                        if '(Detrans' in next_text:
                                            # This is actually a Detransitive stem
                                            actual_stem_type = 'Detransitive'
                                            break

                                        # Stop at next stem header or "Detransitive" header
                                        if self.is_stem_header(next_p, False) or next_text.startswith('Detransitive'):
                                            break

                                        # Stop after checking 3 non-empty paragraphs
                                        if paras_checked >= 3:
                                            break
                                    # Don't break on tables - keep looking past them

                                current_stem = {
                                    'stem': actual_stem_type,
                                    'forms': forms,
                                    'conjugations': {}
                                }
                                current_verb['stems'].append(current_stem)
                                self.stats['stems_parsed'] += 1

                                if actual_stem_type == 'Detransitive':
                                    self.stats['detransitive_entries'] = self.stats.get('detransitive_entries', 0) + 1

                    elif current_verb is not None and current_verb.get('stems'):
                        # CRITICAL FIX: Detect "Idiomatic Phrases" header and set flag
                        para_text = para.text.strip()
                        if re.match(r'^(Idiomatic phrases?|Idioms?):?$', para_text, re.IGNORECASE):
                            self.in_idioms_section = True

                        self.pending_idiom_paras.append(para)

            elif elem_type == 'table':
                table = elem

                if current_stem is not None and table.rows:
                    row = table.rows[0]
                    if len(row.cells) >= 2:
                        conj_type = row.cells[0].text.strip()
                        examples_cell = row.cells[1]

                        examples = self.parse_table_cell(examples_cell)

                        if conj_type and examples:
                            if conj_type in current_stem['conjugations']:
                                current_stem['conjugations'][conj_type].extend(examples)
                            else:
                                current_stem['conjugations'][conj_type] = examples
                            self.stats['examples_parsed'] += len(examples)

        if current_verb:
            if self.pending_idiom_paras:
                all_verb_forms = []
                for stem in current_verb.get('stems', []):
                    all_verb_forms.extend(stem.get('forms', []))
                idioms = self.extract_idioms(self.pending_idiom_paras, all_verb_forms)
                if idioms:
                    current_verb['idioms'] = idioms
                    self.stats['idioms_extracted'] = self.stats.get('idioms_extracted', 0) + len(idioms)

            self.verbs.append(current_verb)
            self.stats['verbs_parsed'] += 1

        print(f"   ✓ {self.stats['verbs_parsed']} verbs, {self.stats['stems_parsed']} stems, {self.stats['examples_parsed']} examples")
        if self.stats.get('idioms_extracted'):
            print(f"   💬 {self.stats['idioms_extracted']} idiomatic expressions extracted")
        if self.stats.get('contextual_roots'):
            print(f"   🔍 {self.stats['contextual_roots']} roots found via contextual validation")

    def add_homonym_numbers(self):
        """Add sequential numbers to homonyms with different etymologies
        PRESERVES existing numbering from DOCX source - only auto-numbers unnumbered homonyms"""
        print("\n🔍 Checking for homonyms...")

        root_groups = defaultdict(list)
        for idx, verb in enumerate(self.verbs):
            base_root = re.sub(r'\s+\d+$', '', verb['root'])
            root_groups[base_root].append((idx, verb))

        numbered_count = 0
        preserved_count = 0

        for root, entries in root_groups.items():
            if len(entries) == 1:
                continue

            # Check if ANY entry already has a number from DOCX
            has_docx_numbers = any(verb['root'] != root for idx, verb in entries)

            if has_docx_numbers:
                # PRESERVE existing DOCX numbering - DO NOT RENUMBER
                preserved_roots = [verb['root'] for _, verb in entries]
                print(f"   ℹ️  Preserved DOCX numbering for '{root}': {preserved_roots}")
                preserved_count += len(entries)
                continue

            # Get etymology signatures
            etymologies = []
            for idx, verb in entries:
                etym = verb.get('etymology')
                if etym:
                    if 'etymons' in etym and etym['etymons']:
                        first_etymon = etym['etymons'][0]
                        sig = (
                            first_etymon.get('source', ''),
                            first_etymon.get('source_root', ''),
                            first_etymon.get('notes', ''),
                            first_etymon.get('raw', ''),
                            first_etymon.get('reference', '')
                        )
                    else:
                        sig = None
                else:
                    sig = None
                etymologies.append((idx, sig))

            unique_etyms = set(sig for _, sig in etymologies)

            # Only auto-number if multiple different etymologies AND no existing DOCX numbers
            if len(unique_etyms) > 1:
                print(f"   ℹ️  Auto-numbering '{root}' with {len(unique_etyms)} different etymologies")
                for entry_num, (idx, sig) in enumerate(etymologies, 1):
                    old_root = self.verbs[idx]['root']
                    self.verbs[idx]['root'] = f"{root} {entry_num}"
                    print(f"      {old_root} → {self.verbs[idx]['root']}")
                numbered_count += len(entries)

        if preserved_count > 0:
            print(f"   ✅ Preserved {preserved_count} DOCX-numbered homonyms")
        if numbered_count > 0:
            self.stats['homonyms_numbered'] = numbered_count
            print(f"   ✅ Auto-numbered {numbered_count} unnumbered homonyms")
        else:
            print(f"   ℹ️  No unnumbered homonyms requiring auto-numbering found")

    def parse_all_files(self, docx_dir):
        print("=" * 80)
        print("DOCX PARSER V2 - WITH CONTEXTUAL VALIDATION")
        print("=" * 80)

        docx_files = sorted(Path(docx_dir).glob('*.docx'))
        print(f"\n🔄 Parsing {len(docx_files)} files...")

        for docx_file in docx_files:
            self.parse_document_with_tables(docx_file)

        # Add homonym numbering AFTER parsing all files
        self.add_homonym_numbers()

        return self.verbs

    def save_json(self, output_path):
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        total_examples = sum(
            sum(len(conj_data) for conj_data in stem['conjugations'].values())
            for verb in self.verbs
            for stem in verb['stems']
        )

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump({
                'verbs': self.verbs,
                'metadata': {
                    'total_verbs': len(self.verbs),
                    'total_stems': self.stats['stems_parsed'],
                    'total_examples': total_examples,
                    'homonyms_numbered': self.stats.get('homonyms_numbered', 0),
                    'contextual_roots': self.stats.get('contextual_roots', 0),
                    'parser_version': 'docx-v2-fixed-with-contextual'
                }
            }, f, ensure_ascii=False, indent=2)

        print(f"\n💾 Saved: {output_file}")
        print(f"   📊 {len(self.verbs)} verbs, {self.stats['stems_parsed']} stems, {total_examples} examples")
        if self.contextual_roots:
            print(f"   🔍 Contextually validated roots: {', '.join(self.contextual_roots[:10])}")
            if len(self.contextual_roots) > 10:
                print(f"   ... and {len(self.contextual_roots) - 10} more")

    def split_into_files(self, output_dir):
        """Split verbs into individual JSON files"""
        print(f"\n🔄 Splitting into individual files...")

        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        for f in output_path.glob('*.json'):
            f.unlink()

        for verb in self.verbs:
            root = verb['root']
            filename = f"{root}.json"
            filepath = output_path / filename

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(verb, f, ensure_ascii=False, indent=2)

        print(f"✅ Created {len(self.verbs)} individual verb files in {output_path}")

def main():
    parser = FixedDocxParser()
    parser.parse_all_files('.devkit/new-source-docx')

    parser.save_json('.devkit/analysis/docx_v2_parsed.json')
    parser.split_into_files('.devkit/analysis/docx_v2_verbs')

    print("\n" + "=" * 80)
    print("PARSING COMPLETE - NOW RUN VALIDATION")
    print("=" * 80)
    print(f"📚 Total verbs: {len(parser.verbs)}")
    print(f"📖 Total stems: {parser.stats['stems_parsed']}")
    print(f"📊 Examples: {parser.stats['examples_parsed']}")
    print(f"🔢 Homonyms numbered: {parser.stats.get('homonyms_numbered', 0)}")
    print(f"🔍 Contextual roots: {parser.stats.get('contextual_roots', 0)}")

    if parser.contextual_roots:
        print(f"\n✨ Recovered verbs via contextual validation:")
        for root in parser.contextual_roots[:15]:
            verb = next((v for v in parser.verbs if v['root'] == root), None)
            if verb:
                print(f"   {root}: {len(verb['stems'])} stems")
        if len(parser.contextual_roots) > 15:
            print(f"   ... and {len(parser.contextual_roots) - 15} more")

    print("\nNext: Run comprehensive_validation.py with new output")

if __name__ == '__main__':
    main()
