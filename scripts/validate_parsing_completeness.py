#!/usr/bin/env python3
"""
Brute-force parsing validator.

Extracts ALL text from DOCX files and verifies it appears in parsed JSON files.
This ensures no data is lost during parsing, even if formatting might be imperfect.

Strategy:
1. Extract every text fragment from DOCX (tables, paragraphs, everything)
2. Extract every text fragment from JSON (all fields recursively)
3. Normalize and tokenize both sets
4. Report coverage and missing data

Usage:
    python3 scripts/validate_parsing_completeness.py
"""

import json
import re
import unicodedata
from pathlib import Path
from collections import Counter, defaultdict
from typing import Set, List, Dict, Any
from docx import Document


class TextExtractor:
    """Extracts and normalizes text for comparison."""

    @staticmethod
    def normalize(text: str) -> str:
        """Normalize text for comparison."""
        if not text:
            return ""

        # Unicode normalization (NFC - composed form)
        text = unicodedata.normalize('NFC', str(text))

        # Normalize whitespace
        text = ' '.join(text.split())

        return text

    @staticmethod
    def tokenize(text: str) -> List[str]:
        """Split text into words, preserving Unicode characters and diacritics."""
        # Extract words with Unicode letters, combining diacritics, numbers
        # This preserves Turoyo characters like ḥ, ṭ, ḏ̣, etc.
        words = re.findall(r'[\w\u0300-\u036F]+', text, re.UNICODE)

        # Filter out single characters and normalize
        return [TextExtractor.normalize(w) for w in words if len(w) > 0]

    @staticmethod
    def extract_meaningful_chunks(text: str, min_length: int = 3) -> Set[str]:
        """Extract meaningful text chunks (3+ characters)."""
        chunks = set()

        # Split by common delimiters
        parts = re.split(r'[,;.!?\n\r\t]+', text)

        for part in parts:
            normalized = TextExtractor.normalize(part)
            if len(normalized) >= min_length:
                chunks.add(normalized)

        return chunks


class DocxTextExtractor:
    """Extracts ALL text from DOCX files."""

    def __init__(self, docx_dir: Path):
        self.docx_dir = docx_dir
        self.all_text: List[str] = []
        self.file_texts: Dict[str, List[str]] = defaultdict(list)

    def extract_all(self) -> List[str]:
        """Extract text from all DOCX files."""
        docx_files = sorted(self.docx_dir.glob('*.docx'))

        if not docx_files:
            raise FileNotFoundError(f"No DOCX files found in {self.docx_dir}")

        print(f"Extracting from {len(docx_files)} DOCX files...")

        for docx_file in docx_files:
            if docx_file.name.startswith('~$'):
                continue

            print(f"  - {docx_file.name}")
            texts = self._extract_from_file(docx_file)
            self.file_texts[docx_file.name] = texts
            self.all_text.extend(texts)

        return self.all_text

    def _extract_from_file(self, docx_file: Path) -> List[str]:
        """Extract all text from a single DOCX file."""
        doc = Document(docx_file)
        texts = []

        # Extract from paragraphs (headers, body text)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)

        # Extract from tables (this is where most verb data lives)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Extract from all paragraphs in cell
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            texts.append(text)

        return texts


class JsonTextExtractor:
    """Extracts ALL text from parsed JSON verb files."""

    def __init__(self, json_dir: Path):
        self.json_dir = json_dir
        self.all_text: List[str] = []
        self.verb_texts: Dict[str, List[str]] = defaultdict(list)

    def extract_all(self) -> List[str]:
        """Extract text from all JSON files."""
        json_files = sorted(self.json_dir.glob('*.json'))

        if not json_files:
            raise FileNotFoundError(f"No JSON files found in {self.json_dir}")

        print(f"Extracting from {len(json_files)} JSON files...")

        for json_file in json_files:
            with open(json_file, 'r', encoding='utf-8') as f:
                verb = json.load(f)

            texts = self._extract_from_verb(verb)
            self.verb_texts[json_file.stem] = texts
            self.all_text.extend(texts)

        return self.all_text

    def _extract_from_verb(self, verb: Dict[str, Any]) -> List[str]:
        """Recursively extract all text from a verb object."""
        texts = []

        def extract_recursive(obj):
            """Recursively extract strings from nested structures."""
            if isinstance(obj, str):
                if obj.strip():
                    texts.append(obj.strip())
            elif isinstance(obj, dict):
                for value in obj.values():
                    extract_recursive(value)
            elif isinstance(obj, list):
                for item in obj:
                    extract_recursive(item)
            elif obj is not None:
                # Convert numbers, booleans to string
                texts.append(str(obj))

        extract_recursive(verb)
        return texts


class ParsingValidator:
    """Validates that all DOCX text appears in JSON output."""

    def __init__(self, docx_dir: str, json_dir: str):
        self.docx_extractor = DocxTextExtractor(Path(docx_dir))
        self.json_extractor = JsonTextExtractor(Path(json_dir))

        self.docx_words = Counter()
        self.json_words = Counter()
        self.docx_chunks: Set[str] = set()
        self.json_chunks: Set[str] = set()

    def validate(self) -> Dict[str, Any]:
        """Run complete validation."""
        print()
        print("=" * 70)
        print("BRUTE-FORCE PARSING VALIDATOR")
        print("=" * 70)
        print()
        print("Strategy:")
        print("  1. Extract ALL text from DOCX files (tables, paragraphs, everything)")
        print("  2. Extract ALL text from JSON files (all fields, recursively)")
        print("  3. Tokenize and compare word frequencies")
        print("  4. Identify missing text and calculate coverage")
        print()
        print("=" * 70)
        print()

        # Extract all text
        docx_texts = self.docx_extractor.extract_all()
        json_texts = self.json_extractor.extract_all()

        print()
        print(f"Extracted {len(docx_texts):,} text fragments from DOCX")
        print(f"Extracted {len(json_texts):,} text fragments from JSON")
        print()

        # Process text
        print("Tokenizing and analyzing...")
        for text in docx_texts:
            self.docx_words.update(TextExtractor.tokenize(text))
            self.docx_chunks.update(TextExtractor.extract_meaningful_chunks(text))

        for text in json_texts:
            self.json_words.update(TextExtractor.tokenize(text))
            self.json_chunks.update(TextExtractor.extract_meaningful_chunks(text))

        print()
        print("=" * 70)
        print("ANALYSIS RESULTS")
        print("=" * 70)
        print()

        # Word-level analysis
        total_docx_words = sum(self.docx_words.values())
        total_json_words = sum(self.json_words.values())
        unique_docx_words = len(self.docx_words)
        unique_json_words = len(self.json_words)

        print(f"Word Statistics:")
        print(f"  DOCX total words:   {total_docx_words:,}")
        print(f"  JSON total words:   {total_json_words:,}")
        print(f"  DOCX unique words:  {unique_docx_words:,}")
        print(f"  JSON unique words:  {unique_json_words:,}")
        print()

        # Find missing words
        missing_words = set(self.docx_words.keys()) - set(self.json_words.keys())
        missing_word_occurrences = sum(self.docx_words[w] for w in missing_words)

        # Calculate word coverage
        word_coverage = ((total_docx_words - missing_word_occurrences) / total_docx_words * 100) if total_docx_words > 0 else 0

        print(f"Word Coverage:")
        print(f"  Missing unique words:       {len(missing_words):,}")
        print(f"  Missing word occurrences:   {missing_word_occurrences:,}")
        print(f"  Coverage:                   {word_coverage:.2f}%")
        print()

        # Chunk-level analysis
        missing_chunks = self.docx_chunks - self.json_chunks
        chunk_coverage = ((len(self.docx_chunks) - len(missing_chunks)) / len(self.docx_chunks) * 100) if self.docx_chunks else 0

        print(f"Text Chunk Statistics:")
        print(f"  DOCX unique chunks: {len(self.docx_chunks):,}")
        print(f"  JSON unique chunks: {len(self.json_chunks):,}")
        print(f"  Missing chunks:     {len(missing_chunks):,}")
        print(f"  Coverage:           {chunk_coverage:.2f}%")
        print()

        # Show missing data
        if missing_words:
            print("=" * 70)
            print("MISSING WORDS (Top 50 by frequency)")
            print("=" * 70)
            sorted_missing = sorted(missing_words, key=lambda w: self.docx_words[w], reverse=True)[:50]
            for word in sorted_missing:
                count = self.docx_words[word]
                print(f"  {word:30s} ({count:3d}x)")
            print()

        if missing_chunks:
            print("=" * 70)
            print("MISSING TEXT CHUNKS (First 30)")
            print("=" * 70)
            for chunk in sorted(missing_chunks)[:30]:
                preview = chunk[:80] + "..." if len(chunk) > 80 else chunk
                print(f"  - {preview}")
            print()

        # Categorize missing words
        categories = self._categorize_missing_words(missing_words)
        if categories:
            print("=" * 70)
            print("MISSING WORDS BY CATEGORY")
            print("=" * 70)
            for category, words in sorted(categories.items()):
                print(f"\n{category}:")
                for word in sorted(words)[:10]:
                    print(f"  - {word} ({self.docx_words[word]}x)")
            print()

        # Final verdict
        print("=" * 70)
        print("VERDICT")
        print("=" * 70)
        print()

        threshold = 99.0
        if word_coverage >= threshold:
            print(f"✅ PASS: Word coverage {word_coverage:.2f}% >= {threshold}%")
            status = "PASS"
        else:
            print(f"❌ FAIL: Word coverage {word_coverage:.2f}% < {threshold}%")
            status = "FAIL"

        print()
        print(f"Summary:")
        print(f"  - {len(self.docx_extractor.file_texts)} DOCX files processed")
        print(f"  - {len(self.json_extractor.verb_texts)} JSON files processed")
        print(f"  - Word coverage: {word_coverage:.2f}%")
        print(f"  - Chunk coverage: {chunk_coverage:.2f}%")
        print()

        return {
            'status': status,
            'word_coverage': word_coverage,
            'chunk_coverage': chunk_coverage,
            'missing_words': missing_words,
            'missing_chunks': missing_chunks,
            'total_docx_words': total_docx_words,
            'total_json_words': total_json_words,
            'missing_word_occurrences': missing_word_occurrences,
        }

    def _categorize_missing_words(self, missing_words: Set[str]) -> Dict[str, Set[str]]:
        """Categorize missing words by type."""
        categories = defaultdict(set)

        for word in missing_words:
            # Check word characteristics
            if word.isdigit():
                categories['Numbers'].add(word)
            elif len(word) <= 2:
                categories['Short words (≤2 chars)'].add(word)
            elif any(c in word for c in ['(', ')', '[', ']', '{', '}']):
                categories['Parentheses/Brackets'].add(word)
            elif word.isupper():
                categories['Uppercase'].add(word)
            elif any(c.isdigit() for c in word):
                categories['Alphanumeric'].add(word)
            else:
                # Check if it's likely Turoyo, German, or reference
                if any(c in word for c in 'ḥṭḏṣšʕʔġ'):
                    categories['Turoyo text'].add(word)
                elif any(c in word for c in 'äöüß'):
                    categories['German text'].add(word)
                else:
                    categories['Other'].add(word)

        return dict(categories)


def main():
    """Run validation."""
    validator = ParsingValidator(
        docx_dir='.devkit/new-source-docx',
        json_dir='.devkit/analysis/docx_v2_verbs'
    )

    try:
        results = validator.validate()

        # Save results to file
        output_file = Path('.devkit/analysis/parsing_validation_report.json')
        output_file.parent.mkdir(parents=True, exist_ok=True)

        # Convert sets to lists for JSON serialization
        json_results = {
            'status': results['status'],
            'word_coverage': results['word_coverage'],
            'chunk_coverage': results['chunk_coverage'],
            'total_docx_words': results['total_docx_words'],
            'total_json_words': results['total_json_words'],
            'missing_word_occurrences': results['missing_word_occurrences'],
            'missing_words_count': len(results['missing_words']),
            'missing_chunks_count': len(results['missing_chunks']),
            'missing_words_sample': sorted(results['missing_words'])[:100],
            'missing_chunks_sample': sorted(results['missing_chunks'])[:50],
        }

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(json_results, f, indent=2, ensure_ascii=False)

        print(f"Detailed report saved to: {output_file}")
        print()

        # Exit with appropriate code
        exit(0 if results['status'] == 'PASS' else 1)

    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        exit(2)


if __name__ == '__main__':
    main()
