#!/usr/bin/env python3
"""
INVESTIGATION: Find other verbs with the same malformed parentheses pattern
"""

from docx import Document
import re
from pathlib import Path

print("=" * 80)
print("FINDING VERBS WITH MALFORMED PARENTHESES")
print("=" * 80)

docx_files = sorted(Path('.devkit/new-source-docx').glob('*.docx'))

total_cases = 0
cases_by_file = {}

for docx_file in docx_files:
    print(f"\n📖 Checking: {docx_file.name}")
    doc = Document(docx_file)

    cases = []

    for i in range(len(doc.paragraphs) - 1):
        para = doc.paragraphs[i]
        next_para = doc.paragraphs[i + 1]

        text = para.text.strip()
        next_text = next_para.text.strip()

        # Pattern: Paragraph with "(<" and ends with ". N) text,"
        # AND next paragraph ends with ")"
        if '(<' in text and next_text.endswith(')'):
            # Check if text has balanced parens but ends with ". N)"
            paren_start = text.find('(<')
            if paren_start >= 0:
                depth = 1
                i_paren = paren_start + 1
                while i_paren < len(text) and depth > 0:
                    if text[i_paren] == '(':
                        depth += 1
                    elif text[i_paren] == ')':
                        depth -= 1
                    i_paren += 1

                if depth == 0:  # Balanced parens
                    etym_content = text[paren_start+2:i_paren-1].strip()
                    text_after = text[i_paren:].strip()

                    # Check if etymology ends with ". N" pattern
                    if re.search(r'\.\s+\d+$', etym_content) and text_after:
                        root_match = re.match(r'^([a-zʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə\u0300-\u036F]{2,12})', text)
                        root = root_match.group(1) if root_match else '???'

                        cases.append({
                            'root': root,
                            'para_idx': i,
                            'text': text[:100] + '...' if len(text) > 100 else text,
                            'next': next_text[:60] + '...' if len(next_text) > 60 else next_text
                        })

    if cases:
        cases_by_file[docx_file.name] = cases
        total_cases += len(cases)
        print(f"   ⚠️  Found {len(cases)} cases:")
        for case in cases[:5]:  # Show first 5
            print(f"      - {case['root']}: {case['text'][:60]}...")
        if len(cases) > 5:
            print(f"      ... and {len(cases) - 5} more")

print(f"\n" + "=" * 80)
print(f"SUMMARY: Found {total_cases} verbs with malformed parentheses")
print("=" * 80)

if total_cases > 0:
    print(f"\n📋 Affected verbs by file:")
    for filename, cases in cases_by_file.items():
        roots = [c['root'] for c in cases]
        print(f"   {filename}: {', '.join(roots[:10])}")
        if len(roots) > 10:
            print(f"      ... and {len(roots) - 10} more")
else:
    print(f"\n✅ No other cases found (pčq might be unique)")
