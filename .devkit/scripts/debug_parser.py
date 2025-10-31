#!/usr/bin/env python3
"""
Debug version of parser to trace pčq specifically
"""

import re
from pathlib import Path
from docx import Document

def is_root_paragraph(para):
    if not para.text.strip():
        return False
    has_italic = any(r.italic for r in para.runs)
    sizes = [r.font.size.pt for r in para.runs if r.font.size]
    has_11pt = 11.0 in sizes
    text = para.text.strip()
    turoyo_chars = r'ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə'
    has_root = re.match(f'^([{turoyo_chars}]{{2,6}})(\s+\d+)?', text)
    return has_italic and has_11pt and has_root

def is_stem_header(para):
    if not para.text.strip():
        return False
    has_bold = any(r.bold for r in para.runs)
    has_italic = any(r.italic for r in para.runs)
    sizes = [r.font.size.pt for r in para.runs if r.font.size]
    has_14pt = 14.0 in sizes
    has_stem = re.match(r'^([IVX]+):\s*', para.text.strip())
    return has_bold and has_italic and has_14pt and has_stem

doc = Document('.devkit/new-source-docx/4. l,m,n,p.docx')

# Build elements
elements = []
for el in doc.element.body:
    tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

    if tag == 'p':
        for para in doc.paragraphs:
            if para._element == el:
                elements.append(('para', para))
                break
    elif tag == 'tbl':
        for table in doc.tables:
            if table._element == el:
                elements.append(('table', table))
                break

current_verb = None
in_pcq = False

for i, (elem_type, elem) in enumerate(elements):
    if elem_type == 'para':
        para = elem
        text = para.text.strip()

        if text.startswith('pčq ('):
            in_pcq = True
            print(f'\n{i}: 🎯 pčq ROOT DETECTED')
            print(f'  is_root_paragraph: {is_root_paragraph(para)}')
            if is_root_paragraph(para):
                current_verb = {'root': 'pčq', 'stems': []}
                print(f'  ✓ current_verb created')

        elif in_pcq and current_verb:
            if is_stem_header(para):
                stem_match = re.match(r'^([IVX]+):\s*', text)
                stem_num = stem_match.group(1) if stem_match else None
                print(f'\n{i}: ✅ STEM HEADER DETECTED: {stem_num}')
                print(f'  Text: {text}')
                current_verb['stems'].append({'stem': stem_num})

            elif text.startswith(('I:', 'II:', 'III:', 'IV:', 'V:')):
                print(f'\n{i}: ⚠️  LOOKS LIKE STEM BUT NOT DETECTED: {text[:50]}')
                has_bold = any(r.bold for r in para.runs)
                has_italic = any(r.italic for r in para.runs)
                sizes = [r.font.size.pt for r in para.runs if r.font.size]
                has_14pt = 14.0 in sizes
                print(f'  has_bold: {has_bold}')
                print(f'  has_italic: {has_italic}')
                print(f'  has_14pt: {has_14pt}')
                print(f'  Runs:')
                for j, run in enumerate(para.runs[:3]):
                    print(f'    {j}: bold={run.bold}, italic={run.italic}, size={run.font.size.pt if run.font.size else None}, text="{run.text[:30]}"')

            elif 'Detransitive' in text:
                print(f'\n{i}: 🔵 DETRANSITIVE: {text}')
                current_verb['stems'].append({'stem': 'Detransitive'})

            next_root = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]{2,6})(\s+\d+|\s+\()', text)
            if next_root and next_root.group(1) != 'pčq':
                print(f'\n{i}: 🛑 NEXT VERB: {next_root.group(1)}')
                break

if current_verb:
    print(f'\n📊 Final result: {len(current_verb["stems"])} stems')
    for stem in current_verb['stems']:
        print(f'  - {stem["stem"]}')
