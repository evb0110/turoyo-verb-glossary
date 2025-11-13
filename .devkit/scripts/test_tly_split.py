#!/usr/bin/env python3
"""
Test concatenated examples fix on tly verb
"""

import sys
sys.path.insert(0, 'parser')

from docx import Document
from parse_docx_production import FixedDocxParser

# Create parser instance
parser = FixedDocxParser()

# Load the DOCX file
doc = Document('.devkit/new-source-docx/6. š,t,ṭ,ṯ.docx')

# Build element map
elements = []
for el in doc.element.body:
    tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

    if tag == 'p':
        for para in doc.paragraphs:
            if para._element == el:
                elements.append(('para', para))
                break
    elif tag == 'tbl':
        for table in doc.tables:
            if table._element == el:
                elements.append(('table', table))
                break

# Find tly in elements
tly_elem_index = None
for i, (elem_type, elem) in enumerate(elements):
    if elem_type == 'para' and elem.text.strip().startswith('tly ('):
        tly_elem_index = i
        break

if not tly_elem_index:
    print('ERROR: tly verb not found')
    exit(1)

print(f'Found tly at element index {tly_elem_index}')

# Find Imperativ table (should be a few elements after tly)
imperativ_table = None
for i in range(tly_elem_index + 1, min(tly_elem_index + 20, len(elements))):
    elem_type, elem = elements[i]
    if elem_type == 'table':
        row = elem.rows[0]
        if len(row.cells) >= 2:
            conj_type = row.cells[0].text.strip()
            if conj_type == 'Imperativ':
                imperativ_table = elem
                print(f'Found Imperativ table at element index {i}')
                break

if not imperativ_table:
    print('ERROR: Imperativ table not found')
    exit(1)

# Parse the Imperativ cell
row = imperativ_table.rows[0]
imperativ_cell = row.cells[1]

print('\n' + '='*80)
print('TESTING CONCATENATION FIX ON TLY IMPERATIV')
print('='*80)

examples = parser.parse_table_cell(imperativ_cell)

print(f'\n✅ Extracted {len(examples)} examples (expected: 5+)\n')

for idx, ex in enumerate(examples, 1):
    print(f'Example {idx}:')
    print(f'  Turoyo: {ex["turoyo"][:80]}{"..." if len(ex["turoyo"]) > 80 else ""}')
    print(f'  Translations: {ex["translations"][:1] if ex["translations"] else []}{"..." if len(ex["translations"]) > 1 else ""}')
    print(f'  References: {ex["references"]}')
    print()

if len(examples) >= 5:
    print('✅ SUCCESS: Found 5 or more examples!')
    print('   The concatenation fix is working correctly.')
else:
    print('⚠️  WARNING: Expected 5+ examples but found', len(examples))
    print('   The fix may need adjustment.')
