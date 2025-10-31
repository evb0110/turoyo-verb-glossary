#!/usr/bin/env python3
"""
Verify critical findings:
1. Does "th" really not exist? (3 stems missing)
2. Are the "false positives" really false? (ʕr, ʕw, ʕtr)
3. Check if there are variant spellings
"""

from pathlib import Path
from docx import Document
import re
import json

def search_for_variants(docx_path, base_root, variant_patterns):
    """Search for a root and its variants"""
    doc = Document(docx_path)

    results = []

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Check for each variant
        for variant in variant_patterns:
            pattern = f'^{re.escape(variant)}(\\s|\\(|$)'
            if re.match(pattern, text):
                # Get context
                context_after = doc.paragraphs[para_idx + 1].text[:100] if para_idx < len(doc.paragraphs) - 1 else ''

                results.append({
                    'variant': variant,
                    'paragraph': text[:150],
                    'context_after': context_after
                })
                break  # Don't match multiple variants for same paragraph

    return results

def check_original_json(root):
    """Check what the original HTML-based JSON has for this root"""
    json_path = Path(f'server/assets/verbs/{root}.json')

    if json_path.exists():
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return {
                'exists': True,
                'stems': len(data.get('stems', [])),
                'etymology': data.get('etymology'),
                'first_stem': data['stems'][0]['stem'] if data.get('stems') else None
            }
    else:
        return {'exists': False}

def main():
    print("="*80)
    print("VERIFYING CRITICAL FINDINGS")
    print("="*80)

    # 1. Check for "th" variants
    print("\n\n### 1. Searching for 'th' (3 stems) - HIGH VALUE")
    print("-" * 80)

    th_variants = ['th', 'thy', 'thw', 'thʕ', 'tho', 'tha']
    results = search_for_variants(
        Path('.devkit/new-source-docx/6. š,t,ṭ,ṯ.docx'),
        'th',
        th_variants
    )

    if results:
        print(f"✓ FOUND {len(results)} possible matches:")
        for r in results:
            print(f"\n  Variant: {r['variant']}")
            print(f"  Para: {r['paragraph']}")
            print(f"  Next: {r['context_after']}")
    else:
        print("✗ NOT FOUND - No variants detected")

    # Check original HTML dataset
    print("\n  Original HTML dataset:")
    th_data = check_original_json('th')
    if th_data['exists']:
        print(f"    ✓ EXISTS: {th_data['stems']} stems")
        print(f"    First stem: {th_data['first_stem']}")
    else:
        print(f"    ✗ DOES NOT EXIST in original")

    # 2. Check "false positive" cases
    print("\n\n### 2. Verifying False Positives")
    print("-" * 80)

    # Check ʕr
    print("\n2a. ʕr (1 stem claimed missing)")
    print("    DOCX has: ʕrḏ̣")
    ar_data = check_original_json('ʕr')
    if ar_data['exists']:
        print(f"    ✓ ʕr EXISTS in original: {ar_data['stems']} stems")
        print(f"    ⚠️ This IS a real missing verb (not false positive)")
    else:
        print(f"    ✗ ʕr does NOT exist in original")
        print(f"    ✓ This IS a false positive (was likely part of ʕrḏ̣)")

    # Check ʕw
    print("\n2b. ʕw (1 stem claimed missing)")
    print("    DOCX has: ʕwʕw (reduplicated)")
    aw_data = check_original_json('ʕw')
    if aw_data['exists']:
        print(f"    ✓ ʕw EXISTS in original: {aw_data['stems']} stems")
        print(f"    ⚠️ This IS a real missing verb (not false positive)")
    else:
        print(f"    ✗ ʕw does NOT exist in original")
        print(f"    ✓ This IS a false positive")

    # Check ʕtr
    print("\n2c. ʕtr (1 stem claimed missing)")
    print("    DOCX has: 'see also ʕtrḏ̣' (cross-reference)")
    atr_data = check_original_json('ʕtr')
    if atr_data['exists']:
        print(f"    ✓ ʕtr EXISTS in original: {atr_data['stems']} stems")
        print(f"    ⚠️ This IS a real missing verb (not false positive)")
    else:
        print(f"    ✗ ʕtr does NOT exist in original")
        print(f"    ✓ This IS a false positive (just a cross-ref)")

    # 3. Check high-value ḏyr
    print("\n\n### 3. Verifying ḏyr (11 stems - HIGHEST VALUE)")
    print("-" * 80)

    dyr_data = check_original_json('ḏyr')
    if dyr_data['exists']:
        print(f"✓ ḏyr EXISTS in original: {dyr_data['stems']} stems")
        print(f"  First stem: {dyr_data['first_stem']}")
        print(f"  Etymology: {dyr_data['etymology']}")
        print(f"\n  Status: CONFIRMED - Parser must handle this!")
    else:
        print(f"✗ ḏyr does NOT exist in original")

    # 4. Summary
    print("\n\n" + "="*80)
    print("VERIFICATION SUMMARY")
    print("="*80)

    print("\nCritical Findings:")
    print(f"1. 'th' search results: {len(results) if results else 0} variants found")
    print(f"2. ʕr is {'REAL MISSING' if ar_data['exists'] else 'FALSE POSITIVE'}")
    print(f"3. ʕw is {'REAL MISSING' if aw_data['exists'] else 'FALSE POSITIVE'}")
    print(f"4. ʕtr is {'REAL MISSING' if atr_data['exists'] else 'FALSE POSITIVE'}")
    print(f"5. ḏyr is {'CONFIRMED (11 stems!)' if dyr_data['exists'] else 'NOT IN ORIGINAL'}")

if __name__ == '__main__':
    main()
