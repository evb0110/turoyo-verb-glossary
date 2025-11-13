#!/usr/bin/env python3
"""
TEST: Run full parser on pčq only and verify etymology is complete
"""

import sys
sys.path.insert(0, 'parser')

from parse_docx_production import FixedDocxParser
from docx import Document
import json

print("=" * 80)
print("FULL PARSER TEST: pčq verb")
print("=" * 80)

# Initialize parser
parser = FixedDocxParser()

# Load the DOCX
docx_path = '.devkit/new-source-docx/4. l,m,n,p.docx'
doc = Document(docx_path)

print(f"\n📖 Loading: {docx_path}")

# Find the pčq paragraph
found = False
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('pčq'):
        found = True
        print(f"\n✅ Found pčq at paragraph {i}")

        # Get next paragraph for multi-para etymology
        next_para = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        next_para_text = next_para.text if next_para else None

        # Extract root and etymology
        root, etymology = parser.extract_root_and_etymology(para.text, next_para_text)

        print(f"\n📊 EXTRACTION RESULT:")
        print(f"   Root: {repr(root)}")
        print(f"   Etymology: {json.dumps(etymology, indent=2, ensure_ascii=False)}")

        if etymology and 'etymons' in etymology and etymology['etymons']:
            etym_raw = etymology['etymons'][0].get('raw', '')
            print(f"\n🔍 VERIFICATION:")
            print(f"   Expected: 'prčq cf. Kurd. p'erçiqandin vt. (-p'erçiq-). 1) to crush, press, smash, squash, KED 107'")
            print(f"   Actual:   {repr(etym_raw)}")

            if "to crush, press, smash, squash, KED 107" in etym_raw:
                print(f"\n🎉 SUCCESS! Complete etymology extracted!")
            else:
                print(f"\n❌ FAILED! Etymology is still truncated!")
        else:
            print(f"\n❌ No etymology extracted!")

        break

if not found:
    print("\n❌ pčq not found in document!")
