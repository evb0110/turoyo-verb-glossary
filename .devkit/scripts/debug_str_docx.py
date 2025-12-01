#!/usr/bin/env python3
"""Debug script to extract str verb structure from DOCX"""

from docx import Document
from docx.table import Table
from pathlib import Path

# Find all DOCX files
docx_dir = Path('.devkit/new-source-docx')
docx_files = sorted(docx_dir.glob('*.docx'))

print("Searching for 'str' verb in DOCX files...")
print("=" * 80)

for docx_file in docx_files:
    doc = Document(docx_file)

    # Combine paragraphs and tables into elements list
    elements = []
    for block in doc.element.body:
        if block.tag.endswith('p'):
            # Find corresponding paragraph object
            for para in doc.paragraphs:
                if para._element == block:
                    elements.append(('para', para))
                    break
        elif block.tag.endswith('tbl'):
            # Find corresponding table object
            for table in doc.tables:
                if table._element == block:
                    elements.append(('table', table))
                    break

    # Look for "str" verb
    for idx, (elem_type, elem) in enumerate(elements):
        if elem_type != 'para':
            continue

        text = elem.text.strip()

        # Check for "str" root
        if text == 'str':
            print(f"\n*** FOUND in: {docx_file.name} ***\n")

            # Extract 50 elements after the root
            for j in range(idx, min(idx + 60, len(elements))):
                elem_type, elem = elements[j]

                if elem_type == 'para':
                    para_text = elem.text.strip()
                    marker = ">>>" if j == idx else "   "
                    print(f"{marker} [{j}] PARA: {para_text[:120]}")

                    # Show run details for important lines
                    if any(keyword in para_text for keyword in ['Detransitive', 'II:', 'msatər', 'misatər']):
                        for run in elem.runs:
                            if run.text.strip():
                                print(f"         Run: italic={run.italic}, text=\"{run.text}\"")

                elif elem_type == 'table':
                    print(f"    [{j}] TABLE: {len(elem.rows)} rows x {len(elem.columns)} cols")
                    # Show first row content as preview
                    if elem.rows:
                        first_row_text = ' | '.join(cell.text.strip()[:30] for cell in elem.rows[0].cells)
                        print(f"         First row: {first_row_text}")

            exit(0)  # Stop after finding str

print("\n'str' not found in any DOCX file!")
