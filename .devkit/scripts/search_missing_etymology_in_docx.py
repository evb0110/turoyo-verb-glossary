#!/usr/bin/env python3
"""Search for the 5 missing etymologies in DOCX source"""

import re
from pathlib import Path
from docx import Document

# The 5 roots missing etymology
missing_roots = ['dyq 1', 'frʕ 2', 'gwlʕ', 'ʕyr 1', 'ḏyr']

# Map roots to base form for searching
search_map = {
    'dyq 1': 'dyq',
    'frʕ 2': 'frʕ',
    'gwlʕ': 'gwlʕ',
    'ʕyr 1': 'ʕyr',
    'ḏyr': 'ḏyr'
}

def search_root_in_docx(root, docx_file):
    """Search for root paragraph and its content"""
    doc = Document(docx_file)

    found_paras = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Check if this paragraph starts with the root
        if text.startswith(root):
            # Collect this paragraph and next few
            context = []
            context.append(f"[Para {i}] {text}")

            # Get next 5 paragraphs for context
            for j in range(i+1, min(i+6, len(doc.paragraphs))):
                next_text = doc.paragraphs[j].text.strip()
                if next_text:
                    context.append(f"[Para {j}] {next_text}")

                    # Stop if we hit another root (starts with letter and has italic)
                    if doc.paragraphs[j].runs and any(r.italic for r in doc.paragraphs[j].runs):
                        break

            found_paras.append('\n'.join(context))

    return found_paras

print("=" * 80)
print("SEARCHING FOR MISSING ETYMOLOGIES IN DOCX SOURCE")
print("=" * 80)

docx_dir = Path('.devkit/new-source-docx')
docx_files = sorted(docx_dir.glob('*.docx'))

for root in missing_roots:
    base_root = search_map[root]

    print(f"\n{'=' * 80}")
    print(f"ROOT: {root} (searching for '{base_root}')")
    print(f"{'=' * 80}")

    found_in_files = []

    for docx_file in docx_files:
        results = search_root_in_docx(base_root, docx_file)
        if results:
            found_in_files.append((docx_file.name, results))

    if not found_in_files:
        print(f"❌ NOT FOUND in any DOCX file!")
    else:
        for filename, results in found_in_files:
            print(f"\n📄 Found in: {filename}")
            for result in results:
                print(f"\n{result}")
                print("-" * 80)

print("\n" + "=" * 80)
print("SEARCH COMPLETE")
print("=" * 80)
