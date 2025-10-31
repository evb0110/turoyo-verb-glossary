#!/usr/bin/env python3
"""
Investigate 20 sampled verbs with NULL etymology in DOCX source files.
"""

import json
import re
from pathlib import Path
from docx import Document

DOCX_DIR = Path(".devkit/new-source-docx")
JSON_DIR = Path("server/assets/verbs")

DOCX_FILES = {
    "1. ʔ, ʕ, b, č.docx": ['ʔ', 'ʕ', 'b', 'č'],
    "2. d, f, g, ġ, ǧ.docx": ['d', 'f', 'g', 'ġ', 'ǧ'],
    "3. h,ḥ,k.docx": ['h', 'ḥ', 'k'],
    "4. l,m,n,p.docx": ['l', 'm', 'n', 'p'],
    "5. q,r,s,ṣ.docx": ['q', 'r', 's', 'ṣ'],
    "6. š,t,ṭ,ṯ.docx": ['š', 't', 'ṭ', 'ṯ'],
    "7. v, w, x, y, z, ž.docx": ['v', 'w', 'x', 'y', 'z', 'ž']
}

SAMPLE_VERBS = [
    'bʕy 2', 'dng', 'fqṣ', 'grḏ', 'hmz', 'klpt', 'lhṭ', 'mʕr', 'ngl 2', 'pšṭ',
    'qḥl', 'rǧʕ', 'slpx', 'twy', 'xnxl', 'zbʕ', 'čfx', 'ġrf 2', 'šrqm', 'žġl 1'
]


def get_docx_file_for_verb(verb_root):
    """Find which DOCX file contains this verb based on first letter."""
    first_letter = verb_root[0]
    for docx_file, letters in DOCX_FILES.items():
        if first_letter in letters:
            return DOCX_DIR / docx_file
    return None


def is_root_paragraph(para):
    """Check if paragraph is a root (11pt italic text)."""
    if not para.runs:
        return False

    for run in para.runs:
        if run.font.size and run.font.size.pt == 11 and run.font.italic:
            return True
    return False


def extract_root_text(para):
    """Extract the root text from a root paragraph."""
    text = para.text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def find_verb_in_docx(docx_path, verb_root):
    """Find verb entry in DOCX and extract text after root."""
    doc = Document(docx_path)

    clean_root = verb_root.split()[0]
    root_with_number = verb_root

    found_root = False
    root_text = ""
    following_paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        if is_root_paragraph(para):
            para_text = extract_root_text(para)

            if clean_root in para_text or root_with_number in para_text:
                if verb_root.endswith(' 1') and para_text.endswith('I'):
                    found_root = True
                    root_text = para_text
                elif verb_root.endswith(' 2') and para_text.endswith('II'):
                    found_root = True
                    root_text = para_text
                elif verb_root.endswith(' 3') and para_text.endswith('III'):
                    found_root = True
                    root_text = para_text
                elif ' ' not in verb_root and not para_text.endswith(('I', 'II', 'III')):
                    found_root = True
                    root_text = para_text
                elif ' ' not in verb_root and para_text.endswith('I'):
                    found_root = True
                    root_text = para_text

                if found_root:
                    j = i + 1
                    while j < len(doc.paragraphs):
                        next_para = doc.paragraphs[j]
                        if is_root_paragraph(next_para):
                            break
                        para_content = next_para.text.strip()
                        if para_content:
                            following_paragraphs.append(para_content)
                        j += 1
                    break

    return {
        'found': found_root,
        'root_text': root_text,
        'following_text': '\n\n'.join(following_paragraphs[:5]) if following_paragraphs else ""
    }


def categorize_etymology(docx_result, verb_root):
    """Categorize etymology status."""
    if not docx_result['found']:
        return 'X', "Could not find verb in DOCX"

    following = docx_result['following_text'].lower()

    if not following:
        return 'C', "No text after root in DOCX (genuinely missing)"

    if 'unknown' in following or 'unclear' in following:
        return 'B', "Etymology marked as 'unknown' in source (correctly extracted)"

    etymology_markers = ['<', 'cf.', 'see', 'etym', 'arab', 'syr', 'akkad', 'hebrew', 'kurmanji']
    has_etymology = any(marker in following for marker in etymology_markers)

    if has_etymology:
        return 'A', "Etymology present in DOCX but parser failed to extract (FIXABLE)"

    has_parenthesis = '(' in following and ')' in following
    has_angle_bracket = '<' in following

    if has_parenthesis or has_angle_bracket:
        return 'D', "Non-standard etymology format requiring investigation"

    return 'C', "No clear etymology information in source (genuinely missing)"


def main():
    print("=" * 80)
    print("NULL ETYMOLOGY INVESTIGATION - 20 VERB SAMPLE")
    print("=" * 80)
    print()

    results = []

    for verb in SAMPLE_VERBS:
        docx_file = get_docx_file_for_verb(verb)
        if not docx_file:
            print(f"ERROR: Could not determine DOCX file for {verb}")
            continue

        if not docx_file.exists():
            print(f"ERROR: DOCX file not found: {docx_file}")
            continue

        print(f"Investigating: {verb}")
        print(f"DOCX file: {docx_file.name}")

        docx_result = find_verb_in_docx(docx_file, verb)
        category, reason = categorize_etymology(docx_result, verb)

        result = {
            'verb': verb,
            'docx_file': docx_file.name,
            'found_in_docx': docx_result['found'],
            'root_text': docx_result['root_text'],
            'following_text': docx_result['following_text'][:500],
            'category': category,
            'reason': reason
        }
        results.append(result)

        print(f"  Category: {category} - {reason}")
        if docx_result['found']:
            print(f"  Root text: {docx_result['root_text'][:100]}")
            print(f"  Following text (first 200 chars): {docx_result['following_text'][:200]}")
        print()

    category_counts = {'A': 0, 'B': 0, 'C': 0, 'D': 0, 'X': 0}
    for r in results:
        category_counts[r['category']] += 1

    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)
    print(f"Total investigated: {len(results)}")
    print(f"  Category A (Parser failures - FIXABLE): {category_counts['A']} ({category_counts['A']/len(results)*100:.1f}%)")
    print(f"  Category B (Correctly 'unknown'): {category_counts['B']} ({category_counts['B']/len(results)*100:.1f}%)")
    print(f"  Category C (Genuinely missing): {category_counts['C']} ({category_counts['C']/len(results)*100:.1f}%)")
    print(f"  Category D (Non-standard format): {category_counts['D']} ({category_counts['D']/len(results)*100:.1f}%)")
    print(f"  Category X (Not found in DOCX): {category_counts['X']} ({category_counts['X']/len(results)*100:.1f}%)")
    print()

    print("Extrapolation to all 196 NULL etymology verbs:")
    for cat in ['A', 'B', 'C', 'D']:
        if category_counts[cat] > 0:
            estimated = int(196 * category_counts[cat] / len(results))
            print(f"  Category {cat}: ~{estimated} verbs")
    print()

    with open('.devkit/analysis/null_etymology_investigation_raw.json', 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print("Raw results saved to: .devkit/analysis/null_etymology_investigation_raw.json")


if __name__ == '__main__':
    main()
