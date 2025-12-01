#!/usr/bin/env python3
import os
import zipfile
from lxml import etree
from pathlib import Path
import html

DOCX_DIR = ".devkit/new-source-docx"
OUTPUT_FILE = ".devkit/analysis/turoyo_verbs_reference.html"
XMLNS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'v': 'urn:schemas-microsoft-com:office:office',
    'o': 'urn:schemas-microsoft-com:office:office',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
}

def extract_text_from_element(elem):
    """Extract all text content from an XML element."""
    text_parts = []
    for child in elem.iter():
        if child.tag == f"{{{XMLNS['w']}}}t":
            if child.text:
                text_parts.append(child.text)
    return ''.join(text_parts)

def extract_footnotes(docx_path):
    """Extract footnotes from DOCX file, skipping empty/placeholder footnotes."""
    footnotes_map = {}

    try:
        with zipfile.ZipFile(docx_path, 'r') as docx:
            if 'word/footnotes.xml' in docx.namelist():
                with docx.open('word/footnotes.xml') as f:
                    tree = etree.parse(f)
                    root = tree.getroot()
                    footnote_elements = root.findall('.//w:footnote', XMLNS)

                    for fn in footnote_elements:
                        fn_id = fn.get(f"{{{XMLNS['w']}}}id")
                        if fn_id and fn_id not in ['-1', '0']:  # Skip placeholder footnotes
                            text = extract_text_from_element(fn).strip()
                            if text:  # Only include non-empty footnotes
                                footnotes_map[fn_id] = text
    except Exception as e:
        print(f"Warning: Could not extract footnotes from {docx_path}: {e}")

    return footnotes_map

def get_paragraph_style(para_elem):
    """Get paragraph style name."""
    pPr = para_elem.find(f"{{{XMLNS['w']}}}pPr")
    if pPr is not None:
        pStyle = pPr.find(f"{{{XMLNS['w']}}}pStyle")
        if pStyle is not None:
            return pStyle.get(f"{{{XMLNS['w']}}}val")
    return None

def extract_run_content(run_elem, footnotes_map):
    """Extract content from a text run, handling special cases."""
    content_parts = []

    for elem in run_elem.iter():
        if elem.tag == f"{{{XMLNS['w']}}}t":
            if elem.text:
                content_parts.append(html.escape(elem.text))
        elif elem.tag == f"{{{XMLNS['w']}}}br":
            content_parts.append("<br/>")
        elif elem.tag == f"{{{XMLNS['w']}}}footnoteReference":
            fn_id = elem.get(f"{{{XMLNS['w']}}}id")
            if fn_id and fn_id in footnotes_map:
                content_parts.append(f'<sup class="footnote-ref" data-footnote-id="{fn_id}">[{fn_id}]</sup>')

    return ''.join(content_parts)

def convert_paragraph(para_elem, footnotes_map):
    """Convert a paragraph element to HTML."""
    style = get_paragraph_style(para_elem)
    text_content = []

    # Extract all text runs
    runs = para_elem.findall(f"{{{XMLNS['w']}}}r")
    for run in runs:
        # Check for run properties
        rPr = run.find(f"{{{XMLNS['w']}}}rPr")
        is_italic = False
        is_bold = False
        if rPr is not None:
            is_italic = rPr.find(f"{{{XMLNS['w']}}}i") is not None
            is_bold = rPr.find(f"{{{XMLNS['w']}}}b") is not None

        run_text = extract_run_content(run, footnotes_map)

        if is_italic:
            run_text = f"<i>{run_text}</i>"
        if is_bold:
            run_text = f"<b>{run_text}</b>"

        text_content.append(run_text)

    text = ''.join(text_content).strip()

    if not text:
        return None

    # Map Word styles to HTML tags
    if style == 'Heading1':
        return f"<h1>{text}</h1>"
    elif style == 'Heading2':
        return f"<h2>{text}</h2>"
    elif style == 'Heading3':
        return f"<h3>{text}</h3>"
    elif style == 'ListBullet':
        return f"<li>{text}</li>"
    else:
        return f"<p>{text}</p>"

def convert_table(table_elem, footnotes_map):
    """Convert a table element to HTML."""
    html_table = '<table class="verb-table" border="1" cellpadding="5" cellspacing="0">'

    rows = table_elem.findall(f"{{{XMLNS['w']}}}tr")
    for row in rows:
        html_table += '<tr>'
        cells = row.findall(f"{{{XMLNS['w']}}}tc")
        for cell in cells:
            # Extract content from cell
            paras = cell.findall(f"{{{XMLNS['w']}}}p")
            cell_content = []
            for para in paras:
                para_html = convert_paragraph(para, footnotes_map)
                if para_html:
                    cell_content.append(para_html)

            cell_html = ''.join(cell_content) if cell_content else '<p>&nbsp;</p>'
            html_table += f'<td>{cell_html}</td>'
        html_table += '</tr>'

    html_table += '</table>'
    return html_table

def convert_docx_to_html(docx_path):
    """Convert a single DOCX file to HTML content."""
    html_content = []
    file_title = Path(docx_path).stem
    html_content.append(f'<section class="docx-file" data-source="{file_title}">')
    html_content.append(f'<h2 class="file-header">{html.escape(file_title)}</h2>')

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error loading {docx_path}: {e}")
        return ''.join(html_content) + '</section>'

    # Extract footnotes first
    footnotes_map = extract_footnotes(docx_path)

    # Parse document using lxml to preserve formatting
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        with docx_zip.open('word/document.xml') as f:
            tree = etree.parse(f)
            root = tree.getroot()
            body = root.find(f"{{{XMLNS['w']}}}body")

            if body is not None:
                for elem in body:
                    if elem.tag == f"{{{XMLNS['w']}}}p":
                        para_html = convert_paragraph(elem, footnotes_map)
                        if para_html:
                            html_content.append(para_html)
                    elif elem.tag == f"{{{XMLNS['w']}}}tbl":
                        table_html = convert_table(elem, footnotes_map)
                        html_content.append(table_html)

    html_content.append('</section>')
    return '\n'.join(html_content)

def merge_docx_files_to_html(docx_dir, output_path):
    """Merge all DOCX files in directory to a single HTML file."""
    docx_files = sorted([f for f in os.listdir(docx_dir) if f.endswith('.docx')])

    print(f"Found {len(docx_files)} DOCX files to convert:")
    for f in docx_files:
        print(f"  - {f}")

    # Build HTML document
    html_parts = ['''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Turoyo Verb Glossary - Complete Reference</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
            color: #333;
        }

        h1 {
            color: #1a1a1a;
            border-bottom: 3px solid #0066cc;
            padding-bottom: 10px;
            margin-top: 30px;
        }

        h2 {
            color: #0066cc;
            margin-top: 25px;
            background-color: #f0f4ff;
            padding: 10px;
            border-left: 4px solid #0066cc;
        }

        h3 {
            color: #333;
            margin-top: 15px;
        }

        .file-header {
            font-size: 1.8em;
            margin-top: 0;
            background-color: #003366;
            color: white;
            padding: 15px;
            border-radius: 5px;
        }

        .docx-file {
            background-color: white;
            margin: 20px 0;
            padding: 20px;
            border-radius: 5px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }

        p {
            margin: 10px 0;
            text-align: justify;
        }

        li {
            margin: 8px 0;
            margin-left: 20px;
        }

        table {
            margin: 15px 0;
            width: 100%;
            border-collapse: collapse;
            background-color: #fff;
        }

        th, td {
            border: 1px solid #ddd;
            padding: 10px;
            text-align: left;
        }

        th {
            background-color: #f0f0f0;
            font-weight: bold;
        }

        tr:nth-child(even) {
            background-color: #f9f9f9;
        }

        .footnote-ref {
            color: #0066cc;
            cursor: pointer;
            font-weight: bold;
        }

        .footnotes-section {
            margin-top: 50px;
            padding-top: 20px;
            border-top: 2px solid #ccc;
            background-color: #f9f9f9;
            padding: 20px;
            border-radius: 5px;
        }

        .footnote-item {
            margin: 15px 0;
            padding: 10px;
            background-color: white;
            border-left: 3px solid #0066cc;
            padding-left: 15px;
        }

        .toc {
            background-color: #f0f4ff;
            padding: 20px;
            border-radius: 5px;
            margin-bottom: 30px;
            border: 1px solid #0066cc;
        }

        .toc h2 {
            margin-top: 0;
        }

        .toc ul {
            list-style-type: none;
            padding-left: 0;
        }

        .toc li {
            margin: 5px 0;
            padding-left: 20px;
        }

        .toc a {
            color: #0066cc;
            text-decoration: none;
        }

        .toc a:hover {
            text-decoration: underline;
        }

        i {
            font-style: italic;
        }

        b {
            font-weight: bold;
        }
    </style>
</head>
<body>
    <h1>Turoyo Verb Glossary - Complete Reference</h1>
    <p style="text-align: center; color: #666; font-style: italic;">
        Comprehensive HTML reference conversion from original DOCX source files
    </p>

    <div class="toc">
        <h2>Contents</h2>
        <ul>
''']

    # Add TOC entries
    for docx_file in docx_files:
        docx_path = os.path.join(docx_dir, docx_file)
        file_id = docx_file.replace('.docx', '').replace(' ', '_').replace(',', '')
        file_title = docx_file.replace('.docx', '')
        html_parts.append(f'            <li><a href="#{file_id}">{html.escape(file_title)}</a></li>')

    html_parts.append('        </ul>\n    </div>\n')

    # Convert each DOCX file
    total_paragraphs = 0
    total_tables = 0
    all_footnotes = {}

    for idx, docx_file in enumerate(docx_files, 1):
        docx_path = os.path.join(docx_dir, docx_file)
        print(f"\n[{idx}/{len(docx_files)}] Converting: {docx_file}")

        try:
            file_html = convert_docx_to_html(docx_path)
            html_parts.append(file_html)

            # Count elements
            doc = Document(docx_path)
            total_paragraphs += len(doc.paragraphs)
            total_tables += len(doc.tables)

            # Collect footnotes
            footnotes_map = extract_footnotes(docx_path)
            all_footnotes.update(footnotes_map)

            print(f"  ✓ Converted: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(footnotes_map)} footnotes")
        except Exception as e:
            print(f"  ✗ Error converting {docx_file}: {e}")
            import traceback
            traceback.print_exc()

    # Add footnotes section if there are any
    if all_footnotes:
        html_parts.append('\n<div class="footnotes-section">')
        html_parts.append('<h2>Footnotes</h2>')
        for fn_id in sorted(all_footnotes.keys(), key=lambda x: int(x) if x.isdigit() else 0):
            fn_text = all_footnotes[fn_id]
            html_parts.append(f'<div class="footnote-item"><sup>[{fn_id}]</sup> {html.escape(fn_text)}</div>')
        html_parts.append('</div>')

    # Close HTML
    html_parts.append('''
</body>
</html>''')

    html_output = '\n'.join(html_parts)

    # Write to file
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_output)

    print(f"\n{'='*60}")
    print(f"✓ Conversion complete!")
    print(f"{'='*60}")
    print(f"Output file: {output_path}")
    print(f"Total paragraphs: {total_paragraphs}")
    print(f"Total tables: {total_tables}")
    print(f"Total footnotes: {len(all_footnotes)}")
    print(f"HTML file size: {len(html_output) / 1024 / 1024:.2f} MB")

    return output_path

if __name__ == '__main__':
    from docx import Document

    output = merge_docx_files_to_html(DOCX_DIR, OUTPUT_FILE)
    print(f"\n✓ HTML reference file created: {output}")
