#!/usr/bin/env python3
"""
DOCX Parser V3 - Etymology Extraction Fixes
- FIX 1: NULL etymologies - Support '( <Source' pattern (space after opening paren)
- FIX 2: Multi-etymon splitting bug - Don't split on 'or' within definitions
- FIX 3: Incomplete extractions - Better 'cf.' and corpus reference patterns
- Previous: Homonym numbering, translation extraction, contextual validation, Unicode normalization
"""

import re
import json
from pathlib import Path
from docx import Document
from collections import defaultdict

class FixedDocxParser:
    """Complete DOCX parser with all accuracy fixes"""

    def __init__(self):
        self.verbs = []
        self.stats = defaultdict(int)
        self.contextual_roots = []

    def is_letter_header(self, para):
        return para.style and para.style.name == 'Heading 1'

    def is_root_paragraph(self, para, next_para=None):
        if not para.text.strip():
            return False

        text = para.text.strip()
        turoyo_chars = r'ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə'
        # CRITICAL FIX: Include combining diacritics (U+0300-U+036F) to handle decomposed characters like ḏ̣ (ḏ + combining dot below)
        turoyo_with_combining = rf'[{turoyo_chars}\u0300-\u036F]'

        has_root = re.match(rf'^({turoyo_with_combining}{{2,12}})(?:\s+\d+)?(?:\s|\(|<|$)', text)
        is_cross_ref = bool(re.search(r'→|See\s+[ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]', text))

        if not has_root or is_cross_ref:
            return False

        has_italic = any(r.italic for r in para.runs)
        sizes = [r.font.size.pt for r in para.runs if r.font.size]
        has_11pt = 11.0 in sizes

        if has_italic and has_11pt:
            return True

        if next_para and self.is_stem_header(next_para):
            has_etymology = bool(re.search(r'\(<\s*[A-Za-z.]+', text))
            has_turoyo_chars = len(text) >= 3

            if has_italic or has_11pt or has_etymology or has_turoyo_chars:
                return True

        return False

    def is_stem_header(self, para):
        if not para.text.strip():
            return False

        text = para.text.strip()
        has_stem = re.match(r'^([IVX]+):\s*', text)

        if not has_stem:
            if text.startswith('Detransitive'):
                return True
            return False

        has_bold = any(r.bold for r in para.runs)
        has_italic = any(r.italic for r in para.runs)
        sizes = [r.font.size.pt for r in para.runs if r.font.size]
        has_14pt = 14.0 in sizes

        if has_bold and has_italic and has_14pt and has_stem:
            return True

        if has_stem and (has_bold or has_italic):
            return True

        if has_stem and sizes:
            return True

        return False

    def parse_etymology_full(self, text, next_para_text=None):
        """Full etymology parsing with multi-paragraph support and flexible patterns

        FIXES:
        1. Support missing opening paren: 'ngl (< MEA...' becomes 'ngl ( < MEA...'
        2. Don't split on 'or' within definitions (brz bug)
        3. Better 'cf.' pattern matching for incomplete cases
        """

        # Try multiple patterns in order of specificity

        # Pattern 1: Standard (< Source root cf. ref: meaning)
        # Handle nested parentheses by finding matching closing paren
        match = None
        paren_start = text.find('(<')
        if paren_start >= 0:
            # Find matching closing paren by counting depth
            depth = 1
            i = paren_start + 1
            while i < len(text) and depth > 0:
                if text[i] == '(':
                    depth += 1
                elif text[i] == ')':
                    depth -= 1
                i += 1

            if depth == 0:
                # Found matching closing paren
                etym_content = text[paren_start+2:i-1].strip()
                class MatchLike:
                    def __init__(self, content):
                        self._content = content
                    def group(self, n):
                        return self._content if n == 1 else None
                match = MatchLike(etym_content)

        # Pattern 1b: FIX 1 - MISSING OPENING PAREN before < - e.g., 'ngl (< MEA...' with space after (
        # This is the bug causing NULL etymologies for ngl 2, zyr 2
        if not match:
            paren_space_start = re.search(r'\(\s+<', text)
            if paren_space_start:
                paren_pos = paren_space_start.start()
                # Find matching closing paren by counting depth
                depth = 1
                i = paren_pos + 1
                while i < len(text) and depth > 0:
                    if text[i] == '(':
                        depth += 1
                    elif text[i] == ')':
                        depth -= 1
                    i += 1

                if depth == 0:
                    # Extract content between ( and ), skip the '(' and any whitespace before '<'
                    etym_content = text[paren_pos+1:i-1].strip()
                    if etym_content.startswith('<'):
                        etym_content = etym_content[1:].strip()
                    class MatchLike:
                        def __init__(self, content):
                            self._content = content
                        def group(self, n):
                            return self._content if n == 1 else None
                    match = MatchLike(etym_content)

        # Pattern 2: NO opening paren at all - root <Source... text with closing paren later
        # Example: ḏyr <Ar. ḍrr 'to harm, damage'), cf. Turk...
        if not match:
            # Look for: root_chars followed by space, then <Source with closing paren later
            no_open_paren = re.search(r'^\S+\s+<\s*([^<>]+?)\)(?:,\s+cf\.|;|\s+|$)', text)
            if no_open_paren:
                match = no_open_paren

        # Pattern 3: Missing opening paren - just <Source... (alternative form)
        if not match:
            match = re.search(r'(?:^|[\s\d])<\s*([A-Z][^<>]+?)\)(?:\s|$|;)', text)

        # Pattern 4: FIX 3 - 'cf.' without '<' - e.g., 'ʕngr ( cf. Syr. ʕgr...'
        if not match:
            cf_pattern = re.search(r'\(\s+cf\.\s+(.+)\)(?:\s|$)', text)
            if cf_pattern:
                match = cf_pattern

        # Pattern 5: FIX 3 - Space before opening paren for FKD references - e.g., 'sxy ( FKD; 1493...'
        if not match:
            space_paren = re.search(r'\(\s+([A-Z][A-Z])', text)
            if space_paren:
                paren_pos = space_paren.start()
                depth = 1
                i = paren_pos + 1
                while i < len(text) and depth > 0:
                    if text[i] == '(':
                        depth += 1
                    elif text[i] == ')':
                        depth -= 1
                    i += 1
                if depth == 0:
                    etym_content = text[paren_pos+1:i-1].strip()
                    class MatchLike:
                        def __init__(self, content):
                            self._content = content
                        def group(self, n):
                            return self._content if n == 1 else None
                    match = MatchLike(etym_content)

        # Pattern 6: Corpus/text reference - (Text_name ... info)
        if not match:
            paren_start = text.find('(')
            if paren_start >= 0:
                after_paren = text[paren_start+1:].strip()
                if re.match(r'[A-Z]\w+\s+(?:text|corpus|Talay)', after_paren):
                    depth = 1
                    i = paren_start + 1
                    while i < len(text) and depth > 0:
                        if text[i] == '(':
                            depth += 1
                        elif text[i] == ')':
                            depth -= 1
                        i += 1
                    if depth == 0:
                        etym_content = text[paren_start+1:i-1].strip()
                        class MatchLike:
                            def __init__(self, content):
                                self._content = content
                            def group(self, n):
                                return self._content if n == 1 else None
                        match = MatchLike(etym_content)

        # Pattern 7: Alternative start patterns (see X, cf. X, unknown)
        if not match:
            match = re.search(r'\(((?:see|cf\.|unknown)[^)]+)\)', text)

        # Pattern 8: Multi-paragraph - unclosed paren
        if not match and next_para_text:
            paren_match = re.search(r'\(<\s*(.+)$', text)
            if paren_match:
                close_match = re.search(r'^(.+?)\)', next_para_text)
                if close_match:
                    combined = paren_match.group(1) + ' ' + close_match.group(1)
                    etym_text = combined.strip().rstrip(';').strip()

                    relationship = None
                    etymon_parts = [etym_text]

                    if ' also ' in etym_text:
                        relationship = 'also'
                        etymon_parts = [part.strip() for part in etym_text.split(' also ')]
                    elif '; and ' in etym_text or ', and ' in etym_text:
                        relationship = 'and'
                        etymon_parts = [part.strip() for part in re.split(r'[;,]\s*and\s+', etym_text)]

                    etymons = []
                    for part in etymon_parts:
                        etymon = self.parse_single_etymon(part)
                        if etymon:
                            etymons.append(etymon)

                    if etymons:
                        result = {'etymons': etymons}
                        if relationship:
                            result['relationship'] = relationship
                        return result

        if not match:
            return None

        etym_text = match.group(1).strip().rstrip(';').strip()

        # FIX 2: CRITICAL - Don't split on 'or' within definitions
        # Only split if 'or' appears between distinct source patterns
        # The brz bug: "a field, plain, or wide expanse" should NOT split on 'or'
        #
        # Strategy: Only split on ' or ' if both sides look like separate etymons
        # (i.e., they start with a source language name like "Arab.", "Syr.", etc.)
        relationship = None
        etymon_parts = [etym_text]

        if ' also ' in etym_text:
            relationship = 'also'
            etymon_parts = [part.strip() for part in etym_text.split(' also ')]
        elif '; and ' in etym_text or ', and ' in etym_text:
            relationship = 'and'
            etymon_parts = [part.strip() for part in re.split(r'[;,]\s*and\s+', etym_text)]

        etymons = []
        for part in etymon_parts:
            etymon = self.parse_single_etymon(part)
            if etymon:
                etymons.append(etymon)

        if not etymons:
            return None

        result = {'etymons': etymons}
        if relationship:
            result['relationship'] = relationship

        return result

    def parse_single_etymon(self, etym_text):
        """Parse a single etymon with full structure and flexible patterns

        FIX 3: Better extraction for 'cf.' patterns and corpus references
        """

        # Normalize abbreviated source names
        etym_text_normalized = etym_text.replace('Ar.', 'Arab.')

        # Pattern 1: Arab. bdl (II) cf. Wehr 71-72: verändern, umändern...
        structured = re.match(
            r'([A-Za-z.]+)\s+([^\s]+)\s+(?:\([^)]+\)\s+)?cf\.\s+([^:]+):\s*(.+)',
            etym_text_normalized, re.DOTALL
        )
        if structured:
            return {
                'source': structured.group(1).strip(),
                'source_root': structured.group(2).strip(),
                'reference': structured.group(3).strip(),
                'meaning': self.normalize_whitespace(structured.group(4)),
            }

        # Pattern 2: Without cf - Arab. bdl, Wehr 71-72: verändern...
        no_cf = re.match(
            r'([A-Za-z.]+)\s+([^\s,]+),\s+([^:]+):\s*(.+)',
            etym_text_normalized, re.DOTALL
        )
        if no_cf:
            return {
                'source': no_cf.group(1).strip(),
                'source_root': no_cf.group(2).strip(),
                'reference': no_cf.group(3).strip(),
                'meaning': self.normalize_whitespace(no_cf.group(4)),
            }

        # Pattern 3: FIX 3 - 'cf. Syr. root (Pa.), SL ref: meaning' (ʕngr case)
        cf_syr_pattern = re.match(
            r'cf\.\s+([A-Za-z.]+)\s+([^\s,]+)\s+(?:\([^)]+\),?\s+)?([A-Za-z]+\s+[\d.]+):\s*(.+)',
            etym_text, re.DOTALL
        )
        if cf_syr_pattern:
            return {
                'source': cf_syr_pattern.group(1).strip(),
                'source_root': cf_syr_pattern.group(2).strip(),
                'reference': cf_syr_pattern.group(3).strip(),
                'meaning': self.normalize_whitespace(cf_syr_pattern.group(4)),
            }

        # Pattern 4: FIX 3 - FKD references like 'FKD; 1493 sexî adj. ar. généreux'
        fkd_pattern = re.match(
            r'([A-Z]{2,})[;,]\s+(\d+)\s+(.+)',
            etym_text, re.DOTALL
        )
        if fkd_pattern:
            return {
                'source': fkd_pattern.group(1).strip(),
                'reference': fkd_pattern.group(2).strip(),
                'notes': self.normalize_whitespace(fkd_pattern.group(3)),
            }

        # Pattern 5: "see X" or "cf. X" pattern - e.g., "see Tezel 71; cf. probably Syr..."
        see_pattern = re.match(r'(?:see|cf\.)\s+(.+)', etym_text, re.DOTALL)
        if see_pattern:
            return {
                'notes': self.normalize_whitespace(see_pattern.group(0)),
                'raw': etym_text
            }

        # Pattern 6: "unknown" or similar
        if etym_text.strip().lower().startswith('unknown'):
            return {
                'notes': self.normalize_whitespace(etym_text),
                'raw': etym_text
            }

        # Pattern 7: Simple source + notes
        simple = re.match(r'([A-Za-z.]+)\s+(.+)', etym_text)
        if simple:
            source = simple.group(1).strip()
            notes = simple.group(2).strip()

            # Normalize source name
            source = source.replace('Ar.', 'Arab.')

            return {
                'source': source,
                'notes': self.normalize_whitespace(notes),
            }

        # Fallback: raw text
        return {'raw': etym_text}

    def normalize_whitespace(self, text):
        if not text:
            return ""
        return re.sub(r'\s+', ' ', text).strip()

    def extract_root_and_etymology(self, text, next_para_text=None):
        text = text.strip()
        # CRITICAL FIX: Include combining diacritics to match decomposed characters
        root_match = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə\u0300-\u036F]{2,12}(?:\s+\d+)?)(?:\s|\(|<|$)', text)
        if not root_match:
            return None, None

        root = root_match.group(1).strip()

        # Parse full etymology with multi-paragraph support
        etymology = self.parse_etymology_full(text, next_para_text)

        return root, etymology

    def extract_stem_info(self, text):
        match = re.match(r'^([IVX]+):\s*(.+)', text.strip())
        if match:
            stem_num = match.group(1)
            forms_text = match.group(2).strip()
            forms_match = re.match(r'^([^\s]+(?:/[^\s]+)*)', forms_text)
            if forms_match:
                forms_str = forms_match.group(1)
                forms = [f.strip() for f in forms_str.split('/') if f.strip()]
                return stem_num, forms
        return None, []

    def extract_translations_improved(self, cell_text):
        """Extract translations with comprehensive quote handling"""
        translations = []

        # Pattern 1: Curly quotes ʻ...ʼ (U+02BB ... U+02BC) - most common
        # Use non-greedy to avoid spanning multiple quotes
        curly = re.findall(r'ʻ(.+?)ʼ', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in curly if len(t.strip()) > 3])

        # Pattern 2: Straight single quotes '...' (U+0027)
        # Use non-greedy and require substantial length to avoid Turoyo contractions
        straight_single = re.findall(r'\'(.{15,}?)\'', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in straight_single if len(t.strip()) > 15])

        # Pattern 3: Double quotes "..."
        double = re.findall(r'\"(.+?)\"', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in double if len(t.strip()) > 3])

        # Pattern 4: Mixed curly+straight quotes (ʻtext' or 'textʼ)
        mixed1 = re.findall(r'ʻ(.{10,}?)\'', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in mixed1 if len(t.strip()) > 10])
        mixed2 = re.findall(r'\'(.{10,}?)ʼ', cell_text, re.DOTALL)
        translations.extend([self.normalize_whitespace(t) for t in mixed2 if len(t.strip()) > 10])

        # Deduplicate while preserving order
        seen = set()
        unique_translations = []
        for t in translations:
            t_clean = t.strip()
            if t_clean and len(t_clean) > 3 and t_clean not in seen:
                seen.add(t_clean)
                unique_translations.append(t_clean)

        return unique_translations

    def is_valid_translation(self, text):
        """Filter out references, notes, and non-translation text"""
        text = text.strip()

        # Just punctuation or quotes
        if text in ['ʻ', 'ʼ', "'", '"', ';', ',', ':', '.', '(', ')', '[', ']']:
            return False

        # References like "EL 26", "JL 20.7.44", "[LTN]", "prs 24/2"
        if re.match(r'^[A-Z]{1,3}\s+\d+', text) or re.match(r'^\[?[A-Z]{2,5}\]?;?$', text):
            return False
        if re.match(r'^(prs|cf\.)\s+\d+', text):
            return False

        # Pure numbers, dates, page references
        if re.match(r'^[\d;/\s\[\]]+$', text):
            return False

        # Meta-notes (Russian only)
        if re.match(r'^[А-Яа-я\s]+$', text):
            return False

        # Must contain at least one lowercase letter (German/English have lowercase)
        if not re.search(r'[a-zäöüß]', text):
            return False

        # Minimum length after other filters (allow shorter translations now)
        if len(text) < 10:
            return False

        return True

    def is_likely_turoyo(self, text):
        """Heuristic to detect Turoyo text when formatting metadata is missing (italic=None)"""
        if not text or not isinstance(text, str):
            return False

        text = text.strip()

        # Turoyo-specific characters that rarely appear in German/English
        turoyo_chars = 'ʔʕḏṯṣṭḥǧġšžəāēīū'
        turoyo_char_count = sum(1 for c in text if c in turoyo_chars)

        # If has significant Turoyo characters, likely Turoyo
        if len(text) > 0 and (turoyo_char_count / len(text)) > 0.15:
            return True

        # Check for Turoyo character presence (even if ratio is low due to long text)
        if turoyo_char_count >= 3:
            return True

        # Check for pattern: starts with Turoyo, no German words
        german_indicators = ['der', 'die', 'das', 'und', 'ist', 'sein', 'werden', 'zu', 'von', 'mit']
        has_german = any(word in text.lower() for word in german_indicators)

        if turoyo_char_count > 0 and not has_german:
            return True

        return False

    def is_form_only_entry(self, text):
        """Detect if text is a single conjugated form with optional reference"""
        text = text.strip()

        simple_form_with_ref = r'^[a-zāēīūəʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓ]+-?[a-zāēīūəʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓ=]*\s*;\s*\d+\s*;\s*$'
        if re.match(simple_form_with_ref, text, re.IGNORECASE):
            return True

        variant_pattern = r'^[a-zāēīūəʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓ]+-?[a-zāēīūəʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓ=]*(?:\s*;\s*[a-zāēīūəʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓ]+-?[a-zāēīūəʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓ=]*)+\s*;\s*\d+\s*;\s*$'
        if re.match(variant_pattern, text, re.IGNORECASE):
            return True

        return False

    def is_reference_only(self, text):
        """Detect if text is just a reference/metadata line"""
        text = text.strip()

        ref_patterns = [
            r'^[A-Z][a-z]+\s+p\.c\.',
            r'^\d{4}_\d{2}_\d{2}',
            r'^[A-Z][a-z]+_[A-Z]',
            r'^[A-Z][a-z]+-[A-Z]',
            r'^\([A-Za-z\s]+\s+\d+',
        ]

        for pattern in ref_patterns:
            if re.match(pattern, text):
                return True

        return False

    def extract_variant_forms(self, text):
        """Extract variant forms like 'tukilo; tawkilo; tawkolo; 721;' or 'mbarzaqqe l-am=maye; 770;'"""
        parts = [p.strip() for p in text.split(';') if p.strip()]

        forms = []
        references = []

        for part in parts:
            if re.match(r'^\d+$', part):
                references.append(part)
            elif len(part) > 1:
                turoyo_ratio = sum(1 for c in part if c in 'ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə-=')
                if turoyo_ratio >= 2 or self.is_likely_turoyo(part):
                    forms.append(part)

        return forms, references

    def parse_table_cell(self, cell):
        """Parse examples from table cell using run formatting to separate Turoyo from translations

        ENHANCED VERSION (2025-10-31): Handle 5 categories of empty Turoyo cases:
        1. Form-only entries (single form + reference)
        2. Numbered list items with italic Turoyo
        3. Reference-only lines
        4. Story excerpts
        5. Variant form lists
        """
        examples = []

        for para in cell.paragraphs:
            full_para_text = para.text.strip()

            if not full_para_text:
                continue

            if self.is_reference_only(full_para_text):
                continue

            turoyo_parts = []
            translation_parts = []

            for run in para.runs:
                text = run.text.strip()
                if not text:
                    continue

                if run.italic is True:
                    turoyo_parts.append(run.text)
                elif run.italic is False:
                    if len(text) > 2:
                        translation_parts.append(text)
                else:
                    if self.is_likely_turoyo(text):
                        turoyo_parts.append(run.text)
                    elif len(text) > 2:
                        translation_parts.append(text)

            turoyo_text = ''.join(turoyo_parts)

            translations = []
            for t in translation_parts:
                normalized = self.normalize_whitespace(t)
                if self.is_valid_translation(normalized):
                    translations.append(normalized)

            if not turoyo_text and not translations:
                if self.is_form_only_entry(full_para_text):
                    forms, refs = self.extract_variant_forms(full_para_text)
                    if forms:
                        turoyo_text = '; '.join(forms)
                        example = {
                            'turoyo': turoyo_text,
                            'translations': [],
                            'references': refs
                        }
                        examples.append(example)
                        continue

            elif not turoyo_text and translations:
                first_translation = translations[0] if translations else ''

                if self.is_form_only_entry(first_translation):
                    forms, refs = self.extract_variant_forms(first_translation)
                    if forms:
                        example = {
                            'turoyo': '; '.join(forms),
                            'translations': translations[1:] if len(translations) > 1 else [],
                            'references': refs
                        }
                        examples.append(example)
                        continue

                if self.is_likely_turoyo(first_translation):
                    example = {
                        'turoyo': first_translation,
                        'translations': translations[1:] if len(translations) > 1 else [],
                        'references': []
                    }
                    examples.append(example)
                    continue

            if turoyo_text or translations:
                references = re.findall(r'\d+(?:;\s*\d+)?(?:/\d+)?', turoyo_text)

                example = {
                    'turoyo': self.normalize_whitespace(turoyo_text),
                    'translations': translations,
                    'references': references[:2] if references else []
                }

                if example['turoyo'] or example['translations']:
                    examples.append(example)

        return examples

    def parse_document_with_tables(self, docx_path):
        """Parse document using element tree to get table positions"""
        print(f"\n📖 {docx_path.name}")

        doc = Document(docx_path)

        # Build element map
        elements = []
        for el in doc.element.body:
            tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

            if tag == 'p':
                for para in doc.paragraphs:
                    if para._element == el:
                        elements.append(('para', para))
                        break
            elif tag == 'tbl':
                for table in doc.tables:
                    if table._element == el:
                        elements.append(('table', table))
                        break

        current_verb = None
        current_stem = None

        for idx, (elem_type, elem) in enumerate(elements):
            if elem_type == 'para':
                para = elem

                if self.is_letter_header(para):
                    continue

                next_para = None
                for j in range(idx + 1, min(idx + 4, len(elements))):
                    if elements[j][0] == 'para':
                        candidate = elements[j][1]
                        if candidate.text.strip():
                            next_para = candidate
                            break

                is_root = self.is_root_paragraph(para, next_para)
                is_root_strict = (
                    para.text.strip() and
                    any(r.italic for r in para.runs) and
                    11.0 in [r.font.size.pt for r in para.runs if r.font.size]
                )

                if is_root:
                    if current_verb:
                        self.verbs.append(current_verb)
                        self.stats['verbs_parsed'] += 1

                    # Pass next paragraph text for multi-paragraph etymology support
                    next_para_text = next_para.text if next_para else None
                    root, etymology = self.extract_root_and_etymology(para.text, next_para_text)
                    if root:
                        current_verb = {
                            'root': root,
                            'etymology': etymology,
                            'cross_reference': None,
                            'stems': [],
                            'uncertain': '???' in para.text
                        }
                        current_stem = None

                        if not is_root_strict:
                            self.contextual_roots.append(root)
                            self.stats['contextual_roots'] += 1

                elif self.is_stem_header(para):
                    stem_num, forms = self.extract_stem_info(para.text)
                    if stem_num and current_verb is not None:
                        current_stem = {
                            'stem': stem_num,
                            'forms': forms,
                            'conjugations': {}
                        }
                        current_verb['stems'].append(current_stem)
                        self.stats['stems_parsed'] += 1

                elif 'Detransitive' in para.text and current_verb:
                    if not any(s['stem'] == 'Detransitive' for s in current_verb['stems']):
                        current_stem = {
                            'stem': 'Detransitive',
                            'forms': [],
                            'conjugations': {}
                        }
                        current_verb['stems'].append(current_stem)
                        self.stats['detransitive_entries'] += 1

            elif elem_type == 'table':
                table = elem

                if current_stem is not None and table.rows:
                    row = table.rows[0]
                    if len(row.cells) >= 2:
                        conj_type = row.cells[0].text.strip()
                        examples_cell = row.cells[1]

                        examples = self.parse_table_cell(examples_cell)

                        if conj_type and examples:
                            current_stem['conjugations'][conj_type] = examples
                            self.stats['examples_parsed'] += len(examples)

        if current_verb:
            self.verbs.append(current_verb)
            self.stats['verbs_parsed'] += 1

        print(f"   ✓ {self.stats['verbs_parsed']} verbs, {self.stats['stems_parsed']} stems, {self.stats['examples_parsed']} examples")
        if self.stats.get('contextual_roots'):
            print(f"   🔍 {self.stats['contextual_roots']} roots found via contextual validation")

    def add_homonym_numbers(self):
        """Add sequential numbers to homonyms with different etymologies"""
        print("\n🔍 Checking for homonyms...")

        root_groups = defaultdict(list)
        for idx, verb in enumerate(self.verbs):
            base_root = re.sub(r'\s+\d+$', '', verb['root'])
            root_groups[base_root].append((idx, verb))

        numbered_count = 0
        for root, entries in root_groups.items():
            if len(entries) == 1:
                idx, verb = entries[0]
                if self.verbs[idx]['root'] != root:
                    old_root = self.verbs[idx]['root']
                    self.verbs[idx]['root'] = root
                    print(f"   ℹ️  Stripped manual number: {old_root} → {root}")
                continue

            # Get etymology signatures
            etymologies = []
            for idx, verb in entries:
                etym = verb.get('etymology')
                if etym:
                    if 'etymons' in etym and etym['etymons']:
                        first_etymon = etym['etymons'][0]
                        sig = (
                            first_etymon.get('source', ''),
                            first_etymon.get('source_root', ''),
                            first_etymon.get('notes', ''),
                            first_etymon.get('raw', ''),
                            first_etymon.get('reference', '')
                        )
                    else:
                        sig = None
                else:
                    sig = None
                etymologies.append((idx, sig))

            unique_etyms = set(sig for _, sig in etymologies)

            # Number if multiple different etymologies
            if len(unique_etyms) > 1:
                print(f"   ℹ️  Found homonyms for '{root}' with {len(unique_etyms)} different etymologies")
                for entry_num, (idx, sig) in enumerate(etymologies, 1):
                    old_root = self.verbs[idx]['root']
                    self.verbs[idx]['root'] = f"{root} {entry_num}"
                    print(f"      {old_root} → {self.verbs[idx]['root']}")
                numbered_count += len(entries)
            elif len(entries) == 1:
                idx, verb = entries[0]
                if self.verbs[idx]['root'] != root:
                    old_root = self.verbs[idx]['root']
                    self.verbs[idx]['root'] = root
                    print(f"   ℹ️  Stripped manual number: {old_root} → {root}")

        if numbered_count > 0:
            self.stats['homonyms_numbered'] = numbered_count
            print(f"   ✅ Auto-numbered {numbered_count} homonym entries")
        else:
            print(f"   ℹ️  No homonyms requiring numbering found")

    def parse_all_files(self, docx_dir):
        print("=" * 80)
        print("DOCX PARSER V3 - ETYMOLOGY EXTRACTION FIXES")
        print("=" * 80)

        docx_files = sorted(Path(docx_dir).glob('*.docx'))
        print(f"\n🔄 Parsing {len(docx_files)} files...")

        for docx_file in docx_files:
            self.parse_document_with_tables(docx_file)

        # Add homonym numbering AFTER parsing all files
        self.add_homonym_numbers()

        return self.verbs

    def save_json(self, output_path):
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        total_examples = sum(
            sum(len(conj_data) for conj_data in stem['conjugations'].values())
            for verb in self.verbs
            for stem in verb['stems']
        )

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump({
                'verbs': self.verbs,
                'metadata': {
                    'total_verbs': len(self.verbs),
                    'total_stems': self.stats['stems_parsed'],
                    'total_examples': total_examples,
                    'homonyms_numbered': self.stats.get('homonyms_numbered', 0),
                    'contextual_roots': self.stats.get('contextual_roots', 0),
                    'parser_version': 'docx-v3-etymology-fixes'
                }
            }, f, ensure_ascii=False, indent=2)

        print(f"\n💾 Saved: {output_file}")
        print(f"   📊 {len(self.verbs)} verbs, {self.stats['stems_parsed']} stems, {total_examples} examples")
        if self.contextual_roots:
            print(f"   🔍 Contextually validated roots: {', '.join(self.contextual_roots[:10])}")
            if len(self.contextual_roots) > 10:
                print(f"   ... and {len(self.contextual_roots) - 10} more")

    def split_into_files(self, output_dir):
        """Split verbs into individual JSON files"""
        print(f"\n🔄 Splitting into individual files...")

        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        for f in output_path.glob('*.json'):
            f.unlink()

        for verb in self.verbs:
            root = verb['root']
            filename = f"{root}.json"
            filepath = output_path / filename

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(verb, f, ensure_ascii=False, indent=2)

        print(f"✅ Created {len(self.verbs)} individual verb files in {output_path}")

def main():
    parser = FixedDocxParser()
    parser.parse_all_files('.devkit/new-source-docx')

    parser.save_json('.devkit/analysis/docx_v3_parsed.json')
    parser.split_into_files('.devkit/analysis/docx_v3_verbs')

    print("\n" + "=" * 80)
    print("PARSING COMPLETE - NOW RUN VALIDATION")
    print("=" * 80)
    print(f"📚 Total verbs: {len(parser.verbs)}")
    print(f"📖 Total stems: {parser.stats['stems_parsed']}")
    print(f"📊 Examples: {parser.stats['examples_parsed']}")
    print(f"🔢 Homonyms numbered: {parser.stats.get('homonyms_numbered', 0)}")
    print(f"🔍 Contextual roots: {parser.stats.get('contextual_roots', 0)}")

    if parser.contextual_roots:
        print(f"\n✨ Recovered verbs via contextual validation:")
        for root in parser.contextual_roots[:15]:
            verb = next((v for v in parser.verbs if v['root'] == root), None)
            if verb:
                print(f"   {root}: {len(verb['stems'])} stems")
        if len(parser.contextual_roots) > 15:
            print(f"   ... and {len(parser.contextual_roots) - 15} more")

    print("\nNext: Run comprehensive_validation.py with new output")

if __name__ == '__main__':
    main()
