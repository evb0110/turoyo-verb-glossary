#!/usr/bin/env python3
"""Comprehensive verification: check BOTH table and non-table extraction for hyw 1"""

import json
import re
from docx import Document
from pathlib import Path

print("="*80)
print("PART 1: VERIFY TABLE EXTRACTION")
print("="*80)

# Extract from DOCX tables
doc = Document('.devkit/new-source-docx/3. h,ḥ,k.docx')

root_pattern = re.compile(r'^(hyw\s+1)\s*[<(]', re.UNICODE)
in_hyw1 = False
docx_tables = []

for el in doc.element.body:
    tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

    if tag == 'p':
        for para in doc.paragraphs:
            if para._element == el:
                text = para.text.strip()
                if root_pattern.match(text):
                    in_hyw1 = True
                elif in_hyw1 and re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə]+(?:\s+\d+)?)\s*[<(]', text, re.UNICODE):
                    in_hyw1 = False
                break

    elif tag == 'tbl' and in_hyw1:
        for table in doc.tables:
            if table._element == el:
                if table.rows and len(table.rows[0].cells) >= 2:
                    conj_type = table.rows[0].cells[0].text.strip()
                    content = table.rows[0].cells[1].text.strip()
                    docx_tables.append({
                        'type': conj_type,
                        'content': content,
                        'length': len(content)
                    })
                break

# Extract from JSON
with open('.devkit/analysis/docx_v2_verbs/hyw 1.json') as f:
    verb = json.load(f)

json_tables = []
for stem in verb['stems']:
    for conj_type, examples in stem['conjugations'].items():
        for example in examples:
            # Reconstruct the content
            parts = []
            if example.get('turoyo'):
                parts.append(example['turoyo'])
            if example.get('translations'):
                parts.extend(example['translations'])
            content = ' '.join(parts)
            json_tables.append({
                'type': conj_type,
                'content': content,
                'length': len(content)
            })

print(f"\nDOCX Tables: {len(docx_tables)}")
print(f"JSON Tables: {len(json_tables)}")

print(f"\n📊 DETAILED TABLE COMPARISON:")
print(f"\nDOCX:")
for i, t in enumerate(docx_tables, 1):
    print(f"{i:2}. {t['type']:15} ({t['length']:5} chars): {t['content'][:80]}...")

print(f"\nJSON:")
for i, t in enumerate(json_tables, 1):
    print(f"{i:2}. {t['type']:15} ({t['length']:5} chars): {t['content'][:80]}...")

# Calculate data loss
docx_total_chars = sum(t['length'] for t in docx_tables)
json_total_chars = sum(t['length'] for t in json_tables)
table_loss = docx_total_chars - json_total_chars
table_loss_pct = (table_loss / docx_total_chars * 100) if docx_total_chars > 0 else 0

print(f"\n⚖️  TABLE DATA:")
print(f"   DOCX total chars: {docx_total_chars}")
print(f"   JSON total chars: {json_total_chars}")
print(f"   Data loss: {table_loss} chars ({table_loss_pct:.1f}%)")

print("\n" + "="*80)
print("PART 2: VERIFY NON-TABLE EXTRACTION")
print("="*80)

# Extract non-table paragraphs from DOCX
doc = Document('.devkit/new-source-docx/3. h,ḥ,k.docx')

in_hyw1 = False
has_stem = False
stem_pattern = re.compile(r'^([IVX]+|Pa\.|Af\.|Št\.|Šaf\.):\s*', re.UNICODE)
docx_paragraphs = []

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    if root_pattern.match(text):
        in_hyw1 = True
        has_stem = False
        continue

    if in_hyw1 and re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə]+(?:\s+\d+)?)\s*[<(]', text, re.UNICODE):
        break

    if in_hyw1:
        if stem_pattern.match(text) or text.startswith('Detransitive'):
            has_stem = True
        elif has_stem and not para._element.getparent().tag.endswith('tbl'):
            # Non-table paragraph after stems
            if text not in ['Idiomatic phrases', 'Idioms:', 'Examples:']:
                # Skip numbered meaning lists
                if not re.match(r'^\d+\)\s+.+;\s*\d+\)\s+.+;', text):
                    docx_paragraphs.append(text)

# Get idioms from JSON
json_paragraphs = verb.get('idioms', []) or []

print(f"\nDOCX non-table paragraphs: {len(docx_paragraphs)}")
print(f"JSON idioms: {len(json_paragraphs)}")

print(f"\n📝 FIRST 5 DOCX PARAGRAPHS:")
for i, text in enumerate(docx_paragraphs[:5], 1):
    print(f"{i}. {text[:100]}...")

print(f"\n📝 FIRST 5 JSON IDIOMS:")
for i, text in enumerate(json_paragraphs[:5], 1):
    if isinstance(text, str):
        print(f"{i}. {text[:100]}...")
    else:
        print(f"{i}. [object] {str(text)[:100]}...")

# Calculate data loss
docx_para_chars = sum(len(p) for p in docx_paragraphs)
json_para_chars = sum(len(p) if isinstance(p, str) else len(str(p)) for p in json_paragraphs)
para_loss = docx_para_chars - json_para_chars
para_loss_pct = (para_loss / docx_para_chars * 100) if docx_para_chars > 0 else 0

print(f"\n⚖️  NON-TABLE DATA:")
print(f"   DOCX total chars: {docx_para_chars}")
print(f"   JSON total chars: {json_para_chars}")
print(f"   Data loss: {para_loss} chars ({para_loss_pct:.1f}%)")

print("\n" + "="*80)
print("OVERALL SUMMARY")
print("="*80)
total_docx = docx_total_chars + docx_para_chars
total_json = json_total_chars + json_para_chars
total_loss = total_docx - total_json
total_loss_pct = (total_loss / total_docx * 100) if total_docx > 0 else 0

print(f"TOTAL DOCX: {total_docx} chars")
print(f"TOTAL JSON: {total_json} chars")
print(f"TOTAL LOSS: {total_loss} chars ({total_loss_pct:.1f}%)")

if total_loss_pct > 10:
    print(f"\n❌ CRITICAL: {total_loss_pct:.1f}% data loss!")
elif total_loss_pct > 5:
    print(f"\n⚠️  WARNING: {total_loss_pct:.1f}% data loss")
else:
    print(f"\n✅ ACCEPTABLE: {total_loss_pct:.1f}% data loss")
