#!/usr/bin/env python3
"""
Complete DOCX parser with table association
Uses document element tree to properly link tables to stems
"""

import re
import json
from pathlib import Path
from docx import Document
from collections import defaultdict
import copy

class CompleteDocxParser:
    """Parse verbs from DOCX with proper table handling"""

    def __init__(self):
        self.verbs = []
        self.stats = defaultdict(int)

    def is_letter_header(self, para):
        return para.style and para.style.name == 'Heading 1'

    def is_root_paragraph(self, para):
        if not para.text.strip():
            return False
        has_italic = any(r.italic for r in para.runs)
        sizes = [r.font.size.pt for r in para.runs if r.font.size]
        has_11pt = 11.0 in sizes
        text = para.text.strip()
        turoyo_chars = 'ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə'
        has_root = re.match(f'^([{turoyo_chars}]{{2,6}})(\s+\d+)?', text)
        return has_italic and has_11pt and has_root

    def is_stem_header(self, para):
        if not para.text.strip():
            return False
        has_bold = any(r.bold for r in para.runs)
        has_italic = any(r.italic for r in para.runs)
        sizes = [r.font.size.pt for r in para.runs if r.font.size]
        has_14pt = 14.0 in sizes
        has_stem = re.match(r'^([IVX]+):\s*', para.text.strip())
        return has_bold and has_italic and has_14pt and has_stem

    def extract_root_and_etymology(self, text):
        match = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]+(?:\s+\d+)?)\s*\((.+?)\)', text)
        if match:
            root = match.group(1).strip()
            etym_text = match.group(2).strip().lstrip('<').strip()
            parts = etym_text.split(None, 1)
            etymology = {
                'etymons': [{
                    'source': parts[0] if parts else '',
                    'raw': etym_text
                }]
            } if parts else None
            return root, etymology

        root_match = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]+(?:\s+\d+)?)', text)
        if root_match:
            return root_match.group(1).strip(), None
        return None, None

    def extract_stem_info(self, text):
        match = re.match(r'^([IVX]+):\s*(.+)', text.strip())
        if match:
            stem_num = match.group(1)
            forms_text = match.group(2).strip()
            forms_match = re.match(r'^([^\s]+(?:/[^\s]+)*)', forms_text)
            if forms_match:
                forms_str = forms_match.group(1)
                forms = [f.strip() for f in forms_str.split('/') if f.strip()]
                return stem_num, forms
        return None, []

    def normalize_header(self, header):
        """Normalize conjugation headers to canonical keys"""
        if header is None:
            return []

        h = self.normalize_whitespace(header)
        if not h:
            return []

        mapping = {
            'Infectum - wa': 'Infectum-wa',
            'Infectum – wa': 'Infectum-wa',
            'Infectum — wa': 'Infectum-wa',
            'Infectum–wa': 'Infectum-wa',
            'Infectum — Transitive': 'Infectum-Transitive',
            ' Infectum': 'Infectum',
            'Imperativ': 'Imperative',
            'Infinitiv': 'Infinitive',
            'Preterite': 'Preterit',
            'Part act.': 'Participle_Active',
            'Part. act.': 'Participle_Active',
            'Part Act.': 'Participle_Active',
            'Part act': 'Participle_Active',
            'Part. Act.': 'Participle_Active',
            'Part act.': 'Participle_Active',
            'Part pass.': 'Participle_Passive',
            'Part. pass.': 'Participle_Passive',
            'Part. Pass.': 'Participle_Passive',
            'Part.Pass': 'Participle_Passive',
            'Part pass': 'Participle_Passive',
            'Pass. Part.': 'Participle_Passive',
            'Passive Part.': 'Participle_Passive',
            'Part': 'Participle',
            'Participle': 'Participle',
            'Nomen Patiens': 'Nomen Patiens',
            'Nomen Patientis?': 'Nomen Patiens',
            'Nomen Actionis': 'Nomen Actionis',
            'Nomen agentis': 'Nomen agentis',
        }

        normalized = mapping.get(h, h)
        if ' and ' in normalized:
            parts = [
                mapping.get(part.strip(), part.strip())
                for part in normalized.split(' and ')
                if part.strip()
            ]
            return parts

        return [normalized]

    def normalize_whitespace(self, text):
        if not text:
            return ""
        return re.sub(r'\s+', ' ', text).strip()

    def parse_table_cell(self, cell_text):
        """Parse examples from table cell text"""
        cell_text = cell_text.strip()
        if not cell_text:
            return []

        examples = []
        parts = re.split(r'(?:^|\n)\s*\d+\)\s*', cell_text)

        for part in parts:
            part = part.strip()
            if not part:
                continue

            translations = re.findall(r'[ʻ\'\"“”‘’]([^ʻ\'\"“”‘’]{3,})[ʼ\'\"“”‘’]', part)
            references = re.findall(r'\d+(?:/\d+)?', part)

            example = {
                'turoyo': self.normalize_whitespace(part),
                'translations': [self.normalize_whitespace(t) for t in translations if t.strip()],
                'references': [self.normalize_whitespace(r) for r in references[:3]]
            }

            examples.append(example)

        return examples

    def parse_document_with_tables(self, docx_path):
        """Parse document using element tree to get table positions"""
        print(f"\n📖 {docx_path.name}")

        doc = Document(docx_path)

        # Build element map: paragraph/table elements in order
        elements = []
        for el in doc.element.body:
            tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

            if tag == 'p':
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element == el:
                        elements.append(('para', para))
                        break
            elif tag == 'tbl':
                # Find corresponding table object
                for table in doc.tables:
                    if table._element == el:
                        elements.append(('table', table))
                        break

        # Now parse through elements sequentially
        current_verb = None
        current_stem = None
        collecting_gloss_for_stem = False
        collecting_idioms = False

        for elem_type, elem in elements:
            if elem_type == 'para':
                para = elem
                text = para.text.strip()

                # Letter header
                if self.is_letter_header(para):
                    continue

                # Root (new verb)
                if self.is_root_paragraph(para):
                    if current_verb:
                        idioms = current_verb.get('idioms')
                        if idioms == []:
                            current_verb['idioms'] = None
                        self.verbs.append(current_verb)
                        self.stats['verbs_parsed'] += 1

                    collecting_idioms = False
                    collecting_gloss_for_stem = False

                    root, etymology = self.extract_root_and_etymology(para.text)
                    if root:
                        current_verb = {
                            'root': root,
                            'etymology': etymology,
                            'cross_reference': None,
                            'stems': [],
                            'uncertain': '???' in para.text
                        }
                        current_stem = None

                # Stem header
                elif self.is_stem_header(para):
                    stem_num, forms = self.extract_stem_info(para.text)
                    if stem_num and current_verb is not None:
                        current_stem = {
                            'stem': stem_num,
                            'forms': forms,
                            'conjugations': {}
                        }
                        current_verb['stems'].append(current_stem)
                        self.stats['stems_parsed'] += 1
                        collecting_gloss_for_stem = True
                        collecting_idioms = False

                # Detransitive marker
                elif 'Detransitive' in para.text and current_verb:
                    if not any(s['stem'] == 'Detransitive' for s in current_verb['stems']):
                        note_text = para.text.split('Detransitive', 1)[1].strip() if 'Detransitive' in para.text else ''
                        note_text = note_text.lstrip(':-').strip()
                        current_stem = {
                            'stem': 'Detransitive',
                            'forms': [],
                            'conjugations': {}
                        }
                        if note_text:
                            current_stem['label_gloss_tokens'] = [{'italic': False, 'text': note_text}]
                        current_verb['stems'].append(current_stem)
                        self.stats['detransitive_entries'] += 1
                        collecting_gloss_for_stem = bool(note_text)
                        collecting_idioms = False

                elif not text:
                    if collecting_gloss_for_stem:
                        collecting_gloss_for_stem = False
                    continue

                elif current_verb is not None:
                    lowered = text.lower()
                    if lowered in {'idiomatic phrases', 'idiomatic phrases:', 'idioms', 'idioms:', 'collocations', 'collocations:'}:
                        if current_verb.get('idioms') is None:
                            current_verb['idioms'] = []
                        collecting_idioms = True
                        collecting_gloss_for_stem = False
                        continue

                    if collecting_idioms:
                        if current_verb.get('idioms') is None:
                            current_verb['idioms'] = []
                        current_verb['idioms'].append(text)
                        continue

                    if current_stem is not None and collecting_gloss_for_stem:
                        tokens = current_stem.setdefault('label_gloss_tokens', [])
                        if tokens:
                            tokens.append({'italic': False, 'text': ' '})
                        tokens.append({'italic': False, 'text': text})
                        continue

                    collecting_gloss_for_stem = False

            elif elem_type == 'table':
                table = elem

                if current_stem is not None and table.rows:
                    last_headers = None
                    for row in table.rows:
                        if len(row.cells) < 2:
                            continue

                        raw_header = row.cells[0].text
                        examples_text = row.cells[1].text

                        headers = self.normalize_header(raw_header) if raw_header.strip() else last_headers
                        if not headers:
                            continue

                        examples = self.parse_table_cell(examples_text)
                        if not examples:
                            last_headers = headers
                            continue

                        for header in headers:
                            existing = current_stem['conjugations'].setdefault(header, [])
                            existing.extend(copy.deepcopy(examples))
                            self.stats['examples_parsed'] += len(examples)

                        last_headers = headers

                collecting_gloss_for_stem = False
                collecting_idioms = False

        # Save last verb
        if current_verb:
            idioms = current_verb.get('idioms')
            if idioms == []:
                current_verb['idioms'] = None
            self.verbs.append(current_verb)
            self.stats['verbs_parsed'] += 1

        print(f"   ✓ {self.stats['verbs_parsed']} verbs, {self.stats['stems_parsed']} stems, {self.stats['examples_parsed']} examples")

    def parse_all_files(self, docx_dir):
        print("=" * 80)
        print("COMPLETE DOCX PARSER (with tables)")
        print("=" * 80)

        docx_files = sorted(Path(docx_dir).glob('*.docx'))
        print(f"\n🔄 Parsing {len(docx_files)} files...")

        for docx_file in docx_files:
            self.parse_document_with_tables(docx_file)

        return self.verbs

    def save_json(self, output_path):
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        total_examples = sum(
            sum(len(conj_data) for conj_data in stem['conjugations'].values())
            for verb in self.verbs
            for stem in verb['stems']
        )

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump({
                'verbs': self.verbs,
                'metadata': {
                    'total_verbs': len(self.verbs),
                    'total_stems': self.stats['stems_parsed'],
                    'total_examples': total_examples,
                    'parser_version': 'docx-complete-1.0.0'
                }
            }, f, ensure_ascii=False, indent=2)

        print(f"\n💾 Saved: {output_file}")
        print(f"   📊 {len(self.verbs)} verbs, {self.stats['stems_parsed']} stems, {total_examples} examples")

    def split_into_files(self, output_dir):
        """Split verbs into individual JSON files"""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        for f in output_path.glob('*.json'):
            f.unlink()

        for verb in self.verbs:
            root = verb['root']
            filename = f"{root}.json"
            filepath = output_path / filename

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(verb, f, ensure_ascii=False, indent=2)

        print(f"✅ Created {len(self.verbs)} individual verb files in {output_path}")

def main():
    parser = CompleteDocxParser()
    parser.parse_all_files('.devkit/new-source-docx')

    parser.save_json('.devkit/analysis/docx_complete_parsed.json')
    parser.split_into_files('.devkit/analysis/docx_verbs')
    parser.split_into_files('server/assets/verbs')

    print("\n" + "=" * 80)
    print("FINAL STATISTICS")
    print("=" * 80)
    print(f"📚 Total verbs: {len(parser.verbs)}")
    print(f"📖 Total stems: {parser.stats['stems_parsed']}")
    print(f"📊 Examples parsed: {parser.stats['examples_parsed']}")
    print(f"🔗 Detransitive entries: {parser.stats.get('detransitive_entries', 0)}")
    print("=" * 80)

    # Show sample
    if parser.verbs:
        print("\n📄 Sample verb (first with examples):")
        sample = next((v for v in parser.verbs if any(s['conjugations'] for s in v['stems'])), parser.verbs[0])
        print(f"Root: {sample['root']}")
        print(f"Stems: {len(sample['stems'])}")
        for stem in sample['stems'][:2]:
            print(f"\n  {stem['stem']}: {', '.join(stem['forms'])}")
            for conj_type, examples in list(stem['conjugations'].items())[:2]:
                print(f"    {conj_type}: {len(examples)} examples")
                if examples:
                    ex = examples[0]
                    print(f"      - {ex['turoyo'][:60]}...")
                    if ex['translations']:
                        print(f"        Translations: {ex['translations']}")

if __name__ == '__main__':
    main()
