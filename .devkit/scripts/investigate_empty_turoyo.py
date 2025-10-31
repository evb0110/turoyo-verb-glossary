#!/usr/bin/env python3
"""
Investigate Empty Turoyo Cases - Systematic DOCX Source Analysis

For each of the 26 empty Turoyo cases, examine DOCX source structure:
- Find the exact table cell in DOCX
- Analyze run.italic metadata
- Check paragraph structure
- Identify why parser failed
"""

import json
import os
import re
from pathlib import Path
from docx import Document
from collections import defaultdict

# Map roots to their DOCX files
ROOT_TO_FILE = {
    'ʔ': '1. ʔ, ʕ, b, č.docx',
    'ʕ': '1. ʔ, ʕ, b, č.docx',
    'b': '1. ʔ, ʕ, b, č.docx',
    'č': '1. ʔ, ʕ, b, č.docx',
    'd': '2. d, f, g, ġ, ǧ.docx',
    'f': '2. d, f, g, ġ, ǧ.docx',
    'g': '2. d, f, g, ġ, ǧ.docx',
    'ġ': '2. d, f, g, ġ, ǧ.docx',
    'ǧ': '2. d, f, g, ġ, ǧ.docx',
    'h': '3. h,ḥ,k.docx',
    'ḥ': '3. h,ḥ,k.docx',
    'k': '3. h,ḥ,k.docx',
    'l': '4. l,m,n,p.docx',
    'm': '4. l,m,n,p.docx',
    'n': '4. l,m,n,p.docx',
    'p': '4. l,m,n,p.docx',
    'q': '5. q,r,s,ṣ.docx',
    'r': '5. q,r,s,ṣ.docx',
    's': '5. q,r,s,ṣ.docx',
    'ṣ': '5. q,r,s,ṣ.docx',
    'š': '6. š,t,ṭ,ṯ.docx',
    't': '6. š,t,ṭ,ṯ.docx',
    'ṭ': '6. š,t,ṭ,ṯ.docx',
    'ṯ': '6. š,t,ṭ,ṯ.docx',
    'v': '7. v, w, x, y, z, ž.docx',
    'w': '7. v, w, x, y, z, ž.docx',
    'x': '7. v, w, x, y, z, ž.docx',
    'y': '7. v, w, x, y, z, ž.docx',
    'z': '7. v, w, x, y, z, ž.docx',
    'ž': '7. v, w, x, y, z, ž.docx',
}

def get_docx_file_for_root(root):
    """Determine which DOCX file contains this root"""
    first_char = root[0]
    return ROOT_TO_FILE.get(first_char, None)

def find_verb_in_docx(docx_path, target_root):
    """Find a specific verb in DOCX and return its structure"""
    doc = Document(docx_path)

    # Normalize target root (remove homonym numbers)
    target_base = re.sub(r'\s+\d+$', '', target_root)

    # Build element map
    elements = []
    for el in doc.element.body:
        tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

        if tag == 'p':
            for para in doc.paragraphs:
                if para._element == el:
                    elements.append(('para', para))
                    break
        elif tag == 'tbl':
            for table in doc.tables:
                if table._element == el:
                    elements.append(('table', table))
                    break

    # Find the verb
    current_verb = None
    current_stem = None
    verb_data = None

    for idx, (elem_type, elem) in enumerate(elements):
        if elem_type == 'para':
            para = elem
            text = para.text.strip()

            # Check if this is a root paragraph
            turoyo_with_combining = r'[ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə\u0300-\u036F]'
            has_root = re.match(rf'^({turoyo_with_combining}{{2,12}})(?:\s+\d+)?(?:\s|\(|<|$)', text)

            if has_root:
                root = has_root.group(1).strip()
                root_base = re.sub(r'\s+\d+$', '', root)

                if root_base == target_base or root == target_root:
                    verb_data = {
                        'root': root,
                        'text': text,
                        'stems': []
                    }
                    current_verb = root

            # Check for stem headers
            if current_verb and verb_data:
                stem_match = re.match(r'^([IVX]+):\s*', text)
                if stem_match:
                    current_stem = stem_match.group(1)
                    verb_data['stems'].append({
                        'stem': current_stem,
                        'text': text,
                        'conjugations': []
                    })

        elif elem_type == 'table' and current_stem and verb_data:
            table = elem
            if table.rows and len(table.rows) > 0:
                row = table.rows[0]
                if len(row.cells) >= 2:
                    conj_type = row.cells[0].text.strip()
                    examples_cell = row.cells[1]

                    # Analyze cell structure
                    cell_analysis = analyze_cell_structure(examples_cell)

                    verb_data['stems'][-1]['conjugations'].append({
                        'type': conj_type,
                        'cell_analysis': cell_analysis
                    })

    return verb_data

def analyze_cell_structure(cell):
    """Analyze the structure of a table cell paragraph by paragraph"""
    paragraphs_data = []

    for para in cell.paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue

        runs_data = []
        for run in para.runs:
            run_text = run.text
            if not run_text:
                continue

            runs_data.append({
                'text': run_text,
                'italic': run.italic,
                'bold': run.bold,
                'font_size': run.font.size.pt if run.font.size else None,
            })

        paragraphs_data.append({
            'full_text': para_text,
            'runs': runs_data
        })

    return paragraphs_data

def load_empty_cases():
    """Load all empty Turoyo cases from JSON"""
    verbs_dir = 'server/assets/verbs'
    empty_cases = []

    for filename in os.listdir(verbs_dir):
        if not filename.endswith('.json'):
            continue

        with open(os.path.join(verbs_dir, filename), 'r', encoding='utf-8') as f:
            verb = json.load(f)

        root = verb.get('root', '')

        for stem in verb.get('stems', []):
            stem_name = stem.get('stem', '')
            for conj_type, examples in stem.get('conjugations', {}).items():
                for ex in examples:
                    if not ex.get('turoyo', '').strip():
                        empty_cases.append({
                            'root': root,
                            'stem': stem_name,
                            'conjugation': conj_type,
                            'translation': ex.get('translations', [''])[0] if ex.get('translations') else '',
                            'reference': '; '.join(ex.get('references', []))
                        })

    return sorted(empty_cases, key=lambda x: x['root'])

def main():
    print("=" * 80)
    print("INVESTIGATING EMPTY TUROYO CASES")
    print("=" * 80)

    empty_cases = load_empty_cases()
    print(f"\n📋 Found {len(empty_cases)} empty Turoyo cases\n")

    docx_dir = Path('.devkit/new-source-docx')

    # Group by root for efficiency
    cases_by_root = defaultdict(list)
    for case in empty_cases:
        cases_by_root[case['root']].append(case)

    report_lines = []
    report_lines.append("# Empty Turoyo Investigation Report\n")
    report_lines.append(f"Total cases: {len(empty_cases)}\n")
    report_lines.append(f"Affected verbs: {len(cases_by_root)}\n\n")

    for root, cases in sorted(cases_by_root.items()):
        print(f"\n{'=' * 80}")
        print(f"ROOT: {root} ({len(cases)} empty cases)")
        print('=' * 80)

        report_lines.append(f"\n## {root} ({len(cases)} cases)\n\n")

        # Find DOCX file
        docx_filename = get_docx_file_for_root(root)
        if not docx_filename:
            print(f"⚠️  Could not determine DOCX file for root '{root}'")
            report_lines.append(f"⚠️  Could not determine DOCX file\n\n")
            continue

        docx_path = docx_dir / docx_filename
        if not docx_path.exists():
            print(f"⚠️  DOCX file not found: {docx_path}")
            report_lines.append(f"⚠️  DOCX file not found: {docx_path}\n\n")
            continue

        print(f"📖 Searching in: {docx_filename}")

        # Find verb in DOCX
        verb_data = find_verb_in_docx(docx_path, root)

        if not verb_data:
            print(f"⚠️  Verb '{root}' not found in DOCX")
            report_lines.append(f"⚠️  Verb not found in DOCX\n\n")
            continue

        print(f"✓ Found verb: {verb_data['root']}")
        print(f"  Stems: {len(verb_data['stems'])}")

        # For each empty case, find the matching conjugation
        for case in cases:
            stem_name = case['stem']
            conj_type = case['conjugation']

            print(f"\n  {stem_name} - {conj_type}")
            report_lines.append(f"### {stem_name} - {conj_type}\n\n")

            # Find matching stem
            matching_stem = None
            for stem in verb_data['stems']:
                if stem['stem'] == stem_name:
                    matching_stem = stem
                    break

            if not matching_stem:
                print(f"    ⚠️  Stem '{stem_name}' not found in DOCX data")
                report_lines.append(f"⚠️  Stem not found in DOCX data\n\n")
                continue

            # Find matching conjugation
            matching_conj = None
            for conj in matching_stem['conjugations']:
                if conj['type'] == conj_type:
                    matching_conj = conj
                    break

            if not matching_conj:
                print(f"    ⚠️  Conjugation '{conj_type}' not found in DOCX data")
                report_lines.append(f"⚠️  Conjugation not found in DOCX data\n\n")
                continue

            # Analyze the cell structure
            cell_analysis = matching_conj['cell_analysis']

            print(f"    📊 Cell structure: {len(cell_analysis)} paragraphs")
            report_lines.append(f"**JSON output:**\n")
            report_lines.append(f"- Translation: {case['translation'][:80]}\n")
            report_lines.append(f"- Reference: {case['reference']}\n\n")

            report_lines.append(f"**DOCX structure:** {len(cell_analysis)} paragraphs\n\n")

            for i, para_data in enumerate(cell_analysis, 1):
                print(f"      Para {i}: {len(para_data['runs'])} runs")
                print(f"        Full text: {para_data['full_text'][:100]}")

                report_lines.append(f"**Paragraph {i}:**\n")
                report_lines.append(f"```\n{para_data['full_text']}\n```\n\n")
                report_lines.append(f"**Runs ({len(para_data['runs'])}):**\n")

                for j, run in enumerate(para_data['runs'], 1):
                    print(f"          Run {j}: italic={run['italic']}, bold={run['bold']}, size={run['font_size']}")
                    print(f"            Text: {run['text'][:80]}")

                    report_lines.append(f"{j}. `italic={run['italic']}`, `bold={run['bold']}`, `size={run['font_size']}`\n")
                    report_lines.append(f"   ```\n   {run['text']}\n   ```\n")

                report_lines.append("\n")

            # Pattern analysis
            print("\n    🔍 Pattern Analysis:")
            report_lines.append("**Pattern Analysis:**\n")

            # Check for numbered list items
            has_numbered_list = any(re.match(r'^\d+\)', para['full_text']) for para in cell_analysis)
            if has_numbered_list:
                print("      ✓ Contains numbered list items (1), 2), etc.)")
                report_lines.append("- ✓ Contains numbered list items (1), 2), etc.)\n")

            # Check for italic=None runs
            none_italic_runs = [
                run for para in cell_analysis
                for run in para['runs']
                if run['italic'] is None
            ]
            if none_italic_runs:
                print(f"      ✓ Has {len(none_italic_runs)} runs with italic=None")
                report_lines.append(f"- ✓ Has {len(none_italic_runs)} runs with italic=None\n")

            # Check for form-only pattern
            has_form_only = any(
                re.match(r'^[a-zāēīūəʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓ][a-zāēīūəʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓ\s\-=]+;\s*\d+\s*;', para['full_text'], re.IGNORECASE)
                for para in cell_analysis
            )
            if has_form_only:
                print("      ✓ Matches form-only pattern (form; number;)")
                report_lines.append("- ✓ Matches form-only pattern (form; number;)\n")

            # Check for long story excerpt
            has_long_text = any(len(para['full_text']) > 200 for para in cell_analysis)
            if has_long_text:
                print("      ✓ Contains long text (> 200 chars) - likely story excerpt")
                report_lines.append("- ✓ Contains long text (> 200 chars) - likely story excerpt\n")

            report_lines.append("\n---\n\n")

    # Save report
    report_path = Path('.devkit/analysis/empty_turoyo_investigation.md')
    report_path.parent.mkdir(parents=True, exist_ok=True)

    with open(report_path, 'w', encoding='utf-8') as f:
        f.writelines(report_lines)

    print(f"\n{'=' * 80}")
    print(f"💾 Report saved: {report_path}")
    print('=' * 80)

if __name__ == '__main__':
    main()
