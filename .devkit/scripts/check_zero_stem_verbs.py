#!/usr/bin/env python3
"""Check verbs with 0 stems"""

import json
import re
from pathlib import Path
from docx import Document

def check_verb_in_docx(docx_path, target_root):
    """Check if a verb has stems in the DOCX source"""
    doc = Document(docx_path)

    elements = []
    for el in doc.element.body:
        tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

        if tag == 'p':
            for para in doc.paragraphs:
                if para._element == el:
                    elements.append(('para', para))
                    break
        elif tag == 'tbl':
            for table in doc.tables:
                if table._element == el:
                    elements.append(('table', table))
                    break

    for i, (elem_type, elem) in enumerate(elements):
        if elem_type == 'para':
            text = elem.text.strip()

            if text.startswith(f'{target_root} (') or text.startswith(f'{target_root}<'):
                print(f'\nFound {target_root} at element {i}:')
                print(f'  {text[:100]}')

                stem_count = 0
                for j in range(1, 15):
                    if i+j >= len(elements):
                        break

                    next_type, next_elem = elements[i+j]
                    if next_type == 'para':
                        next_text = next_elem.text.strip()
                        if not next_text:
                            continue

                        is_stem = re.match(r'^([IVX]+):\s*', next_text)
                        if is_stem:
                            stem_count += 1
                            print(f'  +{j}: STEM {is_stem.group(1)}: {next_text[:50]}')

                        next_root = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]{2,6})(?:\s+\d+)?(?:\s|\()', next_text)
                        if next_root and next_root.group(1) != target_root:
                            print(f'  +{j}: Next verb: {next_root.group(1)}')
                            break

                print(f'  Total stems in DOCX: {stem_count}')
                return stem_count

    print(f'\n{target_root} not found in {docx_path.name}')
    return None

docx_dir = Path('.devkit/new-source-docx')
output_dir = Path('.devkit/analysis/docx_v2_verbs')

zero_stem_verbs = []
for f in output_dir.glob('*.json'):
    with open(f, 'r', encoding='utf-8') as file:
        verb = json.load(file)
        if len(verb['stems']) == 0:
            zero_stem_verbs.append(verb['root'])

print(f'Total verbs with 0 stems in output: {len(zero_stem_verbs)}')
print('\nChecking first 5 in DOCX source:\n')

for root in sorted(zero_stem_verbs)[:5]:
    for docx_file in docx_dir.glob('*.docx'):
        result = check_verb_in_docx(docx_file, root)
        if result is not None:
            break
