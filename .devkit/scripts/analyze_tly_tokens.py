#!/usr/bin/env python3
"""
Analyze tokens in tly Imperativ to understand concatenation pattern
"""

import sys
sys.path.insert(0, 'parser')

from docx import Document
from parse_docx_production import FixedDocxParser

# Create parser instance
parser = FixedDocxParser()

# Load the DOCX file
doc = Document('.devkit/new-source-docx/6. š,t,ṭ,ṯ.docx')

# Build element map
elements = []
for el in doc.element.body:
    tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

    if tag == 'p':
        for para in doc.paragraphs:
            if para._element == el:
                elements.append(('para', para))
                break
    elif tag == 'tbl':
        for table in doc.tables:
            if table._element == el:
                elements.append(('table', table))
                break

# Find tly Imperativ
tly_elem_index = None
for i, (elem_type, elem) in enumerate(elements):
    if elem_type == 'para' and elem.text.strip().startswith('tly ('):
        tly_elem_index = i
        break

imperativ_table = None
for i in range(tly_elem_index + 1, min(tly_elem_index + 20, len(elements))):
    elem_type, elem = elements[i]
    if elem_type == 'table':
        row = elem.rows[0]
        if len(row.cells) >= 2:
            conj_type = row.cells[0].text.strip()
            if conj_type == 'Imperativ':
                imperativ_table = elem
                break

# Parse the cell paragraphs
row = imperativ_table.rows[0]
imperativ_cell = row.cells[1]

para = imperativ_cell.paragraphs[0]
para_text = para.text

# Use internal tokenizer
tokens = parser._split_raw_to_tokens(para_text)

# Convert text tokens to turoyo
for tkn in tokens:
    if tkn.get('kind') == 'text':
        tkn['kind'] = 'turoyo'

print('TOKEN ANALYSIS FOR TLY IMPERATIV')
print('='*80)
print(f'Total tokens: {len(tokens)}\n')

# Find the concatenation point (around example 5)
print('Tokens around example 5 (indices 80-130):')
print()

for i in range(80, min(130, len(tokens))):
    token = tokens[i]
    kind = token.get('kind', '')
    value = token.get('value', '')[:50]

    marker = ''
    # Check if this matches our split pattern
    if (i > 3 and
        tokens[i-4].get('kind') == 'translation' and
        tokens[i-3].get('kind') == 'punct' and tokens[i-3].get('value') == ';' and
        tokens[i-2].get('kind') == 'ref' and
        tokens[i-1].get('kind') == 'punct' and tokens[i-1].get('value') == ';' and
        tokens[i].get('kind') in ['turoyo', 'translation']):
        marker = ' <-- SPLIT POINT'

    print(f'{i:3d}: [{kind:12s}] {repr(value):55s} {marker}')

print('\n' + '='*80)
