#!/usr/bin/env python3
"""
INVESTIGATION: Why is pčq etymology truncated?
Expected: "prčq cf. Kurd. p'erçiqandin vt. (-p'erçiq-). 1) to crush, press, smash, squash, KED 107"
Actual: "prčq cf. Kurd. p'erçiqandin vt. (-p'erçiq-). 1"
"""

from docx import Document
import re

docx_path = '.devkit/new-source-docx/4. l,m,n,p.docx'
doc = Document(docx_path)

print("=" * 80)
print("INVESTIGATING pčq ETYMOLOGY TRUNCATION")
print("=" * 80)

found_pcq = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    if text.startswith('pčq'):
        found_pcq = True
        print(f"\n✅ FOUND pčq at paragraph {i}")
        print(f"Text: {repr(text)}")
        print(f"Length: {len(text)}")

        print("\n📋 NEXT 5 PARAGRAPHS:")
        for j in range(i+1, min(i+6, len(doc.paragraphs))):
            next_text = doc.paragraphs[j].text.strip()
            print(f"\n  Para {j}: {repr(next_text)}")
            print(f"  Length: {len(next_text)}")

            if doc.paragraphs[j].runs:
                print(f"  Runs:")
                for run in doc.paragraphs[j].runs:
                    print(f"    - {repr(run.text)} (italic={run.italic})")

        break

if not found_pcq:
    print("\n❌ pčq not found in document!")
    print("Searching for similar roots:")
    for para in doc.paragraphs:
        if 'pč' in para.text or 'prč' in para.text:
            print(f"  Found: {para.text[:100]}")
