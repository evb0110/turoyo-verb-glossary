#!/usr/bin/env python3
"""
Search for the 25 genuinely missing verbs in DOCX source files.
Determines if they exist but failed to parse, or are genuinely absent.
"""

import re
from pathlib import Path
from docx import Document
from collections import defaultdict

class MissingVerbSearcher:
    """Search for missing verbs in DOCX files"""

    def __init__(self):
        self.missing_verbs = [
            ('ḏyr', 11, 'HIGH PRIORITY'),
            ('šry', 3, 'High value'),
            ('th', 3, 'High value'),
            ('gʕgs', 2, ''),
            ('qrmč', 2, ''),
            ('šbḥ', 2, ''),
            ('wṣy 1407', 2, 'Has reference number'),
            ('znzl', 2, ''),
            ('bzr', 1, ''),
            ('kšy', 1, ''),
            ('kətl', 1, ''),
            ('nḥt 2', 1, 'Has homonym number'),
            ('rhm', 1, ''),
            ('rǧ', 1, ''),
            ('tkn', 1, ''),
            ('tky', 1, ''),
            ('tzbṭ', 1, ''),
            ('zwq', 1, ''),
            ('zyr 2', 1, 'Has homonym number'),
            ('šym', 1, ''),
            ('ʕr', 1, ''),
            ('ʕtr', 1, ''),
            ('ʕw', 1, ''),
            ('ḏyʕ', 1, ''),
            ('ṣlṭ', 1, ''),
        ]

        # Map roots to expected DOCX files
        self.file_mapping = {
            '1. ʔ, ʕ, b, č.docx': ['ʕr', 'ʕtr', 'ʕw', 'bzr'],
            '2. d, f, g, ġ, ǧ.docx': ['ḏyr', 'ḏyʕ', 'gʕgs'],
            '3. h,ḥ,k.docx': ['kšy', 'kətl'],
            '4. l,m,n,p.docx': ['nḥt 2'],
            '5. q,r,s,ṣ.docx': ['qrmč', 'rhm', 'rǧ', 'ṣlṭ'],
            '6. š,t,ṭ,ṯ.docx': ['šry', 'šbḥ', 'šym', 'th', 'tkn', 'tky', 'tzbṭ'],
            '7. v, w, x, y, z, ž.docx': ['wṣy 1407', 'znzl', 'zwq', 'zyr 2'],
        }

        self.results = {
            'found_in_docx': [],
            'not_in_docx': [],
            'different_format': [],
        }

    def clean_root_for_search(self, root):
        """Clean root for searching - remove numbers and spaces"""
        # Extract just the root letters, without homonym numbers or references
        match = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]+)', root)
        if match:
            return match.group(1)
        return root

    def search_in_paragraph(self, para, root_clean, root_original):
        """Search for root in paragraph text"""
        text = para.text.strip()
        if not text:
            return None

        # Pattern 1: Exact match at start (with optional homonym number)
        # e.g., "ḏyr", "ḏyr 2", "wṣy 1407"
        pattern1 = f'^{re.escape(root_clean)}(\\s+\\d+)?\\b'
        if re.match(pattern1, text):
            return {
                'type': 'exact_match',
                'text': text[:100],
                'full_text': text
            }

        # Pattern 2: Root appears in text (could be in a form or cross-reference)
        if root_clean in text:
            return {
                'type': 'appears_in_text',
                'text': text[:100],
                'full_text': text
            }

        return None

    def is_root_paragraph(self, para):
        """Check if paragraph is likely a root heading"""
        if not para.text.strip():
            return False
        has_italic = any(r.italic for r in para.runs)
        sizes = [r.font.size.pt for r in para.runs if r.font.size]
        has_11pt = 11.0 in sizes
        return has_italic and has_11pt

    def search_docx_file(self, docx_path, roots_to_find):
        """Search a DOCX file for specific roots"""
        doc = Document(docx_path)
        findings = {}

        for root_original, stems, note in [v for v in self.missing_verbs if v[0] in roots_to_find]:
            root_clean = self.clean_root_for_search(root_original)

            # Search through all paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                result = self.search_in_paragraph(para, root_clean, root_original)

                if result:
                    # Check if it's a root paragraph
                    is_root = self.is_root_paragraph(para)

                    # Get context (previous and next paragraphs)
                    context_before = doc.paragraphs[para_idx - 1].text[:80] if para_idx > 0 else ''
                    context_after = doc.paragraphs[para_idx + 1].text[:80] if para_idx < len(doc.paragraphs) - 1 else ''

                    findings[root_original] = {
                        'root': root_original,
                        'stems': stems,
                        'note': note,
                        'found': True,
                        'match_type': result['type'],
                        'paragraph_text': result['text'],
                        'full_paragraph': result['full_text'],
                        'is_root_paragraph': is_root,
                        'context_before': context_before,
                        'context_after': context_after,
                        'paragraph_index': para_idx
                    }
                    break

            # If not found
            if root_original not in findings:
                findings[root_original] = {
                    'root': root_original,
                    'stems': stems,
                    'note': note,
                    'found': False
                }

        return findings

    def search_all_files(self, docx_dir):
        """Search all DOCX files"""
        print("=" * 80)
        print("SEARCHING FOR 25 MISSING VERBS IN DOCX FILES")
        print("=" * 80)

        all_findings = {}

        for filename, roots in self.file_mapping.items():
            docx_path = Path(docx_dir) / filename

            if not docx_path.exists():
                print(f"\n⚠️  File not found: {filename}")
                continue

            print(f"\n📖 Searching: {filename}")
            print(f"   Looking for: {', '.join(roots)}")

            findings = self.search_docx_file(docx_path, roots)

            # Print immediate results
            for root, result in findings.items():
                if result['found']:
                    match_type = result['match_type']
                    is_root = "✅ ROOT PARA" if result['is_root_paragraph'] else "⚠️  Not root para"
                    print(f"   ✓ Found: {root} ({result['stems']} stems) - {match_type} - {is_root}")
                else:
                    print(f"   ✗ NOT FOUND: {root} ({result['stems']} stems)")

            all_findings.update(findings)

        # Categorize results
        for root, result in all_findings.items():
            if result['found']:
                if result['is_root_paragraph'] and result['match_type'] == 'exact_match':
                    self.results['found_in_docx'].append(result)
                else:
                    self.results['different_format'].append(result)
            else:
                self.results['not_in_docx'].append(result)

        return all_findings

    def analyze_parser_failures(self):
        """Analyze why parser failed for found verbs"""
        print("\n" + "=" * 80)
        print("ANALYZING PARSER FAILURES")
        print("=" * 80)

        for result in self.results['found_in_docx']:
            print(f"\n❌ PARSER FAILED: {result['root']} ({result['stems']} stems)")
            print(f"   Match type: {result['match_type']}")
            print(f"   Is root para: {result['is_root_paragraph']}")
            print(f"   Paragraph: {result['paragraph_text']}")

            # Check for parser detection issues
            full_text = result['full_paragraph']

            # Check for special characters
            has_numbers = bool(re.search(r'\d+', result['root']))
            has_spaces = ' ' in result['root']

            issues = []
            if has_numbers:
                issues.append("Has reference/homonym number")
            if has_spaces:
                issues.append("Has spaces in root")
            if '???' in full_text:
                issues.append("Marked as uncertain (???)")
            if not result['is_root_paragraph']:
                issues.append("Not detected as root paragraph (formatting issue)")

            if issues:
                print(f"   Possible issues: {', '.join(issues)}")

    def generate_report(self, output_path):
        """Generate comprehensive markdown report"""
        found_count = len(self.results['found_in_docx'])
        different_count = len(self.results['different_format'])
        not_found_count = len(self.results['not_in_docx'])

        report = [
            "# Missing Verbs DOCX Search Results\n",
            "## Executive Summary\n",
            f"**Total missing verbs investigated: 25**\n",
            f"**Found in DOCX (parser failed): {found_count} verbs**\n",
            f"**Found but different format: {different_count} verbs**\n",
            f"**Not found in DOCX: {not_found_count} verbs**\n",
            "\n---\n",
        ]

        # Summary table
        report.append("\n## Summary Table\n")
        report.append("| Category | Count | Percentage | Action Required |\n")
        report.append("|----------|-------|------------|----------------|\n")
        report.append(f"| Found in DOCX (parser failed) | {found_count} | {found_count/25*100:.1f}% | Fix parser |\n")
        report.append(f"| Different format | {different_count} | {different_count/25*100:.1f}% | Investigate |\n")
        report.append(f"| Not in DOCX | {not_found_count} | {not_found_count/25*100:.1f}% | Source difference |\n")
        report.append("\n---\n")

        # High-priority verbs
        report.append("\n## High-Priority Verbs (Most Stems)\n")
        high_priority = sorted(
            [v for v in self.missing_verbs if v[1] >= 3],
            key=lambda x: x[1],
            reverse=True
        )

        for root, stems, note in high_priority:
            # Find result
            result = None
            category = None
            for r in self.results['found_in_docx']:
                if r['root'] == root:
                    result = r
                    category = "Found in DOCX (parser failed)"
                    break
            if not result:
                for r in self.results['different_format']:
                    if r['root'] == root:
                        result = r
                        category = "Different format"
                        break
            if not result:
                for r in self.results['not_in_docx']:
                    if r['root'] == root:
                        result = r
                        category = "NOT in DOCX"
                        break

            report.append(f"\n### {root} ({stems} stems) - {category}\n")

            if result and result['found']:
                report.append(f"**Match type:** {result['match_type']}\n")
                report.append(f"**Is root paragraph:** {result['is_root_paragraph']}\n")
                report.append(f"**Paragraph text:**\n```\n{result['full_paragraph']}\n```\n")

                if result['context_before']:
                    report.append(f"**Context before:** {result['context_before']}\n")
                if result['context_after']:
                    report.append(f"**Context after:** {result['context_after']}\n")
            else:
                report.append("**Status:** NOT FOUND in expected DOCX file\n")
                report.append("**Implication:** May be legitimately missing from DOCX source\n")

        report.append("\n---\n")

        # Category 1: Found in DOCX (parser failed)
        if self.results['found_in_docx']:
            report.append("\n## Category 1: Found in DOCX (Parser Failed)\n")
            report.append(f"**Count:** {found_count} verbs\n")
            report.append(f"**Total stems:** {sum(r['stems'] for r in self.results['found_in_docx'])} stems\n")
            report.append("\n### Detailed Findings\n")

            for result in sorted(self.results['found_in_docx'], key=lambda x: x['stems'], reverse=True):
                report.append(f"\n#### {result['root']} ({result['stems']} stems)\n")
                report.append(f"- **Match type:** {result['match_type']}\n")
                report.append(f"- **Root paragraph:** {result['is_root_paragraph']}\n")
                report.append(f"- **Paragraph:** {result['paragraph_text']}\n")

                # Identify issues
                issues = []
                if re.search(r'\d+', result['root']):
                    issues.append("Contains number (reference/homonym)")
                if not result['is_root_paragraph']:
                    issues.append("Not detected as root paragraph (formatting)")
                if '???' in result['full_paragraph']:
                    issues.append("Marked as uncertain (???)")

                if issues:
                    report.append(f"- **Likely issues:** {', '.join(issues)}\n")

        # Category 2: Different format
        if self.results['different_format']:
            report.append("\n## Category 2: Found but Different Format\n")
            report.append(f"**Count:** {different_count} verbs\n")
            report.append("\n### Detailed Findings\n")

            for result in self.results['different_format']:
                report.append(f"\n#### {result['root']} ({result['stems']} stems)\n")
                report.append(f"- **Match type:** {result['match_type']}\n")
                report.append(f"- **Root paragraph:** {result['is_root_paragraph']}\n")
                report.append(f"- **Paragraph:** {result['paragraph_text']}\n")

        # Category 3: Not found in DOCX
        if self.results['not_in_docx']:
            report.append("\n## Category 3: Not Found in DOCX\n")
            report.append(f"**Count:** {not_found_count} verbs\n")
            report.append(f"**Total stems:** {sum(r['stems'] for r in self.results['not_in_docx'])} stems\n")
            report.append("\n**Implication:** These verbs exist in the original HTML-based dataset but not in the DOCX source files.\n")
            report.append("This represents a genuine source data difference, not a parser issue.\n")
            report.append("\n### List\n")

            for result in sorted(self.results['not_in_docx'], key=lambda x: x['stems'], reverse=True):
                report.append(f"- **{result['root']}** ({result['stems']} stems)")
                if result['note']:
                    report.append(f" - {result['note']}")
                report.append("\n")

        # Recommendations
        report.append("\n---\n")
        report.append("\n## Recommendations\n")

        if found_count > 0:
            report.append(f"\n### Priority 1: Fix Parser for {found_count} Verbs\n")
            report.append("These verbs exist in DOCX but the parser failed to extract them.\n")
            report.append("\n**Root causes to address:**\n")

            # Analyze common issues
            issues_count = defaultdict(int)
            for result in self.results['found_in_docx']:
                if re.search(r'\d+', result['root']):
                    issues_count['Has reference/homonym number'] += 1
                if not result['is_root_paragraph']:
                    issues_count['Not detected as root paragraph'] += 1
                if '???' in result['full_paragraph']:
                    issues_count['Marked as uncertain'] += 1

            for issue, count in sorted(issues_count.items(), key=lambda x: x[1], reverse=True):
                report.append(f"- {issue}: {count} verbs\n")

            total_stems = sum(r['stems'] for r in self.results['found_in_docx'])
            report.append(f"\n**Impact:** {total_stems} stems of data can be recovered\n")
            report.append(f"**Effort:** Medium - requires parser regex/logic improvements\n")

        if not_found_count > 0:
            report.append(f"\n### Priority 2: Document Source Differences ({not_found_count} Verbs)\n")
            report.append("These verbs are not in DOCX source files.\n")
            total_stems = sum(r['stems'] for r in self.results['not_in_docx'])
            report.append(f"**Impact:** {total_stems} stems cannot be recovered from DOCX\n")
            report.append(f"**Action:** Document as legitimate source difference\n")

        # Accuracy impact
        report.append("\n### Accuracy Impact\n")
        recoverable_stems = sum(r['stems'] for r in self.results['found_in_docx'])
        unrecoverable_stems = sum(r['stems'] for r in self.results['not_in_docx'])

        report.append(f"- **Recoverable:** {recoverable_stems} stems ({found_count} verbs)\n")
        report.append(f"- **Unrecoverable:** {unrecoverable_stems} stems ({not_found_count} verbs)\n")
        report.append(f"- **Total impact:** {recoverable_stems + unrecoverable_stems} stems across 25 verbs\n")

        # Write report
        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(report)

        print(f"\n📄 Report saved: {output_path}")

def main():
    searcher = MissingVerbSearcher()

    # Search all files
    findings = searcher.search_all_files('.devkit/new-source-docx')

    # Analyze failures
    searcher.analyze_parser_failures()

    # Generate report
    searcher.generate_report('.devkit/analysis/MISSING_VERBS_DOCX_SEARCH.md')

    # Print summary
    print("\n" + "=" * 80)
    print("SEARCH COMPLETE")
    print("=" * 80)
    print(f"✅ Found in DOCX (parser failed): {len(searcher.results['found_in_docx'])} verbs")
    print(f"⚠️  Different format: {len(searcher.results['different_format'])} verbs")
    print(f"❌ Not in DOCX: {len(searcher.results['not_in_docx'])} verbs")

if __name__ == '__main__':
    main()
