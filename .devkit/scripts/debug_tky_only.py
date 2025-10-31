#!/usr/bin/env python3
"""
Debug parser for tky issue
"""

import re
from docx import Document

def is_letter_header(para):
    return para.style and para.style.name == 'Heading 1'

def is_root_paragraph(para):
    if not para.text.strip():
        return False
    has_italic = any(r.italic for r in para.runs)
    sizes = [r.font.size.pt for r in para.runs if r.font.size]
    has_11pt = 11.0 in sizes
    text = para.text.strip()
    turoyo_chars = r'ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə'
    has_root = re.match(f'^([{turoyo_chars}]{{2,6}})(?:\s+\d+)?(?:\s|\(|$)', text)
    is_cross_ref = bool(re.search(r'→|See\s+[ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]', text))
    return has_italic and has_11pt and has_root and not is_cross_ref

def is_stem_header(para):
    if not para.text.strip():
        return False
    has_bold = any(r.bold for r in para.runs)
    has_italic = any(r.italic for r in para.runs)
    sizes = [r.font.size.pt for r in para.runs if r.font.size]
    has_14pt = 14.0 in sizes
    has_stem = re.match(r'^([IVX]+):\s*', para.text.strip())
    return has_bold and has_italic and has_14pt and has_stem

def extract_root(text):
    root_match = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]{2,6}(?:\s+\d+)?)(?:\s|\(|$)', text)
    if root_match:
        return root_match.group(1).strip()
    return None

# Parse the document
doc = Document('.devkit/new-source-docx/6. š,t,ṭ,ṯ.docx')

elements = []
for el in doc.element.body:
    tag = el.tag.split('}')[1] if '}' in el.tag else el.tag
    
    if tag == 'p':
        for para in doc.paragraphs:
            if para._element == el:
                elements.append(('para', para))
                break
    elif tag == 'tbl':
        for table in doc.tables:
            if table._element == el:
                elements.append(('table', table))
                break

current_verb = None
verbs_saved = []
in_tky_region = False

for elem_type, elem in elements:
    if elem_type == 'para':
        para = elem
        
        # Start tracking when we hit tky
        if 'tky' in para.text and not in_tky_region:
            in_tky_region = True
            print('\n=== ENTERING TKY REGION ===\n')
        
        # Stop after tkyf region
        if in_tky_region and 'tkyt' in para.text:
            print('\n=== LEAVING TKY REGION ===\n')
            break
        
        if not in_tky_region:
            continue
        
        print(f'Para: {para.text.strip()[:60]}')
        
        if is_letter_header(para):
            print('  → SKIP: letter header')
            continue
        
        if is_root_paragraph(para):
            root = extract_root(para.text)
            print(f'  → ROOT PARAGRAPH: "{root}"')
            
            if current_verb:
                print(f'  → SAVING previous verb: {current_verb["root"]} ({len(current_verb["stems"])} stems)')
                verbs_saved.append(current_verb)
            
            current_verb = {
                'root': root,
                'stems': []
            }
            print(f'  → Created new current_verb: "{root}"')
            
        elif is_stem_header(para):
            stem_match = re.match(r'^([IVX]+):', para.text.strip())
            stem_num = stem_match.group(1) if stem_match else None
            print(f'  → STEM HEADER: {stem_num}')
            
            if current_verb is not None:
                current_verb['stems'].append({'stem': stem_num})
                print(f'  → Added stem to "{current_verb["root"]}" (now {len(current_verb["stems"])} stems)')
            else:
                print(f'  → ERROR: current_verb is None!')
        
        print()

if current_verb:
    print(f'\nFinal: Saving current_verb: {current_verb["root"]} ({len(current_verb["stems"])} stems)')
    verbs_saved.append(current_verb)

print(f'\n=== RESULTS ===')
print(f'Verbs saved: {len(verbs_saved)}')
for v in verbs_saved:
    print(f'  - {v["root"]}: {len(v["stems"])} stems')
