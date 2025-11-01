#!/usr/bin/env python3
"""Extract table content from hyw 1 to verify"""

from docx import Document
import re

doc = Document('.devkit/new-source-docx/3. h,ḥ,k.docx')

root_pattern = re.compile(r'^(hyw\s+1)\s*[<(]', re.UNICODE)

in_hyw1 = False
tables_found = []

# Track elements
for el in doc.element.body:
    tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

    if tag == 'p':
        for para in doc.paragraphs:
            if para._element == el:
                text = para.text.strip()
                if root_pattern.match(text):
                    in_hyw1 = True
                    print(f"✓ Found hyw 1 root")
                elif in_hyw1 and re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə]+(?:\s+\d+)?)\s*[<(]', text, re.UNICODE):
                    print(f"✓ Found next root, stopping")
                    break
                break

    elif tag == 'tbl' and in_hyw1:
        for table in doc.tables:
            if table._element == el:
                print(f"\n📊 TABLE FOUND ({len(table.rows)} rows x {len(table.columns)} cols)")

                # Show table content
                for i, row in enumerate(table.rows):
                    if len(row.cells) >= 2:
                        col1 = row.cells[0].text.strip()
                        col2 = row.cells[1].text.strip()

                        if col1 and col2:
                            print(f"  Row {i}: [{col1}] → {col2[:80]}...")

                tables_found.append(table)
                break

    if in_hyw1 and len(tables_found) >= 10:  # Stop after finding enough tables
        break

print(f"\n✅ Total tables found for hyw 1: {len(tables_found)}")
