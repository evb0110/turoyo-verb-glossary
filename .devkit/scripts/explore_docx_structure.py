#!/usr/bin/env python3
"""
Explore DOCX structure to understand how to parse verb data
"""

from docx import Document
from docx.shared import Pt
from pathlib import Path

def analyze_paragraph(para, max_text=100):
    """Analyze a paragraph's structure"""
    text = para.text[:max_text]

    result = {
        'text': text,
        'style': para.style.name if para.style else 'None',
        'runs': len(para.runs),
        'bold': any(r.bold for r in para.runs),
        'italic': any(r.italic for r in para.runs),
    }

    # Check font sizes
    sizes = []
    for run in para.runs:
        if run.font.size:
            sizes.append(run.font.size.pt)

    if sizes:
        result['font_sizes'] = list(set(sizes))

    return result

def main():
    docx_path = Path('.devkit/new-source-docx/1. ʔ, ʕ, b, č.docx')

    print("=" * 80)
    print(f"ANALYZING: {docx_path.name}")
    print("=" * 80)

    doc = Document(docx_path)

    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")

    print("\n" + "=" * 80)
    print("FIRST 30 PARAGRAPHS")
    print("=" * 80)

    for i, para in enumerate(doc.paragraphs[:30]):
        info = analyze_paragraph(para)

        if info['text'].strip():
            print(f"\n[{i}] Style: {info['style']}")
            print(f"    Text: {info['text']}")
            print(f"    Bold: {info['bold']}, Italic: {info['italic']}, Runs: {info['runs']}")
            if 'font_sizes' in info:
                print(f"    Font sizes: {info['font_sizes']}")

    print("\n" + "=" * 80)
    print("LOOKING FOR VERB ROOTS (Turoyo letters)")
    print("=" * 80)

    turoyo_chars = 'ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə'

    root_candidates = []
    for i, para in enumerate(doc.paragraphs[:100]):
        text = para.text.strip()
        # Look for short sequences of turoyo characters
        if (2 <= len(text) <= 10 and
            all(c in turoyo_chars or c.isspace() or c.isdigit() for c in text)):
            info = analyze_paragraph(para, max_text=50)
            root_candidates.append((i, text, info))

    print(f"\nFound {len(root_candidates)} potential root paragraphs:")
    for idx, text, info in root_candidates[:15]:
        print(f"[{idx}] '{text}' - Style: {info['style']}, Bold: {info['bold']}")

    print("\n" + "=" * 80)
    print("LOOKING FOR STEM HEADERS (Roman numerals)")
    print("=" * 80)

    import re
    stem_pattern = re.compile(r'^([IVX]+):')

    stems = []
    for i, para in enumerate(doc.paragraphs[:200]):
        text = para.text.strip()
        match = stem_pattern.match(text)
        if match:
            info = analyze_paragraph(para, max_text=100)
            stems.append((i, text, info))

    print(f"\nFound {len(stems)} potential stem headers:")
    for idx, text, info in stems[:10]:
        print(f"[{idx}] '{text}'")
        print(f"      Style: {info['style']}, Bold: {info['bold']}, Sizes: {info.get('font_sizes', [])}")

    print("\n" + "=" * 80)
    print("ANALYZING TABLES")
    print("=" * 80)

    for i, table in enumerate(doc.tables[:3]):
        print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} cols")

        # Show first row
        if table.rows:
            first_row = table.rows[0]
            print(f"  First row cells:")
            for j, cell in enumerate(first_row.cells[:3]):
                text = cell.text[:60].replace('\n', ' ')
                print(f"    [{j}] {text}")

if __name__ == '__main__':
    main()
