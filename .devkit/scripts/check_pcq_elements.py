#!/usr/bin/env python3
import re
from pathlib import Path
from docx import Document

doc = Document('.devkit/new-source-docx/4. l,m,n,p.docx')

elements = []
for el in doc.element.body:
    tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

    if tag == 'p':
        for para in doc.paragraphs:
            if para._element == el:
                elements.append(('para', para))
                break
    elif tag == 'tbl':
        for table in doc.tables:
            if table._element == el:
                elements.append(('table', table))
                break

in_verb = False
count = 0

for i, (elem_type, elem) in enumerate(elements):
    if elem_type == 'para':
        text = elem.text.strip()
        if not text:
            continue

        if re.match(r'^pčq(\s|\()', text):
            in_verb = True
            print(f'{i}: ROOT: {text[:100]}')

        if in_verb:
            count += 1
            if count > 30:
                break

            has_bold = any(r.bold for r in elem.runs)
            has_italic = any(r.italic for r in elem.runs)
            sizes = [r.font.size.pt for r in elem.runs if r.font.size]

            stem_match = re.match(r'^([IVX]+):', text)
            detrans = 'Detransitive' in text

            print(f'{i} PARA: {text[:80]} | bold={has_bold} italic={has_italic} 14pt={14.0 in sizes} stem={stem_match.group(1) if stem_match else None} detrans={detrans}')

            next_root = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]{2,6})(\s+\d+|\s+\()', text)
            if next_root and count > 5 and next_root.group(1) != 'pčq':
                print(f'{i} PARA: NEXT VERB: {next_root.group(1)}')
                break

    elif elem_type == 'table' and in_verb:
        count += 1
        if count > 30:
            break
        print(f'{i} TABLE: {len(elem.rows)} rows, {len(elem.columns)} cols')
