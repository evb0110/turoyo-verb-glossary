#!/usr/bin/env python3
"""
Analyze Specific Empty Turoyo Cases

Focus on high-priority cases to understand patterns
"""

import re
from pathlib import Path
from docx import Document

def analyze_paragraph_runs(para):
    """Analyze runs in a paragraph"""
    print(f"\n    Full text: {para.text[:100]}")

    for i, run in enumerate(para.runs, 1):
        if not run.text:
            continue
        print(f"      Run {i}: italic={run.italic}, bold={run.bold}")
        print(f"         Text: [{run.text[:60]}]")

def find_verb_table(doc_path, root_text, stem_text, conj_text):
    """Find specific verb conjugation table in DOCX"""
    doc = Document(doc_path)

    print(f"\n{'='*80}")
    print(f"Searching for: {root_text} - {stem_text} - {conj_text}")
    print(f"File: {doc_path.name}")
    print('='*80)

    # Build elements
    elements = []
    for el in doc.element.body:
        tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

        if tag == 'p':
            for para in doc.paragraphs:
                if para._element == el:
                    elements.append(('para', para))
                    break
        elif tag == 'tbl':
            for table in doc.tables:
                if table._element == el:
                    elements.append(('table', table))
                    break

    # Find root
    found_root = False
    found_stem = False

    for idx, (elem_type, elem) in enumerate(elements):
        if elem_type == 'para':
            text = elem.text.strip()

            # Check for root
            if not found_root and root_text in text and len(text) < 200:
                print(f"\n✓ Found root paragraph:")
                print(f"  {text[:100]}")
                found_root = True
                continue

            # Check for stem
            if found_root and not found_stem and stem_text in text:
                print(f"\n✓ Found stem paragraph:")
                print(f"  {text[:100]}")
                found_stem = True
                continue

        elif elem_type == 'table' and found_root and found_stem:
            table = elem
            if table.rows and len(table.rows) > 0:
                row = table.rows[0]
                if len(row.cells) >= 2:
                    cell_conj_type = row.cells[0].text.strip()

                    if conj_text in cell_conj_type or cell_conj_type in conj_text:
                        print(f"\n✓ Found conjugation table: {cell_conj_type}")

                        # Analyze cell
                        examples_cell = row.cells[1]
                        print(f"\n  📊 Cell has {len(examples_cell.paragraphs)} paragraphs:")

                        for para_idx, para in enumerate(examples_cell.paragraphs, 1):
                            if not para.text.strip():
                                continue

                            print(f"\n  --- Paragraph {para_idx} ---")
                            analyze_paragraph_runs(para)

                        return True

    print("\n⚠️  Not found")
    return False

# Priority cases
cases = [
    # HIGH PRIORITY
    {
        'file': '6. š,t,ṭ,ṯ.docx',
        'root': 'šfr',
        'stem': 'I',
        'conj': 'Preterit Intransitive',
        'note': 'High priority - 2 empty cases'
    },
    {
        'file': '5. q,r,s,ṣ.docx',
        'root': 'qwy',
        'stem': 'III',
        'conj': 'Infectum',
        'note': 'High priority - 2 empty cases'
    },

    # MEDIUM PRIORITY - Different patterns
    {
        'file': '2. d, f, g, ġ, ǧ.docx',
        'root': 'dwy',
        'stem': 'II',
        'conj': 'Preterit',
        'note': 'Form-only pattern: mdawéle-li; 618;'
    },
    {
        'file': '4. l,m,n,p.docx',
        'root': 'nfq',
        'stem': 'III',
        'conj': 'Infinitiv',
        'note': 'Numbered list: 1) tufiqo; tefiqo; tawfoqo; 639;'
    },
    {
        'file': '4. l,m,n,p.docx',
        'root': 'ngl 2',
        'stem': 'II',
        'conj': 'Infinitiv',
        'note': 'Simple form: Negolo. (İlyas p.c)'
    },
    {
        'file': '1. ʔ, ʕ, b, č.docx',
        'root': 'byq',
        'stem': 'I',
        'conj': 'Preterit Intransitive',
        'note': 'Story excerpt - long text'
    },
    {
        'file': '5. q,r,s,ṣ.docx',
        'root': 'qbḏ',
        'stem': 'I',
        'conj': 'Preterit',
        'note': 'Story excerpt with italic=None'
    },
    {
        'file': '4. l,m,n,p.docx',
        'root': 'nṭr',
        'stem': 'I',
        'conj': 'Infectum',
        'note': 'Numbered list with mixed italic'
    },
]

def main():
    docx_dir = Path('.devkit/new-source-docx')

    print("="*80)
    print("EMPTY TUROYO CASE ANALYSIS")
    print("="*80)

    for case in cases:
        docx_path = docx_dir / case['file']
        print(f"\n\n{'#'*80}")
        print(f"# {case['note']}")
        print(f"{'#'*80}")

        find_verb_table(
            docx_path,
            case['root'],
            case['stem'],
            case['conj']
        )

    print("\n\n" + "="*80)
    print("ANALYSIS COMPLETE")
    print("="*80)

if __name__ == '__main__':
    main()
