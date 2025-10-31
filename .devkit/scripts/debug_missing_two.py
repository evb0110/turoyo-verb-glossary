#!/usr/bin/env python3
"""Debug the 2 still-missing etymologies"""

import re
from pathlib import Path
from docx import Document

roots_to_debug = {
    'gwlʕ': '2. d, f, g, ġ, ǧ.docx',
    'ḏyr': '2. d, f, g, ġ, ǧ.docx'
}

print("=" * 80)
print("DEBUGGING STILL-MISSING ETYMOLOGIES")
print("=" * 80)

for root, filename in roots_to_debug.items():
    docx_path = Path(f'.devkit/new-source-docx/{filename}')
    doc = Document(docx_path)

    print(f"\n{'=' * 80}")
    print(f"ROOT: {root}")
    print(f"{'=' * 80}")

    # Find the root paragraph
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith(root):
            print(f"\nParagraph {i}:")
            print(f"  Text: {text}")
            print(f"  Italic runs: {any(r.italic for r in para.runs)}")
            print(f"  Font sizes: {[r.font.size.pt for r in para.runs if r.font.size]}")

            # Check next paragraph
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                print(f"\nNext paragraph {i+1}:")
                print(f"  Text: {next_text}")

            # Try to parse with our patterns
            print(f"\nTrying pattern 1 (standard):")
            match1 = re.search(r'\(<\s*(.+?)\)(?:\s|$|;)', text)
            print(f"  Match: {match1.group(1) if match1 else 'NONE'}")

            print(f"\nTrying pattern 2 (missing opening paren):")
            match2 = re.search(r'(?:^|[\s\d])<\s*([A-Z][^<>]+?)\)(?:\s|$|;)', text)
            print(f"  Match: {match2.group(1) if match2 else 'NONE'}")

            print(f"\nTrying pattern 3 (alternative start):")
            match3 = re.search(r'\(((?:see|cf\.|unknown)[^)]+)\)', text)
            print(f"  Match: {match3.group(1) if match3 else 'NONE'}")

            # Check for unclosed paren
            print(f"\nChecking for unclosed paren:")
            open_count = text.count('(')
            close_count = text.count(')')
            print(f"  Opening parens: {open_count}, Closing parens: {close_count}")

            if open_count > close_count:
                print(f"  → UNCLOSED! Should combine with next paragraph")

            break

print("\n" + "=" * 80)
print("ANALYSIS COMPLETE")
print("=" * 80)
