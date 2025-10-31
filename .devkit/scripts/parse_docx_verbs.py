#!/usr/bin/env python3
"""
Parse Turoyo verbs from DOCX files
Much cleaner than HTML parsing!
"""

import re
import json
from pathlib import Path
from docx import Document
from collections import defaultdict

class DocxVerbParser:
    """Parse Turoyo verbs from DOCX files"""

    def __init__(self):
        self.verbs = []
        self.stats = defaultdict(int)
        self.current_letter = None

    def is_letter_header(self, para):
        """Check if paragraph is a letter header (Heading 1)"""
        return para.style and para.style.name == 'Heading 1'

    def is_root_paragraph(self, para):
        """
        Root paragraphs have:
        - Body Text style
        - Italic formatting
        - Font size 11pt
        - Start with Turoyo letters
        """
        if not para.text.strip():
            return False

        # Check for italic runs
        has_italic = any(r.italic for r in para.runs)
        if not has_italic:
            return False

        # Check font size (11pt is common for etymology)
        sizes = [r.font.size.pt for r in para.runs if r.font.size]
        has_11pt = 11.0 in sizes

        # Check for Turoyo characters at start
        text = para.text.strip()
        turoyo_chars = 'ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə'

        # Root should be 2-6 Turoyo characters, possibly followed by number
        root_pattern = f'^([{turoyo_chars}]{{2,6}})(\s+\d+)?'
        has_root = re.match(root_pattern, text)

        return has_italic and has_11pt and has_root

    def is_stem_header(self, para):
        """
        Stem headers have:
        - Bold + Italic
        - Font size 14pt
        - Start with Roman numeral followed by colon
        """
        if not para.text.strip():
            return False

        # Check formatting
        has_bold = any(r.bold for r in para.runs)
        has_italic = any(r.italic for r in para.runs)

        # Check font size
        sizes = [r.font.size.pt for r in para.runs if r.font.size]
        has_14pt = 14.0 in sizes

        # Check for roman numeral pattern
        text = para.text.strip()
        stem_pattern = r'^([IVX]+):\s*(.+)'
        has_stem = re.match(stem_pattern, text)

        return has_bold and has_italic and has_14pt and has_stem

    def extract_root_and_etymology(self, text):
        """Extract root and etymology from paragraph text"""
        # Pattern: root (< etymology)
        # Example: ʔbʕ (< MA bʕy cf. SL 169: ...)
        match = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]+(?:\s+\d+)?)\s*\((.+)\)', text)

        if match:
            root = match.group(1).strip()
            etym_text = match.group(2).strip()

            # Parse etymology (simplified)
            etymology = self.parse_etymology_simple(etym_text)

            return root, etymology

        # No etymology pattern, just extract root
        root_match = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]+(?:\s+\d+)?)', text)
        if root_match:
            return root_match.group(1).strip(), None

        return None, None

    def parse_etymology_simple(self, etym_text):
        """Simplified etymology parsing"""
        # Remove leading < if present
        etym_text = etym_text.lstrip('<').strip()

        # Try to extract source language and details
        parts = etym_text.split(None, 1)
        if len(parts) >= 1:
            return {
                'etymons': [{
                    'source': parts[0],
                    'raw': etym_text
                }]
            }

        return None

    def extract_stem_info(self, text):
        """Extract stem number and forms"""
        match = re.match(r'^([IVX]+):\s*(.+)', text)
        if match:
            stem_num = match.group(1)
            forms_text = match.group(2).strip()

            # Split forms by / or whitespace before English text
            # Example: "abəʕ/obəʕ" or "mʔadəb/miʔadəb  to teach"
            forms_match = re.match(r'^([^\s]+(?:/[^\s]+)*)', forms_text)
            if forms_match:
                forms_str = forms_match.group(1)
                forms = [f.strip() for f in forms_str.split('/') if f.strip()]
                return stem_num, forms

        return None, []

    def parse_table(self, table):
        """Parse a conjugation table"""
        if not table.rows or len(table.columns) < 2:
            return None

        row = table.rows[0]
        if len(row.cells) < 2:
            return None

        # Column 0: conjugation type (Preterit, Infectum, etc.)
        conj_type = row.cells[0].text.strip()

        # Column 1: examples (Turoyo text with translations)
        examples_text = row.cells[1].text.strip()

        # Parse examples
        examples = self.parse_examples(examples_text)

        return conj_type, examples

    def parse_examples(self, text):
        """Parse examples from table cell"""
        examples = []

        # Split by numbered items (1), 2), 3), etc.)
        parts = re.split(r'\n?\d+\)\s*', text)

        for part in parts:
            part = part.strip()
            if not part:
                continue

            # Look for German translations in quotes
            translations = re.findall(r'[ʻ\'\"]([\w\s,;.:!?äöüßÄÖÜ]+)[ʼ\'\"]', part)

            example = {
                'turoyo': part,
                'translations': [t.strip() for t in translations if len(t.strip()) > 3],
                'references': []
            }

            examples.append(example)

        return examples

    def parse_document(self, docx_path):
        """Parse a single DOCX file"""
        print(f"\n📖 Parsing: {docx_path.name}")

        doc = Document(docx_path)

        current_verb = None
        current_stem = None
        para_idx = 0

        while para_idx < len(doc.paragraphs):
            para = doc.paragraphs[para_idx]

            # Letter header
            if self.is_letter_header(para):
                self.current_letter = para.text.strip()
                print(f"   Letter: {self.current_letter}")
                para_idx += 1
                continue

            # Root paragraph (start of new verb)
            if self.is_root_paragraph(para):
                # Save previous verb
                if current_verb:
                    self.verbs.append(current_verb)
                    self.stats['verbs_parsed'] += 1

                # Extract root and etymology
                root, etymology = self.extract_root_and_etymology(para.text)

                if root:
                    current_verb = {
                        'root': root,
                        'etymology': etymology,
                        'cross_reference': None,
                        'stems': [],
                        'uncertain': '???' in para.text
                    }

                    current_stem = None

                para_idx += 1
                continue

            # Stem header
            if self.is_stem_header(para):
                stem_num, forms = self.extract_stem_info(para.text)

                if stem_num and current_verb is not None:
                    current_stem = {
                        'stem': stem_num,
                        'forms': forms,
                        'conjugations': {}
                    }
                    current_verb['stems'].append(current_stem)
                    self.stats['stems_parsed'] += 1

                para_idx += 1
                continue

            # Check for "Detransitive" marker
            if 'Detransitive' in para.text and current_verb:
                current_stem = {
                    'stem': 'Detransitive',
                    'forms': [],
                    'conjugations': {}
                }
                current_verb['stems'].append(current_stem)
                self.stats['detransitive_entries'] += 1

            para_idx += 1

        # Save last verb
        if current_verb:
            self.verbs.append(current_verb)
            self.stats['verbs_parsed'] += 1

        # Now parse tables and associate with stems
        self.associate_tables_with_stems(doc)

        print(f"   ✓ Parsed {self.stats['verbs_parsed']} verbs, {self.stats['stems_parsed']} stems")

    def associate_tables_with_stems(self, doc):
        """Associate tables with their stems (second pass)"""
        # This is tricky - tables appear after stem headers
        # We need to track which stem each table belongs to

        stem_positions = []
        for i, para in enumerate(doc.paragraphs):
            if self.is_stem_header(para):
                stem_num, _ = self.extract_stem_info(para.text)
                if stem_num:
                    stem_positions.append((i, stem_num))

        # Now find tables and assign to nearest previous stem
        # Note: In python-docx, tables are separate from paragraphs
        # We need to use the document's element tree

        # For now, just count tables as a sanity check
        table_count = len(doc.tables)
        self.stats['tables_found'] = table_count

        # TODO: Implement proper table-to-stem association
        # This requires walking the document's XML element tree

    def parse_all_files(self, docx_dir):
        """Parse all DOCX files in directory"""
        docx_files = sorted(Path(docx_dir).glob('*.docx'))

        print(f"🔄 Parsing {len(docx_files)} DOCX files...")

        for docx_file in docx_files:
            self.parse_document(docx_file)

        return self.verbs

    def save_json(self, output_path):
        """Save parsed verbs to JSON"""
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump({
                'verbs': self.verbs,
                'metadata': {
                    'total_verbs': len(self.verbs),
                    'total_stems': self.stats['stems_parsed'],
                    'tables_found': self.stats.get('tables_found', 0),
                    'parser_version': 'docx-1.0.0'
                }
            }, f, ensure_ascii=False, indent=2)

        print(f"\n💾 Saved: {output_file}")
        print(f"   📊 {len(self.verbs)} verbs, {self.stats['stems_parsed']} stems")

def main():
    print("=" * 80)
    print("DOCX VERB PARSER")
    print("=" * 80)

    parser = DocxVerbParser()
    parser.parse_all_files('.devkit/new-source-docx')

    parser.save_json('.devkit/analysis/docx_parsed.json')

    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)
    print(f"📚 Total verbs: {len(parser.verbs)}")
    print(f"📖 Total stems: {parser.stats['stems_parsed']}")
    print(f"📊 Tables found: {parser.stats.get('tables_found', 0)}")
    print("=" * 80)

    # Sample output
    if parser.verbs:
        print("\nSample verb:")
        sample = parser.verbs[0]
        print(json.dumps(sample, ensure_ascii=False, indent=2))

if __name__ == '__main__':
    main()
