#!/usr/bin/env python3
"""
Deep investigation of why parser fails to detect certain verbs.
Examines paragraph formatting in detail for high-priority missing verbs.
"""

from pathlib import Path
from docx import Document
import re

def analyze_paragraph_formatting(para, root_text):
    """Detailed formatting analysis of a paragraph"""
    print(f"\n{'='*80}")
    print(f"PARAGRAPH: {para.text[:100]}")
    print(f"{'='*80}")

    # Overall paragraph properties
    print(f"\nParagraph Style: {para.style.name if para.style else 'None'}")
    print(f"Paragraph Alignment: {para.alignment}")
    print(f"Number of runs: {len(para.runs)}")

    # Check if it matches root detection logic
    has_italic = any(r.italic for r in para.runs)
    sizes = [r.font.size.pt for r in para.runs if r.font.size]
    has_11pt = 11.0 in sizes
    text = para.text.strip()
    turoyo_chars = r'ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə'
    has_root = re.match(f'^([{turoyo_chars}]{{2,6}})(\s+\d+)?', text)

    print(f"\n✓ DETECTION CHECKS:")
    print(f"  Has italic runs: {has_italic}")
    print(f"  Font sizes present: {sorted(set(sizes)) if sizes else 'None'}")
    print(f"  Has 11pt font: {has_11pt}")
    print(f"  Matches root pattern: {bool(has_root)}")
    print(f"  Would be detected as root: {has_italic and has_11pt and has_root}")

    # Detailed run analysis
    print(f"\n📝 RUN DETAILS:")
    for i, run in enumerate(para.runs):
        if not run.text.strip():
            continue

        print(f"\n  Run {i+1}: '{run.text}'")
        print(f"    Italic: {run.italic}")
        print(f"    Bold: {run.bold}")
        print(f"    Font size: {run.font.size.pt if run.font.size else 'None'}")
        print(f"    Font name: {run.font.name}")
        print(f"    Underline: {run.underline}")

    return has_italic and has_11pt and has_root

def find_and_analyze_verb(docx_path, verb_root):
    """Find a specific verb and analyze its formatting"""
    print(f"\n\n{'#'*80}")
    print(f"# ANALYZING: {verb_root} in {docx_path.name}")
    print(f"{'#'*80}")

    doc = Document(docx_path)

    # Clean root for searching
    root_clean = re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭwxyzžḏṯẓāēīūə]+)', verb_root)
    if root_clean:
        root_clean = root_clean.group(1)
    else:
        root_clean = verb_root

    pattern = f'^{re.escape(root_clean)}(\\s+\\d+)?\\b'

    found = False
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(pattern, text):
            found = True
            would_detect = analyze_paragraph_formatting(para, verb_root)

            # Show surrounding context
            print(f"\n📍 CONTEXT:")
            if para_idx > 0:
                print(f"  Before: {doc.paragraphs[para_idx-1].text[:80]}")
            print(f"  >>> THIS: {text[:80]}")
            if para_idx < len(doc.paragraphs) - 1:
                print(f"  After: {doc.paragraphs[para_idx+1].text[:80]}")

            if not would_detect:
                print(f"\n❌ PARSER WOULD MISS THIS - Analyzing why...")

                # Specific diagnostics
                has_italic = any(r.italic for r in para.runs)
                sizes = [r.font.size.pt for r in para.runs if r.font.size]
                has_11pt = 11.0 in sizes

                issues = []
                if not has_italic:
                    issues.append("Missing italic formatting")
                if not has_11pt:
                    issues.append(f"Wrong font size (has: {sizes}, needs: 11pt)")
                if not any(r.italic for r in para.runs if r.text.strip()):
                    issues.append("Italic runs are empty/whitespace")

                print(f"\n🔍 ROOT CAUSE:")
                for issue in issues:
                    print(f"  - {issue}")
            else:
                print(f"\n✅ PARSER SHOULD DETECT THIS")

            break

    if not found:
        print(f"\n❌ NOT FOUND in document")

def main():
    # High-priority verbs to investigate
    investigations = [
        ('.devkit/new-source-docx/2. d, f, g, ġ, ǧ.docx', 'ḏyr'),  # 11 stems - HIGHEST
        ('.devkit/new-source-docx/6. š,t,ṭ,ṯ.docx', 'šry'),         # 3 stems - should work?
        ('.devkit/new-source-docx/6. š,t,ṭ,ṯ.docx', 'th'),          # 3 stems
        ('.devkit/new-source-docx/2. d, f, g, ġ, ǧ.docx', 'gʕgs'),  # 2 stems
        ('.devkit/new-source-docx/6. š,t,ṭ,ṯ.docx', 'šbḥ'),         # 2 stems
        ('.devkit/new-source-docx/1. ʔ, ʕ, b, č.docx', 'bzr'),      # 1 stem
        ('.devkit/new-source-docx/6. š,t,ṭ,ṯ.docx', 'tky'),         # 1 stem - should work?
        ('.devkit/new-source-docx/7. v, w, x, y, z, ž.docx', 'zyr 2'), # 1 stem - should work?
    ]

    for docx_path, verb_root in investigations:
        find_and_analyze_verb(Path(docx_path), verb_root)

    print(f"\n\n{'='*80}")
    print("INVESTIGATION COMPLETE")
    print(f"{'='*80}")

if __name__ == '__main__':
    main()
