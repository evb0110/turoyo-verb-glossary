#!/usr/bin/env python3
import os
import json
from pathlib import Path
from lxml import etree
import zipfile
from docx import Document

DOCX_DIR = ".devkit/new-source-docx"
HTML_FILE = ".devkit/analysis/turoyo_verbs_reference.html"
VALIDATION_REPORT = ".devkit/analysis/html_conversion_validation.json"

XMLNS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}

def extract_footnotes_from_docx(docx_path):
    """Extract footnotes from DOCX, skipping placeholder footnotes."""
    footnotes = {}
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx:
            if 'word/footnotes.xml' in docx.namelist():
                with docx.open('word/footnotes.xml') as f:
                    tree = etree.parse(f)
                    root = tree.getroot()
                    fn_elements = root.findall('.//w:footnote', XMLNS)
                    for fn in fn_elements:
                        fn_id = fn.get(f"{{{XMLNS['w']}}}id")
                        if fn_id and fn_id not in ['-1', '0']:  # Skip placeholder footnotes
                            text_parts = []
                            for t in fn.findall('.//w:t', XMLNS):
                                if t.text:
                                    text_parts.append(t.text)
                            text = ''.join(text_parts).strip()
                            if text:  # Only include non-empty footnotes
                                footnotes[fn_id] = text
    except Exception as e:
        print(f"Error extracting footnotes from {docx_path}: {e}")
    return footnotes

def analyze_docx_files():
    """Get baseline stats from DOCX files."""
    docx_files = sorted([f for f in os.listdir(DOCX_DIR) if f.endswith('.docx')])
    stats = {
        'files': {},
        'totals': {
            'paragraphs': 0,
            'tables': 0,
            'footnotes': 0,
            'files_count': len(docx_files),
        }
    }

    all_footnotes = {}

    for docx_file in docx_files:
        docx_path = os.path.join(DOCX_DIR, docx_file)
        doc = Document(docx_path)
        footnotes = extract_footnotes_from_docx(docx_path)

        file_stats = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'footnotes': len(footnotes),
            'footnotes_content': footnotes,
        }

        stats['files'][docx_file] = file_stats
        stats['totals']['paragraphs'] += len(doc.paragraphs)
        stats['totals']['tables'] += len(doc.tables)
        stats['totals']['footnotes'] += len(footnotes)

        all_footnotes.update(footnotes)

    stats['totals']['unique_footnotes'] = len(all_footnotes)
    stats['all_footnotes'] = all_footnotes

    return stats

def analyze_html_file():
    """Get stats from generated HTML file."""
    stats = {
        'file_size_bytes': 0,
        'file_size_mb': 0,
        'sections': 0,
        'tables': 0,
        'headings': {
            'h1': 0,
            'h2': 0,
            'h3': 0,
        },
        'paragraphs': 0,
        'footnotes_found': 0,
        'footnotes_referenced': [],
    }

    if not os.path.exists(HTML_FILE):
        return stats, "HTML file not found"

    file_size = os.path.getsize(HTML_FILE)
    stats['file_size_bytes'] = file_size
    stats['file_size_mb'] = file_size / 1024 / 1024

    try:
        parser = etree.HTMLParser()
        with open(HTML_FILE, 'r', encoding='utf-8') as f:
            html_content = f.read()

        tree = etree.fromstring(html_content, parser)

        # Count elements
        stats['sections'] = len(tree.findall('.//section[@class="docx-file"]'))
        stats['tables'] = len(tree.findall('.//table'))
        stats['paragraphs'] = len(tree.findall('.//p'))
        stats['headings']['h1'] = len(tree.findall('.//h1'))
        stats['headings']['h2'] = len(tree.findall('.//h2'))
        stats['headings']['h3'] = len(tree.findall('.//h3'))

        # Check for footnotes
        footnote_refs = tree.findall('.//sup[@class="footnote-ref"]')
        stats['footnotes_referenced'] = sorted(
            list(set([el.get('data-footnote-id') for el in footnote_refs if el.get('data-footnote-id')])),
            key=lambda x: int(x) if x.isdigit() else 0
        )
        stats['footnotes_found'] = len(stats['footnotes_referenced'])

        return stats, None
    except Exception as e:
        return stats, str(e)

def generate_validation_report(docx_stats, html_stats, html_error):
    """Generate comprehensive validation report."""
    report = {
        'timestamp': __import__('datetime').datetime.now().isoformat(),
        'source_docx': docx_stats,
        'output_html': html_stats,
        'html_parsing_error': html_error,
        'validation_results': {},
        'warnings': [],
        'errors': [],
    }

    # Validate paragraph count
    # Note: HTML will have more paragraphs because table cells are converted to <p> tags
    # This is expected behavior, not an error
    expected_paras = docx_stats['totals']['paragraphs']
    actual_paras = html_stats['paragraphs']
    expected_tables = docx_stats['totals']['tables']

    # Estimate expected increase: tables create multiple <p> tags per cell
    # This is expected and not a validation failure
    para_increase = actual_paras - expected_paras
    tables_contribution = expected_tables * 2  # Conservative estimate

    report['validation_results']['paragraphs'] = {
        'expected_minimum': expected_paras,  # Original paragraphs
        'actual': actual_paras,
        'increase_from_table_conversion': para_increase,
        'note': 'HTML table cells are converted to <p> tags, causing expected increase',
        'match': actual_paras >= expected_paras  # Should be at least original count
    }

    # Validate table count
    expected_tables = docx_stats['totals']['tables']
    actual_tables = html_stats['tables']
    report['validation_results']['tables'] = {
        'expected': expected_tables,
        'actual': actual_tables,
        'match': expected_tables == actual_tables,
        'message': f"Expected {expected_tables}, got {actual_tables}"
    }

    if expected_tables != actual_tables:
        report['warnings'].append(
            f"Table count mismatch: expected {expected_tables}, got {actual_tables}"
        )

    # Validate footnotes
    expected_footnotes = docx_stats['totals']['unique_footnotes']
    actual_footnotes = html_stats['footnotes_found']
    report['validation_results']['footnotes'] = {
        'expected': expected_footnotes,
        'actual': actual_footnotes,
        'match': expected_footnotes == actual_footnotes,
        'footnote_ids': html_stats['footnotes_referenced'],
        'message': f"Expected {expected_footnotes}, found {actual_footnotes}"
    }

    if expected_footnotes != actual_footnotes:
        report['errors'].append(
            f"Footnote count mismatch: expected {expected_footnotes}, found {actual_footnotes}"
        )

    # Validate file sections
    expected_files = docx_stats['totals']['files_count']
    actual_sections = html_stats['sections']
    report['validation_results']['file_sections'] = {
        'expected': expected_files,
        'actual': actual_sections,
        'match': expected_files == actual_sections,
    }

    if expected_files != actual_sections:
        report['errors'].append(
            f"File section count mismatch: expected {expected_files}, got {actual_sections}"
        )

    # Validate all footnote content is preserved
    html_footnote_ids = set(html_stats['footnotes_referenced'])
    docx_footnote_ids = set(docx_stats['all_footnotes'].keys())

    missing_footnotes = docx_footnote_ids - html_footnote_ids
    if missing_footnotes:
        report['errors'].append(
            f"Missing footnotes in HTML: {sorted(missing_footnotes)}"
        )

    # File size validation
    report['validation_results']['file_size'] = {
        'bytes': html_stats['file_size_bytes'],
        'megabytes': round(html_stats['file_size_mb'], 2),
        'reasonable': html_stats['file_size_mb'] > 1,  # Should be > 1MB with all content
    }

    # Overall result
    has_errors = len(report['errors']) > 0
    report['validation_results']['overall'] = {
        'status': 'FAILED' if has_errors else 'PASSED',
        'errors_count': len(report['errors']),
        'warnings_count': len(report['warnings']),
    }

    return report

def print_report(report):
    """Print validation report to console."""
    print("\n" + "="*70)
    print("HTML CONVERSION VALIDATION REPORT")
    print("="*70)

    overall = report['validation_results']['overall']
    print(f"\n📊 OVERALL STATUS: {overall['status']}")
    print(f"   Errors: {overall['errors_count']}")
    print(f"   Warnings: {overall['warnings_count']}")

    print("\n✓ DOCX Source Statistics:")
    print(f"   Files: {report['source_docx']['totals']['files_count']}")
    print(f"   Paragraphs: {report['source_docx']['totals']['paragraphs']:,}")
    print(f"   Tables: {report['source_docx']['totals']['tables']:,}")
    print(f"   Footnotes: {report['source_docx']['totals']['unique_footnotes']}")

    print("\n📄 HTML Output Statistics:")
    print(f"   File size: {report['output_html']['file_size_mb']:.2f} MB")
    print(f"   Sections: {report['output_html']['sections']}")
    print(f"   Paragraphs: {report['output_html']['paragraphs']:,}")
    print(f"   Tables: {report['output_html']['tables']:,}")
    print(f"   Footnotes found: {report['output_html']['footnotes_found']}")

    print("\n🔍 Validation Results:")
    for key, result in report['validation_results'].items():
        if key != 'overall':
            status = "✓" if result.get('match', result.get('reasonable', False)) else "✗"
            print(f"   {status} {key}: {result.get('message', str(result))}")

    if report['warnings']:
        print("\n⚠️  Warnings:")
        for warning in report['warnings']:
            print(f"   - {warning}")

    if report['errors']:
        print("\n❌ Errors:")
        for error in report['errors']:
            print(f"   - {error}")

    print("\n" + "="*70)

if __name__ == '__main__':
    print("Analyzing source DOCX files...")
    docx_stats = analyze_docx_files()

    print("Analyzing generated HTML file...")
    html_stats, html_error = analyze_html_file()

    print("Generating validation report...")
    report = generate_validation_report(docx_stats, html_stats, html_error)

    # Save report
    os.makedirs(os.path.dirname(VALIDATION_REPORT), exist_ok=True)
    with open(VALIDATION_REPORT, 'w', encoding='utf-8') as f:
        json.dump(report, f, indent=2, ensure_ascii=False)

    # Print report
    print_report(report)

    print(f"\n📋 Full report saved to: {VALIDATION_REPORT}")
