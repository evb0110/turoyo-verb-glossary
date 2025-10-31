#!/usr/bin/env python3
"""
Investigate 20 sampled verbs with NULL etymology in DOCX source files.
Version 2: Improved search that looks for ANY paragraph containing the root.
"""

import json
import re
from pathlib import Path
from docx import Document

DOCX_DIR = Path(".devkit/new-source-docx")
JSON_DIR = Path("server/assets/verbs")

DOCX_FILES = {
    "1. ʔ, ʕ, b, č.docx": ['ʔ', 'ʕ', 'b', 'č'],
    "2. d, f, g, ġ, ǧ.docx": ['d', 'f', 'g', 'ġ', 'ǧ'],
    "3. h,ḥ,k.docx": ['h', 'ḥ', 'k'],
    "4. l,m,n,p.docx": ['l', 'm', 'n', 'p'],
    "5. q,r,s,ṣ.docx": ['q', 'r', 's', 'ṣ'],
    "6. š,t,ṭ,ṯ.docx": ['š', 't', 'ṭ', 'ṯ'],
    "7. v, w, x, y, z, ž.docx": ['v', 'w', 'x', 'y', 'z', 'ž']
}

SAMPLE_VERBS = [
    'bʕy 2', 'dng', 'fqṣ', 'grḏ', 'hmz', 'klpt', 'lhṭ', 'mʕr', 'ngl 2', 'pšṭ',
    'qḥl', 'rǧʕ', 'slpx', 'twy', 'xnxl', 'zbʕ', 'čfx', 'ġrf 2', 'šrqm', 'žġl 1'
]


def get_docx_file_for_verb(verb_root):
    """Find which DOCX file contains this verb based on first letter."""
    first_letter = verb_root[0]
    for docx_file, letters in DOCX_FILES.items():
        if first_letter in letters:
            return DOCX_DIR / docx_file
    return None


def find_verb_in_docx(docx_path, verb_root):
    """Find verb entry in DOCX by searching all paragraphs."""
    doc = Document(docx_path)

    clean_root = verb_root.split()[0]

    candidates = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if not text:
            continue

        starts_with_root = text.startswith(clean_root + ' ') or text.startswith(clean_root + '(')

        if starts_with_root:
            is_numbered = any(text.startswith(f"{clean_root} {n}") for n in ['1', '2', '3'])
            has_paren = '(' in text and ')' in text

            if verb_root.endswith(' 1') and text.startswith(f"{clean_root} 1"):
                candidates.append((i, text, 1))
            elif verb_root.endswith(' 2') and text.startswith(f"{clean_root} 2"):
                candidates.append((i, text, 2))
            elif verb_root.endswith(' 3') and text.startswith(f"{clean_root} 3"):
                candidates.append((i, text, 3))
            elif ' ' not in verb_root and not is_numbered:
                candidates.append((i, text, 0))

    if not candidates:
        return {
            'found': False,
            'root_text': '',
            'following_text': '',
            'para_index': -1
        }

    candidates.sort(key=lambda x: x[2], reverse=True)
    para_idx, root_text, _ = candidates[0]

    following_paragraphs = []
    j = para_idx + 1
    while j < len(doc.paragraphs) and len(following_paragraphs) < 5:
        next_para = doc.paragraphs[j]
        next_text = next_para.text.strip()

        if next_text and next_text[0] in 'ʔʕbčdfgġǧhḥklmnpqrsṣšṭṯtxyzžʕḥṣṭ':
            other_root_start = next_text.split()[0] if ' ' in next_text else next_text.split('(')[0]
            if other_root_start != clean_root:
                break

        if next_text:
            following_paragraphs.append(next_text)
        j += 1

    return {
        'found': True,
        'root_text': root_text,
        'following_text': '\n\n'.join(following_paragraphs),
        'para_index': para_idx
    }


def categorize_etymology(docx_result, verb_root):
    """Categorize etymology status with detailed analysis."""
    if not docx_result['found']:
        return 'X', "Could not find verb in DOCX", None

    root_text = docx_result['root_text']
    following = docx_result['following_text']

    root_lower = root_text.lower()
    following_lower = following.lower()

    if '(unknown)' in root_lower:
        return 'B', "Etymology correctly marked as 'unknown' in source", None

    if '(uncertain' in root_lower:
        return 'B', "Etymology marked as 'uncertain' in source", None

    etymology_markers_root = [
        '<', 'cf.', 'denom', 'turkish', 'arab', 'syr', 'akkad', 'hebrew',
        'kurmanji', 'see ', 'persian', 'armenian', 'from '
    ]

    for marker in etymology_markers_root:
        if marker in root_lower:
            etymol_text = root_text[root_text.find('('):] if '(' in root_text else root_text
            return 'A', f"Etymology present in root paragraph but parser failed (found '{marker}')", etymol_text

    etymology_markers_following = [
        'see ', 'cf.', 'turkish', 'persian', 'armenian', 'from ', 'arabic root'
    ]

    for marker in etymology_markers_following:
        if marker in following_lower:
            return 'A', f"Etymology present in following text but parser failed (found '{marker}')", following[:200]

    if not following:
        return 'C', "No text after root in DOCX (genuinely missing)", None

    return 'C', "No clear etymology information in source (genuinely missing)", None


def main():
    print("=" * 80)
    print("NULL ETYMOLOGY INVESTIGATION - 20 VERB SAMPLE (V2)")
    print("=" * 80)
    print()

    results = []

    for verb in SAMPLE_VERBS:
        docx_file = get_docx_file_for_verb(verb)
        if not docx_file:
            print(f"ERROR: Could not determine DOCX file for {verb}")
            continue

        if not docx_file.exists():
            print(f"ERROR: DOCX file not found: {docx_file}")
            continue

        print(f"Investigating: {verb}")
        print(f"DOCX file: {docx_file.name}")

        docx_result = find_verb_in_docx(docx_file, verb)
        category, reason, etymology_snippet = categorize_etymology(docx_result, verb)

        result = {
            'verb': verb,
            'docx_file': docx_file.name,
            'found_in_docx': docx_result['found'],
            'root_text': docx_result['root_text'],
            'following_text': docx_result['following_text'][:500],
            'category': category,
            'reason': reason,
            'etymology_snippet': etymology_snippet
        }
        results.append(result)

        print(f"  Category: {category} - {reason}")
        if docx_result['found']:
            print(f"  Root text: {docx_result['root_text'][:150]}")
            if etymology_snippet:
                print(f"  Etymology snippet: {etymology_snippet[:150]}")
        print()

    category_counts = {'A': 0, 'B': 0, 'C': 0, 'D': 0, 'X': 0}
    for r in results:
        category_counts[r['category']] += 1

    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)
    print(f"Total investigated: {len(results)}")
    print(f"  Category A (Parser failures - FIXABLE): {category_counts['A']} ({category_counts['A']/len(results)*100:.1f}%)")
    print(f"  Category B (Correctly 'unknown'): {category_counts['B']} ({category_counts['B']/len(results)*100:.1f}%)")
    print(f"  Category C (Genuinely missing): {category_counts['C']} ({category_counts['C']/len(results)*100:.1f}%)")
    print(f"  Category D (Non-standard format): {category_counts['D']} ({category_counts['D']/len(results)*100:.1f}%)")
    print(f"  Category X (Not found in DOCX): {category_counts['X']} ({category_counts['X']/len(results)*100:.1f}%)")
    print()

    print("Extrapolation to all 196 NULL etymology verbs:")
    for cat in ['A', 'B', 'C', 'D']:
        if category_counts[cat] > 0:
            estimated = int(196 * category_counts[cat] / len(results))
            print(f"  Category {cat}: ~{estimated} verbs")
    print()

    print("\nCategory A cases (FIXABLE):")
    for r in results:
        if r['category'] == 'A':
            print(f"  - {r['verb']}: {r['reason']}")
            print(f"    Root: {r['root_text'][:100]}")

    with open('.devkit/analysis/null_etymology_investigation_v2.json', 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print("\nRaw results saved to: .devkit/analysis/null_etymology_investigation_v2.json")


if __name__ == '__main__':
    main()
