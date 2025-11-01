#!/usr/bin/env python3
"""Debug what's happening with table cell parsing"""

import sys
sys.path.insert(0, 'parser')

from parse_docx_production import FixedDocxParser
from docx import Document
import re

parser = FixedDocxParser()
doc = Document('.devkit/new-source-docx/3. h,ḥ,k.docx')

root_pattern = re.compile(r'^(hyw\s+1)\s*[<(]', re.UNICODE)
in_hyw1 = False
table_count = 0

for el in doc.element.body:
    tag = el.tag.split('}')[1] if '}' in el.tag else el.tag

    if tag == 'p':
        for para in doc.paragraphs:
            if para._element == el:
                text = para.text.strip()
                if root_pattern.match(text):
                    in_hyw1 = True
                    print("✓ Found hyw 1")
                elif in_hyw1 and re.match(r'^([ʔʕbčdfgġǧhḥklmnpqrsṣštṭvwxyzžḏṯẓāēīūə]+(?:\s+\d+)?)\s*[<(]', text, re.UNICODE):
                    print("✓ Found next root, stopping")
                    break
                break

    elif tag == 'tbl' and in_hyw1:
        table_count += 1
        for table in doc.tables:
            if table._element == el:
                if table.rows and len(table.rows[0].cells) >= 2:
                    conj_type = table.rows[0].cells[0].text.strip()
                    cell = table.rows[0].cells[1]

                    # Get raw text
                    raw_text = cell.text.strip()
                    raw_length = len(raw_text)

                    # Parse with parser
                    examples = parser.parse_table_cell(cell)

                    # Reconstruct parsed text
                    parsed_text = []
                    for ex in examples:
                        if ex.get('turoyo'):
                            parsed_text.append(ex['turoyo'])
                        if ex.get('translations'):
                            parsed_text.extend(ex['translations'])
                    parsed_combined = ' '.join(parsed_text)
                    parsed_length = len(parsed_combined)

                    loss = raw_length - parsed_length
                    loss_pct = (loss / raw_length * 100) if raw_length > 0 else 0

                    print(f"\n{'='*70}")
                    print(f"TABLE {table_count}: {conj_type}")
                    print(f"{'='*70}")
                    print(f"Raw text ({raw_length} chars):")
                    print(f"  {raw_text[:200]}...")
                    print(f"\nParsed ({len(examples)} examples, {parsed_length} chars):")
                    for i, ex in enumerate(examples, 1):
                        t = ex.get('turoyo', '')[:80]
                        trans = ', '.join(ex.get('translations', [])[:2])[:80]
                        print(f"  {i}. T: {t}")
                        if trans:
                            print(f"     Translations: {trans}")
                    print(f"\n⚠️  DATA LOSS: {loss} chars ({loss_pct:.1f}%)")

                    if table_count >= 3:  # Just show first 3 tables
                        print("\n[Stopping after 3 tables]")
                        break
                break

        if table_count >= 3:
            break
